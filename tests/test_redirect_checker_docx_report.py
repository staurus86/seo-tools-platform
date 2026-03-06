import shutil
import sys
import types
import unittest
from pathlib import Path

from docx import Document

# Make DOCX generator import independent from runtime-only settings dependency.
if "app.config" not in sys.modules:
    fake_config = types.ModuleType("app.config")
    fake_config.settings = types.SimpleNamespace(REPORTS_DIR=".")
    sys.modules["app.config"] = fake_config

from app.reports.docx_generator import DOCXGenerator


class RedirectCheckerDocxReportTests(unittest.TestCase):
    def test_redirect_checker_docx_contains_plan_and_tz_sections(self):
        data = {
            "url": "https://example.com/",
            "results": {
                "checked_url": "https://example.com/",
                "selected_user_agent": {
                    "key": "googlebot_desktop",
                    "label": "Googlebot Desktop",
                },
                "summary": {
                    "total_scenarios": 13,
                    "passed": 8,
                    "warnings": 3,
                    "errors": 2,
                    "quality_score": 71,
                    "quality_grade": "C",
                    "duration_ms": 1240,
                },
                "scenarios": [
                    {
                        "id": 1,
                        "key": "http_to_https",
                        "title": "HTTP -> HTTPS",
                        "what_checked": "Редирект с http:// на https://",
                        "status": "error",
                        "expected": "301/308 на HTTPS",
                        "actual": "200 на http://example.com/",
                        "recommendation": "Настройте принудительный HTTPS.",
                        "test_url": "http://example.com/",
                        "response_codes": [200],
                        "duration_ms": 120,
                        "final_url": "http://example.com/",
                        "hops": 0,
                        "chain": [{"url": "http://example.com/", "status_code": 200, "location": ""}],
                        "error": "",
                    },
                    {
                        "id": 8,
                        "key": "canonical_tag",
                        "title": "Canonical тег",
                        "what_checked": "Проверка canonical",
                        "status": "warning",
                        "expected": "Canonical присутствует",
                        "actual": "canonical не найден",
                        "recommendation": "Добавьте canonical в <head>.",
                        "test_url": "https://example.com/",
                        "response_codes": [200],
                        "duration_ms": 45,
                        "final_url": "https://example.com/",
                        "hops": 0,
                        "chain": [{"url": "https://example.com/", "status_code": 200, "location": ""}],
                        "error": "",
                    },
                    {
                        "id": 9,
                        "key": "missing_404",
                        "title": "404-страницы",
                        "what_checked": "Проверка кода ответа для несуществующей страницы",
                        "status": "error",
                        "expected": "HTTP 404 для несуществующей страницы",
                        "actual": "200 на несуществующий URL",
                        "recommendation": "Настройте корректный 404.",
                        "test_url": "https://example.com/missing-page",
                        "response_codes": [200],
                        "duration_ms": 80,
                        "final_url": "https://example.com/missing-page",
                        "hops": 0,
                        "chain": [{"url": "https://example.com/missing-page", "status_code": 200, "location": ""}],
                        "error": "",
                    },
                    {
                        "id": 17,
                        "key": "soft_404_detection",
                        "title": "Soft-404 detection",
                        "what_checked": "Проверка soft-404 после редиректа",
                        "status": "warning",
                        "expected": "404/410 вместо soft-404",
                        "actual": "200 и контент похож на валидную страницу",
                        "recommendation": "Проверьте soft-404.",
                        "test_url": "https://example.com/missing-page",
                        "response_codes": [200],
                        "duration_ms": 60,
                        "final_url": "https://example.com/",
                        "hops": 0,
                        "chain": [{"url": "https://example.com/missing-page", "status_code": 200, "location": "https://example.com/"}],
                        "error": "",
                    },
                ],
                "recommendations": [
                    "Настройте принудительный HTTPS.",
                    "Добавьте canonical в <head>.",
                ],
                "checked_at": "2026-02-25T10:00:00Z",
            },
        }

        temp_dir = Path("tests") / ".tmp_redirect_docx"
        if temp_dir.exists():
            shutil.rmtree(temp_dir)
        temp_dir.mkdir(parents=True, exist_ok=True)
        try:
            generator = DOCXGenerator()
            generator.reports_dir = str(temp_dir)
            report_path = generator.generate_redirect_checker_report("redirect-docx-test", data)

            doc = Document(report_path)
            text = "\n".join(p.text for p in doc.paragraphs)
            table_text = "\n".join(
                cell.text
                for table in doc.tables
                for row in table.rows
                for cell in row.cells
            )
            self.assertIn("Redirect Checker Report", text)
            self.assertIn("3. Нарушения и разбор", text)
            self.assertIn("4. План действий при нарушениях", text)
            self.assertIn("5. 404 Risks", text)
            self.assertIn("7. Краткие ТЗ на исправление", text)
            self.assertIn("Приложение. Полная структура результата", text)
            self.assertIn("HTTP -> HTTPS", text)
            self.assertIn("Canonical тег", text)
            self.assertIn("404-страницы", text)
            self.assertIn("Soft-404 detection", text)
            self.assertIn("Критерий приемки:", text)
            self.assertIn("Время сценария: 120 ms", text)
            self.assertIn("Время", table_text)
            self.assertIn("120 ms", table_text)
        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)


if __name__ == "__main__":
    unittest.main()
