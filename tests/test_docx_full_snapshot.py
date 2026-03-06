import shutil
import sys
import types
import unittest
from io import BytesIO
from pathlib import Path

from docx import Document

# Make DOCX generator import independent from runtime-only settings dependency.
if "app.config" not in sys.modules:
    fake_config = types.ModuleType("app.config")
    fake_config.settings = types.SimpleNamespace(REPORTS_DIR=".")
    sys.modules["app.config"] = fake_config

from app.reports.docx_generator import DOCXGenerator
from app.tools.llmCrawler.report_docx import build_docx_v2


def _doc_text(doc: Document) -> str:
    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text:
                        parts.append(paragraph.text)
    return "\n".join(parts)


class DocxFullSnapshotTests(unittest.TestCase):
    def test_shared_generator_appends_full_result_snapshot(self):
        data = {
            "url": "https://example.com/sitemap.xml",
            "results": {
                "valid": True,
                "status_code": 200,
                "issues": [],
                "recommendations": [],
                "ui_only_block": {
                    "rollout_state": "beta",
                    "coverage_note": "Coverage cap reached in UI card",
                },
            },
        }

        temp_dir = Path("tests") / ".tmp_docx_snapshot"
        if temp_dir.exists():
            shutil.rmtree(temp_dir)
        temp_dir.mkdir(parents=True, exist_ok=True)
        try:
            generator = DOCXGenerator()
            generator.reports_dir = str(temp_dir)
            report_path = generator.generate_sitemap_report("snapshot-test", data)

            doc = Document(report_path)
            text = _doc_text(doc)
            self.assertIn("Приложение. Полная структура результата", text)
            self.assertIn("Coverage cap reached in UI card", text)
            self.assertIn("Rollout state", text)
        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)

    def test_llm_docx_appends_full_result_snapshot(self):
        payload = build_docx_v2(
            {
                "result": {
                    "final_url": "https://example.com/page",
                    "score": {"total": 72, "top_issues": ["Missing schema"]},
                    "ui_only_snapshot": {
                        "render_mode": "streaming",
                        "panel_note": "Visible in UI only",
                    },
                }
            },
            "llm-snapshot-test",
            wow_enabled=False,
        )

        doc = Document(BytesIO(payload.getvalue()))
        text = _doc_text(doc)
        self.assertIn("Appendix: Full Result Snapshot", text)
        self.assertIn("Visible in UI only", text)
        self.assertIn("Render mode", text)


if __name__ == "__main__":
    unittest.main()
