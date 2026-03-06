"""
Word Report Generator
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from typing import Dict, Any, List
from datetime import datetime
import os
import json
import tempfile
from urllib.parse import urljoin

import requests

from app.config import settings


class DOCXGenerator:
    """Генератор Word-отчетов."""
    
    def __init__(self):
        self.reports_dir = settings.REPORTS_DIR
        os.makedirs(self.reports_dir, exist_ok=True)

    def _configure_document(self, doc: Document, report_title: str = "", subject: str = "") -> None:
        """Apply shared typography and spacing so reports look consistent across tools."""
        for section in doc.sections:
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)
            section.top_margin = Inches(0.65)
            section.bottom_margin = Inches(0.65)

        normal_style = doc.styles["Normal"]
        normal_style.font.name = "Calibri"
        normal_style.font.size = Pt(10.5)
        normal_style.paragraph_format.space_after = Pt(6)
        normal_style.paragraph_format.line_spacing = 1.12

        for style_name, size, color in (
            ("Title", 28, RGBColor(15, 76, 129)),
            ("Heading 1", 16, RGBColor(15, 76, 129)),
            ("Heading 2", 13, RGBColor(14, 116, 144)),
            ("Heading 3", 11, RGBColor(30, 41, 59)),
        ):
            if style_name not in doc.styles:
                continue
            style = doc.styles[style_name]
            style.font.name = "Calibri"
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = color
            style.paragraph_format.space_before = Pt(10 if style_name != "Heading 3" else 8)
            style.paragraph_format.space_after = Pt(5)
            style.paragraph_format.keep_with_next = True

        doc.core_properties.author = "SEO Tools Platform"
        doc.core_properties.company = "SEO Tools Platform"
        if report_title:
            doc.core_properties.title = self._fix_text(report_title)
        if subject:
            doc.core_properties.subject = self._fix_text(subject)
    
    def _add_heading(self, doc, text: str, level: int = 1):
        """Добавляет заголовок."""
        heading = doc.add_heading(self._fix_text(text), level=level)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return heading
    
    def _add_table(self, doc, headers: List[str], rows: List[List[Any]]):
        """Добавляет таблицу."""
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        # Header row with brand color background
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            cell = hdr_cells[i]
            cell.text = self._fix_text(str(header))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            # Brand color header background
            shading = cell._element.get_or_add_tcPr()
            shd_el = shading.makeelement(qn('w:shd'), {
                qn('w:val'): 'clear',
                qn('w:color'): 'auto',
                qn('w:fill'): '0F4C81'
            })
            shading.append(shd_el)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)

        # Data rows
        for row_index, row_data in enumerate(rows):
            row_cells = table.add_row().cells
            for i, value in enumerate(row_data):
                row_cells[i].text = self._fix_text(str(value))
                row_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                if row_index % 2 == 0:
                    tc_pr = row_cells[i]._element.get_or_add_tcPr()
                    shd_el = tc_pr.makeelement(qn('w:shd'), {
                        qn('w:val'): 'clear',
                        qn('w:color'): 'auto',
                        qn('w:fill'): 'F8FBFF'
                    })
                    tc_pr.append(shd_el)

        return table

    def _fix_text(self, text: str) -> str:
        """Repair common mojibake from UTF-8<->cp1251 mixups."""
        if not isinstance(text, str) or not text:
            return text

        weird_chars = set("Р‚РѓвЂљС“вЂћвЂ¦вЂ вЂЎв‚¬вЂ°Р‰вЂ№РЉРЊР‹РЏС’вЂвЂ™вЂњвЂќвЂўвЂ“вЂ”в„ўС™вЂєСљСњС›Сџ")

        def _quality(value: str) -> int:
            cyr = sum(1 for ch in value if ("А" <= ch <= "я") or ch in ("Ё", "ё"))
            weird = sum(1 for ch in value if ch in weird_chars)
            c1 = sum(1 for ch in value if 0x80 <= ord(ch) <= 0x9F)
            return cyr - (weird * 4) - (value.count("пїЅ") * 8) - (c1 * 6)

        def _looks_mojibake(value: str) -> bool:
            if any(marker in value for marker in ("вЂ", "в„", "Ѓ", "вЂљ", "в‚¬", "в„ў", "љ", "њ", "ћ", "џ")):
                return True
            if any(0x80 <= ord(ch) <= 0x9F for ch in value):
                return True
            letters = [ch for ch in value if ch.isalpha()]
            if not letters:
                return False
            rs_count = sum(1 for ch in letters if ch in ("Р ", "РЎ"))
            return (rs_count / len(letters)) > 0.28

        def _cp1251_byte_for_char(ch: str) -> int | None:
            code = ord(ch)
            if code <= 0xFF:
                return code
            if ch == "Ё":
                return 0xA8
            if ch == "ё":
                return 0xB8
            if "А" <= ch <= "я":
                return code - 0x350
            try:
                raw = ch.encode("cp1251", errors="strict")
                if len(raw) == 1:
                    return raw[0]
            except Exception:
                return None
            return None

        def _decode_mixed_cp1251_utf8(value: str) -> str | None:
            raw = bytearray()
            for ch in value:
                b = _cp1251_byte_for_char(ch)
                if b is None:
                    return None
                raw.append(b)
            try:
                return raw.decode("utf-8", errors="strict")
            except Exception:
                return None

        repaired = text
        for _ in range(3):
            candidates = []
            converted = _decode_mixed_cp1251_utf8(repaired)
            if converted and converted != repaired:
                candidates.append(converted)
            try:
                latin1 = repaired.encode("latin1", errors="strict").decode("utf-8", errors="strict")
                if latin1 != repaired:
                    candidates.append(latin1)
            except Exception:
                pass
            if not candidates:
                break
            best = max(candidates, key=_quality)
            if _looks_mojibake(repaired) or _quality(best) > (_quality(repaired) + 1):
                repaired = best
                continue
            break
        return repaired

    def _format_engine_label(self, engine: Any) -> str:
        value = str(engine or "legacy").lower()
        if value == "legacy":
            return "базовый"
        if value == "legacy-fallback":
            return "базовый (fallback)"
        return str(engine or "legacy")

    def _format_profile_label(self, profile: Any, is_mobile: Any = False) -> str:
        value = str(profile or ("mobile" if is_mobile else "desktop")).lower()
        if value == "mobile":
            return "мобильный"
        if value == "desktop":
            return "десктоп"
        return str(profile or "-")

    def _format_policy_profile(self, value: Any, profile_type: str) -> str:
        v = str(value or "").lower()
        if not v:
            return "н/д"
        if profile_type == "retry":
            if v == "standard":
                return "стандартный"
            if v == "aggressive":
                return "агрессивный"
            if v == "strict":
                return "строгий"
        if profile_type == "criticality":
            if v == "balanced":
                return "сбалансированный"
            if v == "strict":
                return "строгий"
            if v == "aggressive":
                return "агрессивный"
        if profile_type == "sla":
            if v == "standard":
                return "стандартный"
            if v == "strict":
                return "строгий"
        return str(value)

    def _normalize_document_text(self, doc: Document) -> None:
        """Normalize paragraph/table text before saving DOCX."""
        def _fix_paragraph(paragraph) -> None:
            if paragraph.runs:
                for run in paragraph.runs:
                    run.text = self._fix_text(run.text)
            elif paragraph.text:
                paragraph.text = self._fix_text(paragraph.text)

        def _fix_table(table) -> None:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        _fix_paragraph(p)
                    for nested in cell.tables:
                        _fix_table(nested)

        for p in doc.paragraphs:
            _fix_paragraph(p)
        for t in doc.tables:
            _fix_table(t)

    def _save_document(self, doc: Document, filepath: str) -> None:
        """Normalize text and save DOCX."""
        self._normalize_document_text(doc)
        doc.save(filepath)

    def _add_cover_page(self, doc: Document, title: str, subtitle: str, url: str, generated_at: str) -> None:
        """Creates a branded cover page."""
        self._configure_document(doc, title, subtitle)
        # Push content down
        doc.add_paragraph()

        # Title in brand color, Heading 0, center-aligned, 28pt
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title_para.runs:
            run.font.color.rgb = RGBColor(15, 76, 129)  # #0F4C81
            run.font.size = Pt(28)

        # Subtitle — smaller, gray, center-aligned
        subtitle_para = doc.add_paragraph(subtitle)
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if subtitle_para.runs:
            subtitle_para.runs[0].font.size = Pt(14)
            subtitle_para.runs[0].font.color.rgb = RGBColor(107, 114, 128)  # gray

        # URL in brand-2 color
        if url:
            url_para = doc.add_paragraph(url)
            url_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if url_para.runs:
                url_para.runs[0].font.color.rgb = RGBColor(14, 116, 144)  # #0E7490

        # Timestamp in gray
        ts_para = doc.add_paragraph(generated_at)
        ts_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if ts_para.runs:
            ts_para.runs[0].font.color.rgb = RGBColor(107, 114, 128)

        # Logo image
        logo_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'static', 'icon.png')
        if os.path.exists(logo_path):
            try:
                logo_para = doc.add_paragraph()
                logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = logo_para.add_run()
                run.add_picture(logo_path, width=Inches(1))
            except Exception:
                pass

        # Page break
        doc.add_page_break()

    def _is_scalar_value(self, value: Any) -> bool:
        return value is None or isinstance(value, (str, int, float, bool))

    def _format_appendix_label(self, value: Any) -> str:
        raw = str(value or "").replace("_", " ").replace(".", " / ").strip()
        if not raw:
            return "Поле"
        return self._fix_text(raw[:1].upper() + raw[1:])

    def _format_appendix_value(self, value: Any, limit: int = 1600) -> str:
        if value is None:
            return "—"
        if isinstance(value, bool):
            return "Да" if value else "Нет"
        if isinstance(value, float):
            return f"{value:.3f}".rstrip("0").rstrip(".")
        if isinstance(value, (dict, list)):
            rendered = json.dumps(value, ensure_ascii=False)
        else:
            rendered = str(value)
        rendered = self._fix_text(rendered)
        if len(rendered) > limit:
            return rendered[:limit] + " ..."
        return rendered

    def _add_appendix_tree(
        self,
        doc: Document,
        value: Any,
        level: int = 2,
        depth: int = 0,
        max_depth: int = 6,
        max_items: int = 200,
    ) -> None:
        if depth >= max_depth:
            doc.add_paragraph("Достигнута предельная глубина вложенности. Остальные данные скрыты для читаемости.")
            return

        if self._is_scalar_value(value):
            doc.add_paragraph(self._format_appendix_value(value))
            return

        if isinstance(value, dict):
            if not value:
                doc.add_paragraph("Раздел пуст.")
                return

            scalar_rows = []
            nested_items = []
            for key, item in value.items():
                if self._is_scalar_value(item):
                    scalar_rows.append([self._format_appendix_label(key), self._format_appendix_value(item)])
                else:
                    nested_items.append((key, item))

            if scalar_rows:
                self._add_table(doc, ["Поле", "Значение"], scalar_rows[:max_items])
                if len(scalar_rows) > max_items:
                    doc.add_paragraph(f"Показаны первые {max_items} строк из {len(scalar_rows)}.")

            for key, item in nested_items[:max_items]:
                self._add_heading(doc, self._format_appendix_label(key), level=min(level, 3))
                self._add_appendix_tree(
                    doc,
                    item,
                    level=min(level + 1, 4),
                    depth=depth + 1,
                    max_depth=max_depth,
                    max_items=max_items,
                )
            if len(nested_items) > max_items:
                doc.add_paragraph(f"Показаны первые {max_items} вложенных блоков из {len(nested_items)}.")
            return

        if isinstance(value, list):
            if not value:
                doc.add_paragraph("Список пуст.")
                return

            shown = value[:max_items]
            if all(self._is_scalar_value(item) for item in shown):
                for item in shown:
                    doc.add_paragraph(self._format_appendix_value(item), style='List Bullet')
            elif all(isinstance(item, dict) for item in shown):
                scalar_keys: List[str] = []
                for item in shown:
                    for key, sub_value in item.items():
                        if self._is_scalar_value(sub_value) and key not in scalar_keys:
                            scalar_keys.append(key)
                scalar_keys = scalar_keys[:6]
                if scalar_keys:
                    rows = []
                    for item in shown:
                        rows.append([self._format_appendix_value(item.get(key), limit=300) for key in scalar_keys])
                    self._add_table(doc, [self._format_appendix_label(key) for key in scalar_keys], rows)
                for idx, item in enumerate(shown, start=1):
                    nested = {k: v for k, v in item.items() if not self._is_scalar_value(v)}
                    if not nested:
                        continue
                    self._add_heading(doc, f"Элемент {idx}", level=min(level, 4))
                    self._add_appendix_tree(
                        doc,
                        nested,
                        level=min(level + 1, 4),
                        depth=depth + 1,
                        max_depth=max_depth,
                        max_items=max_items,
                    )
            else:
                for idx, item in enumerate(shown, start=1):
                    self._add_heading(doc, f"Элемент {idx}", level=min(level, 4))
                    self._add_appendix_tree(
                        doc,
                        item,
                        level=min(level + 1, 4),
                        depth=depth + 1,
                        max_depth=max_depth,
                        max_items=max_items,
                    )
            if len(value) > max_items:
                doc.add_paragraph(f"Показаны первые {max_items} элементов из {len(value)}.")
            return

        doc.add_paragraph(self._format_appendix_value(value))

    def _add_full_result_appendix(self, doc: Document, data: Dict[str, Any], title: str = "Приложение. Полная структура результата") -> None:
        payload = data if isinstance(data, dict) else {"value": data}
        if not payload:
            return
        doc.add_page_break()
        self._add_heading(doc, title, level=1)
        doc.add_paragraph(
            "Этот раздел сохраняет структурированный снимок результата, чтобы DOCX не терял поля, "
            "которые пользователь видит на странице отчета."
        )
        self._add_appendix_tree(doc, payload, level=2, depth=0)

    def _add_header_footer(self, doc: Document, report_title: str) -> None:
        """Add header and footer to the document."""
        # Header
        section = doc.sections[0]
        header = section.header
        if not header.paragraphs:
            header_para = header.add_paragraph()
        else:
            header_para = header.paragraphs[0]
        header_para.text = f"SEO Tools Platform | {report_title}"
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if header_para.runs:
            header_para.runs[0].font.size = Pt(8)
            header_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)

        # Footer
        footer = section.footer
        if not footer.paragraphs:
            footer_para = footer.add_paragraph()
        else:
            footer_para = footer.paragraphs[0]
        footer_para.text = "SEO Tools Platform"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if footer_para.runs:
            footer_para.runs[0].font.size = Pt(8)
            footer_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    def _add_severity_paragraph(self, doc: Document, text: str, severity: str) -> None:
        """Add a color-coded paragraph based on severity."""
        color_map = {
            'critical': RGBColor(185, 28, 28),   # #B91C1C
            'error':    RGBColor(185, 28, 28),   # #B91C1C
            'warning':  RGBColor(194, 65, 12),   # #C2410C
            'ok':       RGBColor(21, 128, 61),   # #15803D
            'success':  RGBColor(21, 128, 61),   # #15803D
            'info':     RGBColor(21, 128, 61),   # #15803D
        }
        color = color_map.get(str(severity).lower(), RGBColor(15, 23, 42))  # default #0F172A
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.color.rgb = color

    def generate_site_analyze_report(self, task_id: str, data: Dict[str, Any]) -> str:
        """Генерирует клиентский DOCX-отчет общего анализа сайта."""
        doc = Document()
        self._add_cover_page(doc, 'SEO Analysis Report', 'Общий SEO-анализ', data.get('url', ''), datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        self._add_header_footer(doc, 'SEO Analysis Report')

        doc.add_paragraph(f"URL: {data.get('url', 'н/д')}")
        doc.add_paragraph(f"Проверено страниц: {data.get('pages_analyzed', 0)}")
        doc.add_paragraph(f"Завершено: {data.get('completed_at', 'н/д')}")
        doc.add_paragraph(
            "Отчет фиксирует текущее техническое состояние сайта и ключевые SEO-риски, "
            "чтобы определить приоритетные доработки и снизить риски потери органического трафика."
        )

        self._add_heading(doc, 'Ключевые результаты', level=1)
        results = data.get('results', {})
        headers = ['Показатель', 'Значение', 'Статус']
        rows = [
            ['Всего страниц', results.get('total_pages', 0), 'ОК'],
            ['Статус анализа', results.get('status', 'н/д'), 'ОК'],
            ['Сводка', results.get('summary', 'н/д'), 'ОК']
        ]
        self._add_table(doc, headers, rows)

        recs = results.get("recommendations", []) or data.get("recommendations", [])
        if recs:
            doc.add_page_break()
            self._add_heading(doc, 'Рекомендации', level=1)
            for rec in recs[:30]:
                doc.add_paragraph(str(rec), style='List Bullet')

        self._add_full_result_appendix(doc, data)
        filepath = os.path.join(self.reports_dir, f"{task_id}.docx")
        self._save_document(doc, filepath)
        return filepath

    def generate_robots_report(self, task_id: str, data: Dict[str, Any]) -> str:
        """Generate full robots.txt DOCX report with complete result coverage."""
        doc = Document()

        url = data.get('url', 'н/д')
        results = data.get('results', {}) or {}
        generated_at = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

        def yes_no(value: Any) -> str:
            return 'Да' if bool(value) else 'Нет'

        self._add_cover_page(doc, 'Robots.txt Report', 'Проверка Robots.txt', data.get('url', ''), generated_at)
        self._add_header_footer(doc, 'Robots.txt Report')

        doc.add_paragraph(f"URL: {url}")
        doc.add_paragraph(f"Сформирован: {generated_at}")

        self._add_heading(doc, '1. Ключевая сводка', level=1)
        summary_rows = [
            ['robots.txt найден', yes_no(results.get('robots_txt_found'))],
            ['HTTP-статус', str(results.get('status_code', 'н/д'))],
            ['Оценка качества', str(results.get('quality_score', 'н/д'))],
            ['Грейд качества', str(results.get('quality_grade', 'н/д'))],
            ['Готов к продакшн', yes_no(results.get('production_ready'))],
            ['Быстрый статус', str(results.get('quick_status', 'н/д'))],
            ['Размер файла (байт)', str(results.get('content_length', 0))],
            ['Количество строк', str(results.get('lines_count', 0))],
            ['User-agent групп', str(results.get('user_agents', 0))],
            ['Правил Disallow', str(results.get('disallow_rules', 0))],
            ['Правил Allow', str(results.get('allow_rules', 0))],
            ['Объявлено sitemap', str(len(results.get('sitemaps', []) or []))],
            ['Объявлено host', str(len(results.get('hosts', []) or []))],
            ['Директив crawl-delay', str(len(results.get('crawl_delays', {}) or {}))],
            ['Директив clean-param', str(len(results.get('clean_params', []) or []))],
        ]
        self._add_table(doc, ['Метрика', 'Значение'], summary_rows)

        severity = results.get('severity_counts', {}) or {}
        self._add_heading(doc, '2. Обзор критичности', level=1)
        self._add_table(
            doc,
            ['Критично', 'Предупреждение', 'Инфо'],
            [[
                str(severity.get('critical', 0)),
                str(severity.get('warning', 0)),
                str(severity.get('info', 0)),
            ]],
        )

        self._add_heading(doc, '3. Критические проблемы', level=1)
        critical = results.get('critical_issues', []) or results.get('issues', []) or []
        if critical:
            for item in critical:
                doc.add_paragraph(str(item), style='List Bullet')
        else:
            doc.add_paragraph('Критические проблемы не обнаружены.')

        self._add_heading(doc, '4. Предупреждения', level=1)
        warnings = results.get('warning_issues', []) or results.get('warnings', []) or []
        if warnings:
            for item in warnings:
                doc.add_paragraph(str(item), style='List Bullet')
        else:
            doc.add_paragraph('Предупреждения не обнаружены.')

        self._add_heading(doc, '5. Информационные заметки', level=1)
        info_issues = results.get('info_issues', []) or []
        if info_issues:
            for item in info_issues:
                doc.add_paragraph(str(item), style='List Bullet')
        else:
            doc.add_paragraph('Информационные заметки отсутствуют.')

        self._add_heading(doc, '6. Рекомендации', level=1)
        recommendations = results.get('recommendations', []) or []
        if recommendations:
            for item in recommendations:
                doc.add_paragraph(str(item), style='List Bullet')
        else:
            doc.add_paragraph('Рекомендации не сформированы.')

        doc.add_page_break()
        self._add_heading(doc, '7. Топ исправлений', level=1)
        top_fixes = results.get('top_fixes', []) or []
        if top_fixes:
            rows = []
            for fix in top_fixes:
                rows.append([
                    str(fix.get('priority', 'medium')).upper(),
                    str(fix.get('title', '')),
                    str(fix.get('why', '')),
                    str(fix.get('action', '')),
                ])
            self._add_table(doc, ['Приоритет', 'Заголовок', 'Почему', 'Действие'], rows)
        else:
            doc.add_paragraph('Приоритизированные исправления не сформированы.')

        self._add_heading(doc, '8. Проверки Sitemap', level=1)
        sitemap_checks = results.get('sitemap_checks', []) or []
        if sitemap_checks:
            rows = []
            for check in sitemap_checks:
                ok = check.get('ok')
                if ok is True:
                    status = 'OK'
                elif ok is False:
                    status = 'ОШИБКА'
                else:
                    status = 'ПРОПУЩЕНО'
                rows.append([
                    str(check.get('url', '')),
                    status,
                    str(check.get('status_code', '')),
                    str(check.get('error', '')),
                ])
            self._add_table(doc, ['URL', 'Статус', 'HTTP', 'Ошибка'], rows)
        else:
            doc.add_paragraph('Проверки sitemap недоступны.')

        doc.add_page_break()
        self._add_heading(doc, '9. Детали правил групп', level=1)
        groups = results.get('groups_detail', []) or []
        if groups:
            for idx, group in enumerate(groups, start=1):
                uas = ', '.join(group.get('user_agents', []) or [])
                doc.add_paragraph(f"Группа {idx}: {uas}")
                disallow_rows = [
                    [str(item.get('path', '')), str(item.get('line', ''))]
                    for item in (group.get('disallow', []) or [])
                ]
                allow_rows = [
                    [str(item.get('path', '')), str(item.get('line', ''))]
                    for item in (group.get('allow', []) or [])
                ]
                if disallow_rows:
                    doc.add_paragraph('Disallow:')
                    self._add_table(doc, ['Путь', 'Строка'], disallow_rows)
                if allow_rows:
                    doc.add_paragraph('Allow:')
                    self._add_table(doc, ['Путь', 'Строка'], allow_rows)
                if (not disallow_rows) and (not allow_rows):
                    doc.add_paragraph('В этой группе нет правил allow/disallow.')
        else:
            doc.add_paragraph('Детализация групп недоступна.')

        self._add_heading(doc, '10. Синтаксические ошибки', level=1)
        syntax_errors = results.get('syntax_errors', []) or []
        if syntax_errors:
            rows = []
            for err in syntax_errors:
                rows.append([
                    str(err.get('line', '')),
                    str(err.get('error', '')),
                    str(err.get('content', '')),
                ])
            self._add_table(doc, ['Строка', 'Ошибка', 'Содержимое'], rows)
        else:
            doc.add_paragraph('Синтаксические ошибки не обнаружены.')

        self._add_heading(doc, '11. Расширенный анализ', level=1)
        http_status_analysis = results.get('http_status_analysis', {}) or {}
        if http_status_analysis:
            doc.add_paragraph(f"Контекст HTTP-статуса: {http_status_analysis.get('status_code', 'н/д')}")
            for note in (http_status_analysis.get('notes', []) or []):
                doc.add_paragraph(str(note), style='List Bullet')

        unsupported_directives = results.get('unsupported_directives', []) or []
        if unsupported_directives:
            rows = []
            for item in unsupported_directives:
                rows.append([
                    str(item.get('line', '')),
                    str(item.get('directive', '')),
                    str(item.get('value', '')),
                ])
            self._add_table(doc, ['Строка', 'Директива', 'Значение'], rows)

        host_validation = results.get('host_validation', {}) or {}
        if host_validation:
            hosts = ', '.join(host_validation.get('hosts', []) or []) or 'н/д'
            doc.add_paragraph(f"Директивы Host: {hosts}")
            for note in (host_validation.get('warnings', []) or []):
                doc.add_paragraph(str(note), style='List Bullet')

        directive_conflicts = results.get('directive_conflicts', {}) or {}
        conflict_items = directive_conflicts.get('details', []) or []
        if conflict_items:
            rows = []
            for item in conflict_items:
                rows.append([
                    str(item.get('type', '')),
                    str(item.get('user_agent', '')),
                    str(item.get('path', item.get('groups', ''))),
                ])
            self._add_table(doc, ['Тип конфликта', 'User-agent', 'Путь/значение'], rows)

        longest_match = results.get('longest_match_analysis', {}) or {}
        longest_notes = longest_match.get('notes', []) or []
        if longest_notes:
            doc.add_paragraph('Примечания по longest-match:')
            for note in longest_notes:
                doc.add_paragraph(str(note), style='List Bullet')

        param_recommendations = results.get('param_recommendations', []) or []
        if param_recommendations:
            doc.add_paragraph('Рекомендации по Yandex Clean-param:')
            for rec in param_recommendations:
                doc.add_paragraph(str(rec), style='List Bullet')

        self._add_heading(doc, '12. Покрытие по ботам', level=1)
        present_agents = results.get('present_agents', []) or []
        missing_bots = results.get('missing_bots', []) or []
        if present_agents:
            doc.add_paragraph('Обнаруженные группы user-agent:')
            for ua in present_agents:
                doc.add_paragraph(str(ua), style='List Bullet')
        else:
            doc.add_paragraph('Явные группы ботов не обнаружены.')
        if missing_bots:
            doc.add_paragraph('Рекомендуемые боты для явных правил:')
            for bot in missing_bots:
                doc.add_paragraph(str(bot), style='List Bullet')
        else:
            doc.add_paragraph('Key bots coverage looks complete.')

        doc.add_page_break()
        self._add_heading(doc, '13. Raw robots.txt (построчный просмотр)', level=1)
        raw = str(results.get('raw_content', '') or '')
        if raw:
            for idx, line in enumerate(raw.splitlines(), start=1):
                p_line = doc.add_paragraph()
                run_num = p_line.add_run(f"{idx:4} | ")
                run_num.font.size = Pt(8)
                run_num.font.color.rgb = RGBColor(120, 120, 120)
                run_num.font.name = 'Consolas'
                run_txt = p_line.add_run(line)
                run_txt.font.size = Pt(9)
                run_txt.font.name = 'Consolas'
        else:
            doc.add_paragraph('Содержимое raw robots.txt недоступно.')

        self._add_heading(doc, '14. Снимок дополнительных полей', level=1)
        covered = {
            'robots_txt_found', 'status_code', 'quality_score', 'quality_grade', 'production_ready', 'quick_status',
            'content_length', 'lines_count', 'user_agents', 'disallow_rules', 'allow_rules', 'sitemaps', 'hosts',
            'crawl_delays', 'clean_params', 'severity_counts', 'critical_issues', 'issues', 'warning_issues',
            'warnings', 'info_issues', 'recommendations', 'top_fixes', 'sitemap_checks', 'groups_detail', 'syntax_errors', 'raw_content',
            'http_status_analysis', 'unsupported_directives', 'host_validation', 'directive_conflicts', 'longest_match_analysis',
            'param_recommendations', 'present_agents', 'missing_bots',
            'machine_summary', 'error', 'can_continue',
        }
        extra_rows = []
        for key in sorted(results.keys()):
            if key in covered:
                continue
            value = results.get(key)
            if isinstance(value, (dict, list)):
                rendered = json.dumps(value, ensure_ascii=False)
            else:
                rendered = str(value)
            if len(rendered) > 500:
                rendered = rendered[:500] + ' ...'
            extra_rows.append([key, rendered])
        if extra_rows:
            self._add_table(doc, ['Поле', 'Значение'], extra_rows)
        else:
            doc.add_paragraph('Дополнительных полей сверх базовых разделов не обнаружено.')

        self._add_heading(doc, '15. Официальная документация', level=1)
        doc.add_paragraph('Google Search Central: https://developers.google.com/search/docs/crawling-indexing/robots/robots_txt')
        doc.add_paragraph('Yandex Webmaster: https://yandex.com/support/webmaster/en/robot-workings/allow-disallow')
        doc.add_paragraph('Yandex Clean-param: https://yandex.com/support/webmaster/en/robot-workings/clean-param')

        self._add_full_result_appendix(doc, data)
        filepath = os.path.join(self.reports_dir, f"{task_id}.docx")
        self._save_document(doc, filepath)
        return filepath

    def generate_sitemap_report(self, task_id: str, data: Dict[str, Any]) -> str:
        """Generate extended sitemap validation DOCX report."""
        doc = Document()
        results = data.get('results', {}) or {}
        now_str = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

        self._add_cover_page(doc, 'Sitemap Report', 'Проверка Sitemap', data.get('url', ''), now_str)
        self._add_header_footer(doc, 'Sitemap Report')
        input_url = results.get("input_url") or data.get('url', 'н/д')
        resolved_url = results.get("resolved_sitemap_url") or data.get('url', 'н/д')
        discovery_source = results.get("sitemap_discovery_source", "")
        doc.add_paragraph(f"Входной URL: {input_url}")
        doc.add_paragraph(f"URL sitemap: {resolved_url}")
        if discovery_source:
            source_map = {
                "direct_input": "прямой URL",
                "robots.txt": "robots.txt",
                "common_path": "стандартный путь",
                "auto_discovery": "автоопределение",
            }
            doc.add_paragraph(f"Источник обнаружения: {source_map.get(discovery_source, discovery_source)}")
        doc.add_paragraph(f"Сформировано: {now_str}")

        summary_rows = [
            ["Валиден", "Да" if results.get("valid") else "Нет"],
            ["HTTP статус", results.get("status_code", "н/д")],
            ["Sitemap-файлов просканировано", results.get("sitemaps_scanned", 0)],
            ["Валидных sitemap-файлов", results.get("sitemaps_valid", 0)],
            ["Всего URL", results.get("urls_count", 0)],
            ["Уникальных URL", results.get("unique_urls_count", 0)],
            ["Дубли URL", results.get("duplicate_urls_count", 0)],
            ["Оценка качества", f"{results.get('quality_score', 'н/д')} ({results.get('quality_grade', 'н/д')})"],
        ]
        self._add_heading(doc, "1. Сводка", level=1)
        self._add_table(doc, ["Метрика", "Значение"], summary_rows)

        severity = results.get("severity_counts", {}) or {}
        self._add_heading(doc, "2. Распределение по критичности", level=1)
        self._add_table(
            doc,
            ["Критично", "Предупреждение", "Инфо"],
            [[severity.get("critical", 0), severity.get("warning", 0), severity.get("info", 0)]],
        )

        issues = results.get("issues", []) or []
        self._add_heading(doc, "3. Приоритизированные проблемы", level=1)
        if issues:
            rows = []
            for it in issues[:40]:
                rows.append([
                    it.get("severity", ""),
                    it.get("title", ""),
                    it.get("details", ""),
                    it.get("action", ""),
                ])
            self._add_table(doc, ["Критичность", "Проблема", "Детали", "Действие"], rows)
        else:
            doc.add_paragraph("Приоритизированные проблемы отсутствуют.")

        action_plan = results.get("action_plan", []) or []
        self._add_heading(doc, "4. План исправлений", level=1)
        if action_plan:
            rows = []
            for item in action_plan[:30]:
                rows.append([
                    item.get("priority", ""),
                    item.get("owner", ""),
                    item.get("issue", ""),
                    item.get("action", ""),
                    item.get("sla", ""),
                ])
            self._add_table(doc, ["Приоритет", "Ответственный", "Проблема", "Действие", "SLA"], rows)
        else:
            doc.add_paragraph("Пункты плана исправлений отсутствуют.")

        hreflang = results.get("hreflang", {}) or {}
        freshness = results.get("freshness", {}) or {}
        media = results.get("media_extensions", {}) or {}
        self._add_heading(doc, "5. Расширенная валидация", level=1)
        self._add_table(
            doc,
            ["Проверка", "Значение"],
            [
                ["Hreflang обнаружен", "Да" if hreflang.get("detected") else "Нет"],
                ["Hreflang ссылок", hreflang.get("links_count", 0)],
                ["Hreflang: некорректный код", hreflang.get("invalid_code_count", 0)],
                ["Hreflang: некорректный href", hreflang.get("invalid_href_count", 0)],
                ["Hreflang: дубль языка", hreflang.get("duplicate_lang_count", 0)],
                ["Lastmod отсутствует", freshness.get("lastmod_missing_count", 0)],
                ["Lastmod устаревший", freshness.get("stale_lastmod_count", 0)],
                ["Lastmod в будущем", freshness.get("lastmod_future_count", 0)],
                ["Image теги / без loc", f"{media.get('image_tags_count', 0)} / {media.get('image_missing_loc_count', 0)}"],
                ["Video теги / без обязательных", f"{media.get('video_tags_count', 0)} / {media.get('video_missing_required_count', 0)}"],
                ["News теги / без обязательных", f"{media.get('news_tags_count', 0)} / {media.get('news_missing_required_count', 0)}"],
            ],
        )

        sitemap_files = results.get("sitemap_files", []) or []
        self._add_heading(doc, "6. Ошибки и предупреждения по файлам", level=1)
        if sitemap_files:
            rows = []
            for item in sitemap_files[:200]:
                errors_txt = " | ".join((item.get("errors") or [])[:5])
                warnings_txt = " | ".join((item.get("warnings") or [])[:5])
                if not errors_txt and not warnings_txt:
                    continue
                rows.append([
                    item.get("sitemap_url", ""),
                    errors_txt,
                    warnings_txt,
                ])
            if rows:
                self._add_table(doc, ["Sitemap-файл", "Ошибки", "Предупреждения"], rows)
            else:
                doc.add_paragraph("Ошибки или предупреждения по файлам не обнаружены.")
        else:
            doc.add_paragraph("В результате нет sitemap-файлов.")

        tool_notes = results.get("tool_notes", []) or []
        if tool_notes:
            self._add_heading(doc, "7. Служебные заметки (не ошибки sitemap)", level=1)
            for note in tool_notes[:30]:
                doc.add_paragraph(str(note), style='List Bullet')

        checks = results.get("live_indexability_checks", []) or []
        doc.add_page_break()
        self._add_heading(doc, "8. Live-выборка индексируемости", level=1)
        if checks:
            rows = []
            for item in checks[:20]:
                rows.append([
                    item.get("url", ""),
                    item.get("status_code", ""),
                    "Да" if item.get("indexable") else "Нет",
                    f"{item.get('canonical_status', 'н/д')}: {item.get('canonical_url', '')}",
                    "; ".join(item.get("reasons", [])[:2]),
                ])
            self._add_table(doc, ["URL", "HTTP", "Индексируемость", "Canonical", "Причины"], rows)
        else:
            doc.add_paragraph("Live-выборка пуста.")

        recs = results.get("recommendations", []) or []
        self._add_heading(doc, "9. Рекомендации", level=1)
        if recs:
            for rec in recs[:40]:
                doc.add_paragraph(str(rec), style='List Bullet')
        else:
            doc.add_paragraph("Рекомендации отсутствуют.")

        self._add_full_result_appendix(doc, data)
        filepath = os.path.join(self.reports_dir, f"{task_id}.docx")
        self._save_document(doc, filepath)
        return filepath

    def generate_render_report(self, task_id: str, data: Dict[str, Any]) -> str:
        """Генерирует расширенный клиентский отчет аудита рендеринга."""
        doc = Document()

        url = data.get('url', 'н/д')
        results = data.get('results', {}) or {}
        summary = results.get('summary', {}) or {}
        variants = results.get('variants', []) or []
        issues = results.get('issues', []) or []
        recommendations = results.get('recommendations', []) or []
        generated_at = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

        self._add_cover_page(doc, 'Render Audit Report', 'Render-аудит', data.get('url', ''), generated_at)
        self._add_header_footer(doc, 'Render Audit Report')

        doc.add_paragraph(f"Анализируемый URL: {url}")
        doc.add_paragraph(f"Дата аудита: {generated_at}")
        doc.add_paragraph(f"Движок: {self._format_engine_label(results.get('engine', 'legacy'))}")
        doc.add_paragraph("Краткое резюме")
        doc.add_paragraph(
            f"Данный отчёт содержит результаты SEO-аудита страницы {url} с фокусом на сравнение контента, "
            "доступного с JavaScript и без него."
        )

        self._add_heading(doc, '1. Ключевые выводы', level=1)
        totals = [
            ['Профилей проверки', summary.get('variants_total', len(variants)), 'Инфо'],
            ['Общая оценка', summary.get('score', 'н/д'), 'Инфо'],
            ['Критичные проблемы', summary.get('critical_issues', 0), 'Важно'],
            ['Предупреждения', summary.get('warning_issues', 0), 'Важно'],
            ['Всего потерянных элементов', summary.get('missing_total', 0), 'Важно'],
            ['Средний процент потерь', f"{summary.get('avg_missing_pct', 0)}%", 'Инфо'],
        ]
        self._add_table(doc, ['Показатель', 'Значение', 'Статус'], totals)
        doc.add_paragraph(f"• SEO-оценка: {summary.get('score', 'н/д')}/100")
        doc.add_paragraph(f"• Критических проблем: {summary.get('critical_issues', 0)}")
        doc.add_paragraph(f"• Предупреждений: {summary.get('warning_issues', 0)}")
        doc.add_paragraph(f"• Добавлено элементов через JS: {summary.get('missing_total', 0)}")

        primary = variants[0] if variants else {}
        raw_p = primary.get('raw', {}) or {}
        js_p = primary.get('rendered', {}) or {}
        meta_p = ((primary.get('meta_non_seo') or {}).get('comparison') or {})
        raw_meta = ((primary.get('meta_non_seo') or {}).get('raw') or {})
        js_meta = ((primary.get('meta_non_seo') or {}).get('rendered') or {})
        raw_schema = ", ".join(raw_p.get('schema_types', []) or []) or "Нет"
        js_schema = ", ".join(js_p.get('schema_types', []) or []) or "Нет"

        def _status(a, b) -> str:
            return "ОК" if str(a).strip() == str(b).strip() else "РАЗЛ."

        self._add_heading(doc, '2. Сравнение SEO-элементов', level=1)
        compare_rows = [
            ['Заголовок страницы (title)', raw_p.get('title', ''), js_p.get('title', ''), _status(raw_p.get('title', ''), js_p.get('title', ''))],
            ['Мета-описание (description)', raw_p.get('meta_description', ''), js_p.get('meta_description', ''), _status(raw_p.get('meta_description', ''), js_p.get('meta_description', ''))],
            ['H1 заголовки', f"{raw_p.get('h1_count', 0)} шт", f"{js_p.get('h1_count', 0)} шт", _status(raw_p.get('h1_count', 0), js_p.get('h1_count', 0))],
            ['H2 заголовки', f"{raw_p.get('h2_count', 0)} шт", f"{js_p.get('h2_count', 0)} шт", _status(raw_p.get('h2_count', 0), js_p.get('h2_count', 0))],
            ['Изображения', f"{raw_p.get('images_count', 0)} шт", f"{js_p.get('images_count', 0)} шт", _status(raw_p.get('images_count', 0), js_p.get('images_count', 0))],
            ['Ссылки', f"{raw_p.get('links_count', 0)} шт", f"{js_p.get('links_count', 0)} шт", _status(raw_p.get('links_count', 0), js_p.get('links_count', 0))],
            ['Canonical', raw_p.get('canonical', ''), js_p.get('canonical', ''), _status(raw_p.get('canonical', ''), js_p.get('canonical', ''))],
            ['Schema-разметка', raw_schema, js_schema, _status(raw_schema, js_schema)],
            ['Мета viewport', 'Есть' if raw_meta.get('meta:viewport') else 'Нет', 'Есть' if js_meta.get('meta:viewport') else 'Нет', _status(bool(raw_meta.get('meta:viewport')), bool(js_meta.get('meta:viewport')))],
        ]
        self._add_table(doc, ['Элемент', 'Без JS', 'С JS', 'Статус'], compare_rows)

        self._add_heading(doc, '3. Анализ по устройствам', level=1)
        if variants:
            device_rows = []
            for variant in variants:
                metrics = variant.get('metrics', {}) or {}
                shots = variant.get('screenshots', {}) or {}
                device_rows.append([
                    variant.get('variant_label') or variant.get('variant_id', 'Профиль'),
                    self._format_profile_label(variant.get('profile_type'), variant.get('mobile')),
                    f"{float(metrics.get('score', 0) or 0):.1f}",
                    f"{int(metrics.get('total_missing', 0) or 0)}",
                    ", ".join(list(shots.keys())) if shots else 'нет',
                ])
            self._add_table(doc, ['Устройство', 'Тип профиля', 'Оценка', 'Потери', 'Скриншоты'], device_rows)
        else:
            doc.add_paragraph("Профили устройств отсутствуют.")

        self._add_heading(doc, '4. Что проверялось и зачем', level=1)
        doc.add_paragraph(
            "Система сравнивает две версии страницы: исходный HTML (без JS) и итоговый DOM после JS-рендера. "
            "Если значимый контент появляется только после JS, поисковый робот может увидеть неполную страницу."
        )
        doc.add_paragraph("Проверяются элементы: title, meta description, canonical, H1-H2, ссылки, изображения, schema.org, видимый текст.")
        doc.add_paragraph("Дополнительно сравниваются не-SEO meta-данные между no-JS и JS: viewport, charset, referrer, theme-color, manifest и др.")

        doc.add_page_break()
        self._add_heading(doc, '5. Результаты по профилям', level=1)
        if variants:
            for variant in variants:
                label = variant.get('variant_label') or variant.get('variant_id', 'Профиль')
                metrics = variant.get('metrics', {}) or {}
                raw = variant.get('raw', {}) or {}
                rendered = variant.get('rendered', {}) or {}
                doc.add_heading(label, level=2)

                rows = [
                    ['Оценка', f"{metrics.get('score', 0):.1f}/100", ''],
                    ['Потерянные элементы', int(metrics.get('total_missing', 0) or 0), ''],
                    ['Потери %', f"{metrics.get('missing_pct', 0):.1f}%", ''],
                    ['H1 (без JS / с JS)', f"{raw.get('h1_count', 0)} / {rendered.get('h1_count', 0)}", ''],
                    ['Ссылки (без JS / с JS)', f"{raw.get('links_count', 0)} / {rendered.get('links_count', 0)}", ''],
                    ['Структурированные данные (без JS / с JS)', f"{raw.get('structured_data_count', 0)} / {rendered.get('structured_data_count', 0)}", ''],
                ]
                self._add_table(doc, ['Параметр', 'Значение', ''], rows)

                seo_required = variant.get('seo_required', {}) or {}
                seo_items = seo_required.get('items', []) or []
                if seo_items:
                    doc.add_paragraph(
                        f"Обязательные SEO-элементы: Пройдено {seo_required.get('pass', 0)}, "
                        f"Предупреждений {seo_required.get('warn', 0)}, Критичных {seo_required.get('fail', 0)}."
                    )
                    seo_rows = []
                    status_map = {'pass': 'Пройдено', 'warn': 'Предупреждение', 'fail': 'Критично'}
                    for item in seo_items:
                        seo_rows.append([
                            item.get('label', ''),
                            item.get('raw', ''),
                            item.get('rendered', ''),
                            status_map.get(item.get('status', ''), item.get('status', '')),
                            item.get('fix', ''),
                        ])
                    self._add_table(doc, ['Элемент', 'Без JS', 'С JS', 'Статус', 'Что исправить'], seo_rows[:80])

                var_issues = variant.get('issues', []) or []
                if var_issues:
                    doc.add_paragraph("Найденные проблемы:", style='List Bullet')
                    for issue in var_issues:
                        sev = str(issue.get('severity', 'info')).upper()
                        title_i = issue.get('title', '')
                        details_i = issue.get('details', '')
                        doc.add_paragraph(f"[{sev}] {title_i}: {details_i}", style='List Bullet')
                        examples = issue.get('examples', []) or []
                        for ex in examples[:5]:
                            doc.add_paragraph(f"Пример: {ex}", style='List Bullet')
                else:
                    doc.add_paragraph("Проблемы не обнаружены.")

                missing = variant.get('missing', {}) or {}
                for key, label in [
                    ('visible_text', 'Текст, который появляется только после JS'),
                    ('headings', 'Заголовки, которые появляются только после JS'),
                    ('links', 'Ссылки, которые появляются только после JS'),
                    ('structured_data', 'Структурированные данные, которые появляются только после JS'),
                ]:
                    values = missing.get(key, []) or []
                    if not values:
                        continue
                    doc.add_paragraph(f"{label}: {len(values)}", style='List Bullet')
                    for item in values[:20]:
                        doc.add_paragraph(str(item), style='List Bullet')

                meta_cmp = ((variant.get('meta_non_seo') or {}).get('comparison') or {})
                meta_items = meta_cmp.get('items', []) or []
                if meta_items:
                    doc.add_paragraph(
                        f"Мета-данные (не SEO): всего {meta_cmp.get('total', 0)}, "
                        f"совпадает {meta_cmp.get('same', 0)}, изменено {meta_cmp.get('changed', 0)}, "
                        f"только JS {meta_cmp.get('only_rendered', 0)}, только без JS {meta_cmp.get('only_raw', 0)}."
                    )
                    meta_rows = []
                    for item in meta_items:
                        status_map = {
                            'same': 'Совпадает',
                            'changed': 'Изменено',
                            'only_rendered': 'Только в JS',
                            'only_raw': 'Только без JS',
                        }
                        meta_rows.append([
                            item.get('key', ''),
                            item.get('raw', ''),
                            item.get('rendered', ''),
                            status_map.get(item.get('status', ''), item.get('status', '')),
                        ])
                    self._add_table(doc, ['Ключ', 'Без JS', 'С JS', 'Статус'], meta_rows[:50])

                var_recs = variant.get('recommendations', []) or []
                if var_recs:
                    doc.add_paragraph("Рекомендации по профилю:", style='List Bullet')
                    for rec in var_recs:
                        doc.add_paragraph(str(rec), style='List Bullet')

                screenshots = variant.get('screenshots', {}) or {}
                for key in ('js', 'nojs', 'js_landscape', 'nojs_landscape'):
                    shot = screenshots.get(key) or {}
                    shot_path = shot.get('path')
                    if shot_path and os.path.exists(shot_path):
                        caption = {
                            'js': 'Скриншот: рендер с JavaScript',
                            'nojs': 'Скриншот: версия без JavaScript',
                            'js_landscape': 'Скриншот: мобильный (горизонтальный), с JavaScript',
                            'nojs_landscape': 'Скриншот: мобильный (горизонтальный), без JavaScript',
                        }.get(key, 'Скриншот')
                        doc.add_paragraph(caption)
                        doc.add_picture(shot_path, width=Inches(6.5))
        else:
            doc.add_paragraph("Данные по профилям отсутствуют.")

        doc.add_page_break()
        self._add_heading(doc, '6. Общий список ошибок', level=1)
        if issues:
            for issue in issues:
                sev = str(issue.get('severity', 'info')).upper()
                profile = issue.get('variant', 'Профиль')
                title_i = issue.get('title', '')
                details_i = issue.get('details', '')
                doc.add_paragraph(f"[{sev}] {profile}: {title_i} - {details_i}", style='List Bullet')
        else:
            doc.add_paragraph("Ошибок не обнаружено.")

        self._add_heading(doc, '7. Что делать для исправления', level=1)
        if recommendations:
            for rec in recommendations:
                doc.add_paragraph(str(rec), style='List Bullet')
        else:
            doc.add_paragraph("Критичных рекомендаций по итогам проверки нет.")

        self._add_full_result_appendix(doc, data)
        filepath = os.path.join(self.reports_dir, f"{task_id}.docx")
        self._save_document(doc, filepath)
        return filepath

    def generate_mobile_report(self, task_id: str, data: Dict[str, Any]) -> str:
        """Extended mobile DOCX report with Russian content."""
        doc = Document()

        issue_guides = {
            "viewport_missing": {
                "name": "Отсутствует meta viewport",
                "why": "Страница отображается как десктопная и плохо адаптируется под мобильные экраны.",
                "impact": "Высокое",
                "fix": [
                    "Добавить тег <meta name=\"viewport\" content=\"width=device-width, initial-scale=1\"> в <head>.",
                    "Проверить наличие тега на всех шаблонах страниц.",
                ],
            },
            "viewport_invalid": {
                "name": "Некорректный viewport",
                "why": "Неправильные параметры viewport ломают масштаб и адаптивность.",
                "impact": "Высокое",
                "fix": [
                    "Исправить viewport на width=device-width, initial-scale=1.",
                    "Убрать дублирующиеся viewport-теги.",
                ],
            },
            "horizontal_overflow": {
                "name": "Горизонтальная прокрутка",
                "why": "Контент выходит за ширину экрана.",
                "impact": "Высокое",
                "fix": [
                    "Найти элементы с шириной больше viewport.",
                    "Добавить max-width: 100% для медиа-элементов.",
                ],
            },
            "small_touch_targets": {
                "name": "Маленькие интерактивные элементы",
                "why": "Элементы < 44x44px неудобны для тача.",
                "impact": "Среднее",
                "fix": [
                    "Увеличить размеры кнопок/ссылок до 44x44px и больше.",
                    "Добавить отступы между соседними элементами.",
                ],
            },
            "small_fonts": {
                "name": "Слишком мелкий текст",
                "why": "Низкая читаемость ухудшает UX-метрики.",
                "impact": "Среднее",
                "fix": [
                    "Привести базовый размер шрифта к 16px+ на мобильных экранах.",
                    "Проверить читаемость ключевых блоков.",
                ],
            },
            "large_images": {
                "name": "Изображения шире экрана",
                "why": "Широкие изображения ломают сетку и провоцируют скролл.",
                "impact": "Среднее",
                "fix": [
                    "Добавить для изображений max-width: 100%; height: auto;.",
                    "Проверить адаптивность баннеров/слайдеров.",
                ],
            },
            "console_errors": {
                "name": "Ошибки JavaScript в консоли",
                "why": "JS-ошибки могут ломать ключевые UI-сценарии.",
                "impact": "Среднее",
                "fix": [
                    "Разобрать ошибки консоли по приоритету.",
                    "Исправить недоступные ресурсы и исключения.",
                ],
            },
            "runtime_error": {
                "name": "Ошибка выполнения проверки",
                "why": "Часть данных по устройству может быть неполной.",
                "impact": "Высокое",
                "fix": [
                    "Проверить доступность сайта и окружение анализа.",
                    "Повторить аудит после устранения причины.",
                ],
            },
            "playwright_unavailable": {
                "name": "Среда Playwright недоступна",
                "why": "Без браузерного движка нельзя выполнить полный аудит.",
                "impact": "Высокое",
                "fix": [
                    "Установить Playwright и Chromium в окружении сервера.",
                    "Перезапустить сервис и повторить аудит.",
                ],
            },
            "mobile_engine_error": {
                "name": "Сбой движка мобильной проверки",
                "why": "Движок не смог корректно завершить анализ.",
                "impact": "Высокое",
                "fix": [
                    "Проверить логи сервиса и окружение выполнения.",
                    "Устранить причину и повторно запустить аудит.",
                ],
            },
        }

        url = data.get("url", "н/д")
        generated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        self._add_cover_page(doc, 'Mobile Audit Report', 'Проверка мобильной версии', data.get('url', ''), generated_at)
        self._add_header_footer(doc, 'Mobile Audit Report')

        doc.add_paragraph(f"Сайт: {url}")
        doc.add_paragraph(f"Отчет сформирован: {generated_at}")

        results = data.get("results", {}) or {}
        summary = results.get("summary", {}) or {}
        devices = results.get("device_results", []) or []
        all_issues = results.get("issues", []) or []
        artifacts = results.get("artifacts", {}) or {}
        screenshot_dir = str(artifacts.get("screenshot_dir") or "").strip()
        server_base_url = str(data.get("server_base_url") or "").strip()
        actionable_issues = [i for i in all_issues if i.get("severity") in ("critical", "warning")]
        info_issues = [i for i in all_issues if i.get("severity") == "info"]
        temp_screenshots: List[str] = []

        def _resolve_mobile_screenshot_path(device: Dict[str, Any]) -> str:
            candidates: List[str] = []
            raw_path = str(device.get("screenshot_path") or "").strip()
            shot_name = str(device.get("screenshot_name") or "").strip()
            shot_url = str(device.get("screenshot_url") or "").strip()

            if raw_path:
                candidates.append(raw_path)
            if shot_name and screenshot_dir:
                candidates.append(os.path.join(screenshot_dir, shot_name))
            if shot_name:
                candidates.append(os.path.join(self.reports_dir, "mobile", task_id, "screenshots", shot_name))
            if raw_path:
                candidates.append(
                    os.path.join(self.reports_dir, "mobile", task_id, "screenshots", os.path.basename(raw_path))
                )

            for candidate in candidates:
                if candidate and os.path.exists(candidate):
                    return candidate

            if shot_url:
                if shot_url.startswith("/"):
                    if server_base_url:
                        shot_url = urljoin(server_base_url, shot_url)
                    else:
                        shot_url = ""
                if shot_url:
                    try:
                        response = requests.get(shot_url, timeout=25)
                        if response.status_code == 200 and response.content:
                            fd, temp_path = tempfile.mkstemp(prefix=f"mobile_docx_{task_id}_", suffix=".png")
                            os.close(fd)
                            with open(temp_path, "wb") as f:
                                f.write(response.content)
                            temp_screenshots.append(temp_path)
                            return temp_path
                    except Exception:
                        pass
            return ""

        self._add_heading(doc, "1. Сводка по проверке", level=1)
        summary_rows = [
            ["Движок проверки", self._format_engine_label(results.get("engine", "legacy"))],
            ["Режим проверки", "Быстрый" if results.get("mode") == "quick" else "Полный"],
            ["Проверено устройств", summary.get("total_devices", len(devices))],
            ["Устройств без критичных проблем", summary.get("mobile_friendly_devices", 0)],
            ["Устройств с проблемами", summary.get("non_friendly_devices", 0)],
            ["Среднее время загрузки, мс", summary.get("avg_load_time_ms", 0)],
            ["Ошибок (critical + warning)", len(actionable_issues)],
            ["Информационных замечаний", len(info_issues)],
            ["Интегральная оценка", results.get("score", "н/д")],
            ["Итог", "Соответствует требованиям" if results.get("mobile_friendly") else "Требуются доработки"],
        ]
        self._add_table(doc, ["Показатель", "Значение"], summary_rows)

        self._add_heading(doc, "2. Технические параметры", level=1)
        tech_rows = [
            ["HTTP", results.get("status_code", "н/д")],
            ["Итоговый URL", results.get("final_url", url)],
            ["Viewport найден", "Да" if results.get("viewport_found") else "Нет"],
            ["Содержимое viewport", results.get("viewport_content") or "-"],
        ]
        self._add_table(doc, ["Параметр", "Значение"], tech_rows)

        self._add_heading(doc, "3. Результаты по устройствам", level=1)
        if devices:
            device_rows = []
            for d in devices:
                category = d.get("category", "")
                if category == "phone":
                    category = "Телефон"
                elif category == "tablet":
                    category = "Планшет"
                device_rows.append([
                    d.get("device_name", ""),
                    category,
                    f"{(d.get('viewport') or {}).get('width', '-')}x{(d.get('viewport') or {}).get('height', '-')}",
                    d.get("status_code", "н/д"),
                    d.get("load_time_ms", 0),
                    d.get("issues_count", 0),
                    "Да" if d.get("mobile_friendly") else "Нет",
                ])
            self._add_table(
                doc,
                ["Устройство", "Тип", "Viewport", "HTTP", "Загрузка, мс", "Ошибок", "ОК для мобильной версии"],
                device_rows,
            )
        else:
            doc.add_paragraph("Данные по устройствам отсутствуют.")

        self._add_heading(doc, "3.1 \u0414\u0435\u0442\u0430\u043b\u0438\u0437\u0430\u0446\u0438\u044f \u043f\u0440\u043e\u0431\u043b\u0435\u043c \u043f\u043e \u0443\u0441\u0442\u0440\u043e\u0439\u0441\u0442\u0432\u0430\u043c", level=2)
        detailed_added = 0
        for d in devices:
            device_issues = d.get("issues", []) or []
            if not device_issues:
                continue
            detailed_added += 1
            device_name = d.get("device_name") or "\u0423\u0441\u0442\u0440\u043e\u0439\u0441\u0442\u0432\u043e"
            doc.add_paragraph(
                f"{device_name} | "
                f"Viewport {(d.get('viewport') or {}).get('width', '-')}x{(d.get('viewport') or {}).get('height', '-')} | "
                f"\u0417\u0430\u043c\u0435\u0447\u0430\u043d\u0438\u0439: {len(device_issues)}"
            )
            for issue in device_issues[:15]:
                sev = str(issue.get("severity", "info")).upper()
                title_i = str(issue.get("title", issue.get("code", "\u041f\u0440\u043e\u0431\u043b\u0435\u043c\u0430")))
                details_i = str(issue.get("details", "") or "").strip()
                if details_i:
                    doc.add_paragraph(f"[{sev}] {title_i}: {details_i}", style="List Bullet")
                else:
                    doc.add_paragraph(f"[{sev}] {title_i}", style="List Bullet")
        if detailed_added == 0:
            doc.add_paragraph("\u0414\u0435\u0442\u0430\u043b\u0438 \u043f\u043e \u043f\u0440\u043e\u0431\u043b\u0435\u043c\u0430\u043c \u0443\u0441\u0442\u0440\u043e\u0439\u0441\u0442\u0432 \u043e\u0442\u0441\u0443\u0442\u0441\u0442\u0432\u0443\u044e\u0442.")

        doc.add_page_break()
        self._add_heading(doc, "4. Выявленные ошибки и план исправления", level=1)
        if not actionable_issues:
            doc.add_paragraph("Критические ошибки и предупреждения не обнаружены.")
        else:
            grouped = {}
            for issue in actionable_issues:
                grouped.setdefault(issue.get("code", "unknown"), []).append(issue)

            for idx, (code, items) in enumerate(grouped.items(), start=1):
                guide = issue_guides.get(code, {
                    "name": items[0].get("title", code),
                    "why": "Проблема влияет на качество мобильной версии.",
                    "impact": "Среднее",
                    "fix": ["Проверить верстку и исправить источник ошибки."],
                })
                self._add_heading(doc, f"4.{idx} {guide['name']}", level=2)
                devices_list = sorted({str(i.get("device", "не указано")) for i in items})
                doc.add_paragraph(f"Серьезность: {items[0].get('severity', 'warning')}")
                doc.add_paragraph(f"Затронуто устройств: {len(devices_list)}")
                doc.add_paragraph(f"Устройства: {', '.join(devices_list)}")
                doc.add_paragraph(f"Почему это важно: {guide['why']}")
                doc.add_paragraph(f"Бизнес-влияние: {guide['impact']}")
                doc.add_paragraph("Что сделать:")
                for step in guide["fix"]:
                    doc.add_paragraph(step, style="List Number")
                doc.add_paragraph("Технические детали из проверки:")
                for example in items[:5]:
                    detail = str(example.get("details", "") or "").strip()
                    if detail:
                        doc.add_paragraph(f"- {detail}")

        self._add_heading(doc, "5. Информационные наблюдения", level=1)
        if info_issues:
            for issue in info_issues:
                info_title = issue.get("title") or "\u041d\u0430\u0431\u043b\u044e\u0434\u0435\u043d\u0438\u0435"
                info_details = issue.get("details", "")
                doc.add_paragraph(
                    f"{info_title}: {info_details}",
                    style="List Bullet",
                )
        else:
            doc.add_paragraph("Дополнительные информационные замечания отсутствуют.")

        doc.add_page_break()
        self._add_heading(doc, "6. Скриншоты проверенных устройств", level=1)
        added = 0
        missing_files = 0
        for d in devices:
            shot = _resolve_mobile_screenshot_path(d)
            if not shot or not os.path.exists(shot):
                if d.get("screenshot_name") or d.get("screenshot_path"):
                    missing_files += 1
                continue
            device_name = d.get("device_name") or "\u0423\u0441\u0442\u0440\u043e\u0439\u0441\u0442\u0432\u043e"
            doc.add_paragraph(
                f"{device_name} | "
                f"Viewport {(d.get('viewport') or {}).get('width', '-')}x{(d.get('viewport') or {}).get('height', '-')} | "
                f"\u0417\u0430\u043c\u0435\u0447\u0430\u043d\u0438\u0439: {d.get('issues_count', 0)}"
            )
            try:
                doc.add_picture(shot, width=Inches(5.8))
                added += 1
            except Exception:
                doc.add_paragraph(f"\u041d\u0435 \u0443\u0434\u0430\u043b\u043e\u0441\u044c \u0432\u0441\u0442\u0440\u043e\u0438\u0442\u044c \u0441\u043a\u0440\u0438\u043d\u0448\u043e\u0442: {shot}")
        if added == 0:
            doc.add_paragraph("Скриншоты отсутствуют.")
            if missing_files > 0:
                doc.add_paragraph(
                    f"\u0424\u0430\u0439\u043b\u044b \u0441\u043a\u0440\u0438\u043d\u0448\u043e\u0442\u043e\u0432 \u043d\u0435 \u043d\u0430\u0439\u0434\u0435\u043d\u044b \u043d\u0430 \u0441\u0435\u0440\u0432\u0435\u0440\u0435: {missing_files}."
                )

        self._add_heading(doc, "7. Итоги", level=1)
        if actionable_issues:
            doc.add_paragraph(
                "Обнаружены ошибки мобильной версии, которые требуют исправления для повышения "
                "качества UX и стабильности SEO-показателей."
            )
            doc.add_paragraph(
                "Рекомендуется выполнить исправления по приоритету (critical -> warning), затем "
                "повторить аудит и сравнить метрики."
            )
        else:
            doc.add_paragraph(
                "Критичных проблем не обнаружено. Мобильная версия сайта соответствует "
                "базовым требованиям удобства и технической корректности."
            )

        filepath = os.path.join(self.reports_dir, f"{task_id}.docx")
        try:
            self._save_document(doc, filepath)
            return filepath
        finally:
            for temp_path in temp_screenshots:
                try:
                    if temp_path and os.path.exists(temp_path):
                        os.remove(temp_path)
                except Exception:
                    pass
    def generate_bot_report(self, task_id: str, data: Dict[str, Any]) -> str:
        """Generate DOCX report for bot accessibility check."""
        doc = Document()
        results = data.get("results", {}) or {}
        summary = results.get("summary", {}) or {}
        blockers = results.get("priority_blockers", []) or []
        playbooks = results.get("playbooks", []) or []
        baseline_diff = results.get("baseline_diff", {}) or {}
        trend = results.get("trend", {}) or {}
        host_consistency = results.get("host_consistency", {}) or {}
        category_stats = results.get("category_stats", []) or []
        alerts = results.get("alerts", []) or []
        robots_linter = ((results.get("robots_linter", {}) or {}).get("findings") or [])
        allowlist_sim = ((results.get("allowlist_simulator", {}) or {}).get("scenarios") or [])
        action_center = ((results.get("action_center", {}) or {}).get("by_owner") or {})
        evidence_pack = ((results.get("evidence_pack", {}) or {}).get("rows") or [])
        batch_runs = results.get("batch_runs", []) or []

        self._add_cover_page(doc, 'Bot Accessibility Report', 'Проверка доступности ботов', data.get('url', ''), datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        self._add_header_footer(doc, 'Bot Accessibility Report')
        doc.add_paragraph(f"URL: {data.get('url', 'н/д')}")
        doc.add_paragraph(f"Сформирован: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Профиль ретраев: {self._format_policy_profile(results.get('retry_profile', 'standard'), 'retry')}")
        doc.add_paragraph(
            f"Профиль критичности: {self._format_policy_profile(results.get('criticality_profile', 'balanced'), 'criticality')} | "
            f"Профиль SLA: {self._format_policy_profile(results.get('sla_profile', 'standard'), 'sla')}"
        )
        doc.add_paragraph(
            f"Режим AI-политики: {'намеренные блокировки AI' if results.get('ai_block_expected') else 'строгая доступность'} "
            f"(ожидаемо заблокировано: {summary.get('expected_ai_policy_blocked', 0)})"
        )

        self._add_heading(doc, "1. Краткая сводка на одной странице", level=1)
        exec_rows = [
            ["Ботов проверено", len(results.get("bots_checked", []) or [])],
            ["Доступно", summary.get("accessible", 0)],
            ["Сканируемо", summary.get("crawlable", 0)],
            ["Рендерится", summary.get("renderable", 0)],
            ["Индексируемо", summary.get("indexable", 0)],
            ["Не индексируемо", summary.get("non_indexable", 0)],
            ["Сигналы WAF/CDN", summary.get("waf_cdn_detected", 0)],
            ["Средний ответ, мс", summary.get("avg_response_time_ms", 0)],
        ]
        self._add_table(doc, ["Метрика", "Значение"], exec_rows)

        business_risk = "Низкий"
        if (summary.get("non_indexable", 0) or 0) > 0 or (summary.get("waf_cdn_detected", 0) or 0) > 0:
            business_risk = "Средний"
        if (summary.get("non_indexable", 0) or 0) >= max(1, int((summary.get("total", 0) or 0) * 0.3)):
            business_risk = "Высокий"
        doc.add_paragraph(f"Бизнес-риск: {business_risk}")

        self._add_heading(doc, "2. Топ блокеров", level=1)
        if blockers:
            blocker_rows = []
            for item in blockers[:10]:
                blocker_rows.append([
                    item.get("code", ""),
                    item.get("title", ""),
                    item.get("affected_bots", 0),
                    item.get("priority_score", 0),
                    ", ".join(item.get("sample_bots", []) or []),
                ])
            self._add_table(doc, ["Код", "Заголовок", "Затронуто", "Приоритет", "Примеры ботов"], blocker_rows)
        else:
            doc.add_paragraph("Приоритетные блокеры не обнаружены.")

        self._add_heading(doc, "3. План спринта (плейбуки)", level=1)
        if playbooks:
            for idx, item in enumerate(playbooks[:8], start=1):
                doc.add_paragraph(
                    f"{idx}. [{item.get('owner', 'Владелец')}] {item.get('title', '')} "
                    f"(приоритет {item.get('priority_score', 0)})"
                )
                for action in (item.get("actions") or [])[:5]:
                    doc.add_paragraph(str(action), style="List Bullet")
        else:
            doc.add_paragraph("Плейбуки для этого запуска не сформированы.")

        self._add_heading(doc, "4. SLA-матрица категорий", level=1)
        if category_stats:
            rows = []
            for c in category_stats:
                rows.append([
                    c.get("category", ""),
                    c.get("indexable", 0),
                    c.get("total", 0),
                    c.get("indexable_pct", 0),
                    c.get("sla_target_pct", 0),
                    "Да" if c.get("sla_met") else "Нет",
                ])
            self._add_table(doc, ["Категория", "Индексируемо", "Всего", "Индексируемо %", "SLA %", "Выполнено"], rows)
        else:
            doc.add_paragraph("Статистика по категориям недоступна.")

        self._add_heading(doc, "5. Консистентность host", level=1)
        doc.add_paragraph(f"Консистентно: {'Да' if host_consistency.get('consistent', True) else 'Нет'}")
        for note in (host_consistency.get("notes") or []):
            doc.add_paragraph(str(note), style="List Bullet")

        self._add_heading(doc, "6. Сравнение с baseline", level=1)
        if baseline_diff.get("has_baseline"):
            rows = [
                [m.get("metric", ""), m.get("current", ""), m.get("baseline", ""), m.get("delta", "")]
                for m in (baseline_diff.get("metrics") or [])
            ]
            if rows:
                self._add_table(doc, ["Метрика", "Текущее", "Базовое", "Дельта"], rows)
            else:
                doc.add_paragraph("Дельты метрик отсутствуют.")
        else:
            doc.add_paragraph(baseline_diff.get("message", "Базовый срез (baseline) не найден."))

        self._add_heading(doc, "7. Trend History", level=1)
        trend_history = trend.get("history", []) or []
        trend_delta = trend.get("delta_vs_previous", {}) or {}
        if trend_history:
            latest = trend.get("latest") or trend_history[0]
            previous = trend.get("previous")
            doc.add_paragraph(
                f"Запусков сохранено для домена: {trend.get('history_count', len(trend_history))}. "
                f"Последний запуск: {latest.get('timestamp', 'н/д')}."
            )
            if previous:
                doc.add_paragraph(f"Предыдущий запуск: {previous.get('timestamp', 'н/д')}.")
                doc.add_paragraph(
                    "Delta vs previous: "
                    f"indexable {trend_delta.get('indexable', 0)}, "
                    f"critical issues {trend_delta.get('critical_issues', 0)}, "
                    f"avg response ms {trend_delta.get('avg_response_time_ms', 0)}."
                )
            rows = []
            for item in trend_history[:10]:
                rows.append([
                    item.get("timestamp", ""),
                    item.get("indexable", 0),
                    item.get("crawlable", 0),
                    item.get("renderable", 0),
                    item.get("avg_response_time_ms", 0),
                    item.get("critical_issues", 0),
                    item.get("warning_issues", 0),
                ])
            self._add_table(
                doc,
                ["Время запуска", "Индексируемо", "Сканируемо", "Рендерится", "Среднее, мс", "Критично", "Предупреждения"],
                rows,
            )
        else:
            doc.add_paragraph("Trend history unavailable.")

        doc.add_page_break()
        self._add_heading(doc, "8. Рекомендации", level=1)
        recs = results.get("recommendations", []) or []
        if recs:
            for rec in recs[:30]:
                doc.add_paragraph(str(rec), style="List Bullet")
        else:
            doc.add_paragraph("Рекомендации отсутствуют.")

        self._add_heading(doc, "9. Alerts", level=1)
        if alerts:
            for a in alerts[:30]:
                doc.add_paragraph(f"[{str(a.get('severity', 'info')).upper()}] {a.get('code', '')}: {a.get('message', '')}", style="List Bullet")
        else:
            doc.add_paragraph("Активные алерты отсутствуют.")

        doc.add_page_break()
        self._add_heading(doc, "10. Линтер Robots Policy", level=1)
        if robots_linter:
            rows = [[str(x.get("severity", "")).upper(), x.get("code", ""), x.get("message", "")] for x in robots_linter[:40]]
            self._add_table(doc, ["Критичность", "Код", "Сообщение"], rows)
        else:
            doc.add_paragraph("Линтер robots не выявил замечаний.")

        self._add_heading(doc, "11. Симулятор allowlist", level=1)
        if allowlist_sim:
            rows = []
            for s in allowlist_sim[:20]:
                rows.append([
                    s.get("category", ""),
                    s.get("affected_bots", 0),
                    s.get("delta_renderable", 0),
                    s.get("delta_indexable", 0),
                    s.get("projected_indexable_pct", 0),
                ])
            self._add_table(doc, ["Категория", "Затронуто", "Дельта рендеринга", "Дельта индексации", "Прогноз %"], rows)
        else:
            doc.add_paragraph("Нет данных симуляции.")

        self._add_heading(doc, "12. Центр действий (по владельцам)", level=1)
        if action_center:
            for owner, rows in action_center.items():
                doc.add_paragraph(str(owner))
                for item in (rows or [])[:8]:
                    doc.add_paragraph(f"{item.get('title', '')} (приоритет {item.get('priority_score', 0)})", style="List Bullet")
        else:
            doc.add_paragraph("Группы действий по владельцам отсутствуют.")

        doc.add_page_break()
        self._add_heading(doc, "13. Пакет доказательств", level=1)
        if evidence_pack:
            rows = []
            for e in evidence_pack[:40]:
                rows.append([
                    e.get("bot", ""),
                    e.get("status", ""),
                    e.get("indexability_reason", ""),
                    f"{'Да' if e.get('waf_detected') else 'Нет'} ({e.get('waf_confidence', 0)})",
                    e.get("waf_reason", ""),
                ])
            self._add_table(doc, ["Бот", "HTTP", "Причина", "WAF", "Причина WAF"], rows)
        else:
            doc.add_paragraph("Строки доказательств отсутствуют.")

        if batch_runs:
            self._add_heading(doc, "14. Пакетные прогоны", level=1)
            rows = []
            for b in batch_runs[:100]:
                rows.append([
                    b.get("url", ""),
                    b.get("indexable", 0),
                    b.get("total", 0),
                    b.get("renderable", 0),
                    b.get("critical_issues", 0),
                    b.get("warning_issues", 0),
                ])
            self._add_table(doc, ["URL", "Индексируемо", "Всего", "Рендерится", "Критично", "Предупреждения"], rows)

        self._add_full_result_appendix(doc, data)
        filepath = os.path.join(self.reports_dir, f"{task_id}.docx")
        self._save_document(doc, filepath)
        return filepath

    def generate_onpage_report(self, task_id: str, data: Dict[str, Any]) -> str:
        """Generate DOCX report for onpage_audit."""
        doc = Document()
        results = data.get("results", {}) or {}
        summary = results.get("summary", {}) or {}
        url = data.get("url", "н/д")
        generated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        self._add_cover_page(doc, 'OnPage Audit Report', 'OnPage-аудит', data.get('url', ''), generated_at)
        self._add_header_footer(doc, 'OnPage Audit Report')

        doc.add_paragraph(f"URL: {url}")
        doc.add_paragraph(f"Сформирован: {generated_at}")
        doc.add_paragraph(f"Engine: {results.get('engine', 'onpage-v1')}")

        self._add_heading(doc, "1. Ключевая сводка", level=1)
        summary_rows = [
            ["Оценка", results.get("score", summary.get("score", 0))],
            ["Spam-оценка", summary.get("spam_score", (results.get("scores", {}) or {}).get("spam_score", 0))],
            ["Оценка покрытия КС", summary.get("keyword_coverage_score", (results.get("scores", {}) or {}).get("keyword_coverage_score", 0))],
            ["Покрытие КС, %", summary.get("keyword_coverage_pct", (results.get("keyword_coverage", {}) or {}).get("coverage_pct", 0))],
            ["AI-риск (композит)", summary.get("ai_risk_composite", (results.get("scores", {}) or {}).get("ai_risk_composite", 0))],
            ["Критические проблемы", summary.get("critical_issues", 0)],
            ["Предупреждения", summary.get("warning_issues", 0)],
            ["Инфо-проблемы", summary.get("info_issues", 0)],
            ["HTTP-статус", results.get("status_code", "н/д")],
            ["Финальный URL", results.get("final_url", url)],
            ["Язык", results.get("language", "auto")],
        ]
        self._add_table(doc, ["Метрика", "Значение"], summary_rows)
        top_risks = sorted((results.get("issues", []) or []), key=lambda x: 0 if x.get("severity") == "critical" else 1)[:5]
        top_actions = (results.get("priority_queue", []) or [])[:5]
        self._add_heading(doc, "1a. Краткая сводка на одной странице", level=2)
        doc.add_paragraph("Топ рисков", style="List Bullet")
        for risk in top_risks:
            doc.add_paragraph(f"{risk.get('severity', '').upper()} | {risk.get('title', '')}", style="List Bullet 2")
        doc.add_paragraph("Топ действий", style="List Bullet")
        for act in top_actions:
                doc.add_paragraph(f"{act.get('bucket', '')}: {act.get('title', '')} (приоритет {act.get('priority_score', 0)})", style="List Bullet 2")

        content = results.get("content", {}) or {}
        content_profile = results.get("content_profile", {}) or {}
        self._add_heading(doc, "2. Метрики контента", level=1)
        content_rows = [
            ["Количество слов", content.get("word_count", 0)],
            ["Уникальные слова", content.get("unique_word_count", 0)],
            ["Количество символов", content.get("char_count", 0)],
            ["Длина чистого текста", content_profile.get("clean_text_length", 0)],
            ["Ядро словаря", content_profile.get("core_vocabulary", 0)],
            ["Водность %", content_profile.get("wateriness_pct", 0)],
            ["Тошнота", content_profile.get("nausea_index", 0)],
            ["Text/HTML %", content_profile.get("text_html_pct", 0)],
        ]
        self._add_table(doc, ["Метрика", "Значение"], content_rows)

        self._add_heading(doc, "3. Мета-теги", level=1)
        title_meta = results.get("title", {}) or {}
        desc_meta = results.get("description", {}) or {}
        h1_meta = results.get("h1", {}) or {}
        meta_rows = [
            ["Title", title_meta.get("text", "")],
            ["Длина title", title_meta.get("length", 0)],
            ["Description", desc_meta.get("text", "")],
            ["Длина description", desc_meta.get("length", 0)],
            ["Количество H1", h1_meta.get("count", 0)],
            ["Значения H1", ", ".join(h1_meta.get("values", []) or [])],
        ]
        self._add_table(doc, ["Поле", "Значение"], meta_rows)

        self._add_heading(doc, "4. Ключевые слова", level=1)
        keyword_rows = []
        for row in (results.get("keywords", []) or [])[:50]:
            keyword_rows.append(
                [
                    row.get("keyword", ""),
                    row.get("occurrences", 0),
                    row.get("density_pct", 0),
                    "Yes" if row.get("in_title") else "No",
                    "Yes" if row.get("in_description") else "No",
                    "Yes" if row.get("in_h1") else "No",
                    str(row.get("status", "ok")).upper(),
                ]
            )
        if keyword_rows:
            self._add_table(doc, ["Ключевое слово", "Частота", "Плотность %", "Title", "Description", "H1", "Статус"], keyword_rows)
        else:
            doc.add_paragraph("Ключевые слова не переданы.")

        self._add_heading(doc, "5. Топ терминов", level=1)
        top_term_rows = []
        for row in (results.get("top_terms", []) or [])[:20]:
            top_term_rows.append([row.get("term", ""), row.get("count", 0), row.get("pct", 0)])
        if top_term_rows:
            self._add_table(doc, ["Термин", "Частота", "Доля %"], top_term_rows)
        else:
            doc.add_paragraph("Топ терминов недоступен.")

        doc.add_page_break()
        self._add_heading(doc, "6. Технические сигналы", level=1)
        technical = results.get("technical", {}) or {}
        technical_rows = [
            ["Canonical href", technical.get("canonical_href", "")],
            ["Canonical self", "Да" if technical.get("canonical_is_self") else "Нет"],
            ["Meta robots", technical.get("robots", "")],
            ["Noindex", "Да" if technical.get("noindex") else "Нет"],
            ["Nofollow", "Да" if technical.get("nofollow") else "Нет"],
            ["Viewport", technical.get("viewport", "")],
            ["HTML lang", technical.get("lang", "")],
            ["Hreflang теги", technical.get("hreflang_count", 0)],
            ["Блоки schema", technical.get("schema_count", 0)],
        ]
        self._add_table(doc, ["Сигнал", "Значение"], technical_rows)

        self._add_heading(doc, "7. Ссылки, медиа и читабельность", level=1)
        links = results.get("links", {}) or {}
        media = results.get("media", {}) or {}
        readability = results.get("readability", {}) or {}
        quality_rows = [
            ["Всего ссылок", links.get("links_total", 0)],
            ["Внутренние ссылки", links.get("internal_links", 0)],
            ["Внешние ссылки", links.get("external_links", 0)],
            ["Nofollow ссылки", links.get("nofollow_links", 0)],
            ["Пустые анкоры", links.get("empty_anchor_links", 0)],
            ["Всего изображений", media.get("images_total", 0)],
            ["Изображения без alt", media.get("images_missing_alt", 0)],
            ["Предложения", readability.get("sentences_count", 0)],
            ["Средняя длина предложения", readability.get("avg_sentence_len", 0)],
            ["Доля длинных предложений", readability.get("long_sentence_ratio", 0)],
            ["Лексическое разнообразие", readability.get("lexical_diversity", 0)],
        ]
        spam_metrics = results.get("spam_metrics", {}) or {}
        quality_rows.extend(
            [
                ["Доля стоп-слов", spam_metrics.get("stopword_ratio", 0)],
                ["Соотношение контента к HTML", spam_metrics.get("content_html_ratio", 0)],
                ["Доля uppercase", spam_metrics.get("uppercase_ratio", 0)],
                ["Доля пунктуации", spam_metrics.get("punctuation_ratio", 0)],
                ["Дубли предложений", spam_metrics.get("duplicate_sentences", 0)],
                ["Доля дублей предложений", spam_metrics.get("duplicate_sentence_ratio", 0)],
            ]
        )
        self._add_table(doc, ["Метрика", "Значение"], quality_rows)

        self._add_heading(doc, "8. N-граммы", level=1)
        ngrams = results.get("ngrams", {}) or {}
        bigrams = ngrams.get("bigrams", []) or []
        if bigrams:
            bigram_rows = [[row.get("term", ""), row.get("count", 0), row.get("pct", 0)] for row in bigrams[:20]]
            self._add_table(doc, ["Биграмма", "Количество", "Доля %"], bigram_rows)
        else:
            doc.add_paragraph("Биграммы недоступны.")
        trigrams = ngrams.get("trigrams", []) or []
        if trigrams:
            trigram_rows = [[row.get("term", ""), row.get("count", 0), row.get("pct", 0)] for row in trigrams[:20]]
            self._add_table(doc, ["Триграмма", "Количество", "Доля %"], trigram_rows)
        else:
            doc.add_paragraph("Триграммы недоступны.")

        doc.add_page_break()
        self._add_heading(doc, "9. Schema и OpenGraph", level=1)
        schema = results.get("schema", {}) or {}
        og = results.get("opengraph", {}) or {}
        schema_rows = [
            ["JSON-LD блоки", schema.get("json_ld_blocks", 0)],
            ["Валидные JSON-LD", schema.get("json_ld_valid_blocks", 0)],
            ["Элементы Microdata", schema.get("microdata_items", 0)],
            ["Элементы RDFa", schema.get("rdfa_items", 0)],
            ["Типы schema", ", ".join([x.get("type", "") for x in (schema.get("types", []) or [])[:10]])],
            ["Теги OpenGraph", og.get("tags_count", 0)],
            ["Обязательные OG найдены", og.get("required_present_count", 0)],
            ["Отсутствующие OG", ", ".join(og.get("required_missing", []) or [])],
        ]
        self._add_table(doc, ["Поле", "Значение"], schema_rows)

        self._add_heading(doc, "10. AI-сигналы", level=1)
        ai = results.get("ai_insights", {}) or {}
        ai_rows = [
            ["Плотность AI-маркеров /1k", ai.get("ai_marker_density_1k", 0)],
            ["Доля хеджирования", ai.get("hedging_ratio", 0)],
            ["Повторяемость шаблонов /1k", ai.get("template_repetition", 0)],
            ["Вариативность (CV)", ai.get("burstiness_cv", 0)],
            ["Прокси perplexity", ai.get("perplexity_proxy", 0)],
            ["Глубина сущностей /1k", ai.get("entity_depth_1k", 0)],
            ["Оценка специфичности утверждений", ai.get("claim_specificity_score", 0)],
            ["Оценка сигнала автора", ai.get("author_signal_score", 0)],
            ["Оценка атрибуции источников", ai.get("source_attribution_score", 0)],
            ["Композитный AI-риск", ai.get("ai_risk_composite", 0)],
        ]
        self._add_table(doc, ["Сигнал", "Значение"], ai_rows)

        link_terms = results.get("link_anchor_terms", []) or []
        if link_terms:
            self._add_heading(doc, "11. Топ терминов анкоров", level=1)
            self._add_table(
                doc,
                ["Термин", "Количество"],
                [[row.get("term", ""), row.get("count", 0)] for row in link_terms[:10]],
            )

        self._add_heading(doc, "11a. Тепловая карта критичности", level=2)
        heatmap = results.get("heatmap", {}) or {}
        heatmap_rows = []
        for cat, payload in heatmap.items():
            heatmap_rows.append([cat, payload.get("score", 0), payload.get("issues", 0), payload.get("critical", 0), payload.get("warning", 0)])
        if heatmap_rows:
            self._add_table(doc, ["Категория", "Оценка", "Проблемы", "Критично", "Предупреждение"], heatmap_rows)

        self._add_heading(doc, "11b. Priority Queue", level=2)
        queue = results.get("priority_queue", []) or []
        if queue:
            queue_rows = [[x.get("bucket", ""), x.get("severity", ""), x.get("code", ""), x.get("title", ""), x.get("priority_score", 0), x.get("effort", 0)] for x in queue[:15]]
            self._add_table(doc, ["Этап", "Критичность", "Код", "Проблема", "Приоритет", "Трудозатраты"], queue_rows)

        self._add_heading(doc, "11c. Цели до/после", level=2)
        targets = results.get("targets", []) or []
        if targets:
            target_rows = [[x.get("metric", ""), x.get("current", 0), x.get("target", 0), x.get("delta", 0)] for x in targets]
            self._add_table(doc, ["Метрика", "Текущее", "Цель", "Дельта"], target_rows)

        doc.add_page_break()
        self._add_heading(doc, "12. Проблемы", level=1)
        issues = results.get("issues", []) or []
        if issues:
            for issue in issues[:80]:
                sev = str(issue.get("severity", "info")).upper()
                title_i = issue.get("title", issue.get("code", "Проблема"))
                details_i = issue.get("details", "")
                doc.add_paragraph(f"[{sev}] {title_i}: {details_i}", style="List Bullet")
        else:
            doc.add_paragraph("Проблемы не обнаружены.")

        self._add_heading(doc, "13. Рекомендации", level=1)
        recs = results.get("recommendations", []) or []
        if recs:
            for rec in recs[:30]:
                doc.add_paragraph(str(rec), style="List Bullet")
        else:
            doc.add_paragraph("Рекомендации недоступны.")

        self._add_full_result_appendix(doc, data)
        filepath = os.path.join(self.reports_dir, f"{task_id}.docx")
        self._save_document(doc, filepath)
        return filepath

    def generate_site_audit_pro_report(self, task_id: str, data: Dict[str, Any]) -> str:
        """Generate compact DOCX report for site_audit_pro."""
        doc = Document()
        title = doc.add_heading("Отчет Site Audit Pro", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        results = data.get("results", {}) or {}
        summary = results.get("summary", {}) or {}
        pipeline = results.get("pipeline", {}) or {}
        pipeline_metrics = pipeline.get("metrics", {}) or {}
        pages = results.get("pages", []) or []
        issues = results.get("issues", []) or []
        url = data.get("url", "н/д")

        doc.add_paragraph(f"URL: {url}")
        doc.add_paragraph(f"Режим: {results.get('mode', 'quick')}")
        doc.add_paragraph(f"Сформирован: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        self._add_heading(doc, "1. Ключевая сводка", level=1)
        summary_rows = [
            ["Всего страниц", summary.get("total_pages", 0)],
            ["Всего проблем", summary.get("issues_total", 0)],
            ["Критично", summary.get("critical_issues", 0)],
            ["Предупреждение", summary.get("warning_issues", 0)],
            ["Инфо", summary.get("info_issues", 0)],
            ["Оценка", summary.get("score", "н/д")],
            ["Средний ответ (мс)", pipeline_metrics.get("avg_response_time_ms", 0)],
            ["Средняя читаемость", pipeline_metrics.get("avg_readability_score", 0)],
            ["Среднее качество ссылок", pipeline_metrics.get("avg_link_quality_score", 0)],
            ["Orphan-страницы", pipeline_metrics.get("orphan_pages", 0)],
        ]
        self._add_table(doc, ["Метрика", "Значение"], summary_rows)

        self._add_heading(doc, "2. Топ проблем", level=1)
        top_issues = issues[:20]
        if top_issues:
            issue_rows = []
            for issue in top_issues:
                issue_rows.append([
                    (issue.get("severity") or "info").upper(),
                    issue.get("code", ""),
                    issue.get("url", ""),
                    issue.get("title", ""),
                ])
            self._add_table(doc, ["Критичность", "Код", "URL", "Проблема"], issue_rows)
        else:
            doc.add_paragraph("Проблемы не обнаружены.")

        self._add_heading(doc, "3. Ключевые сигналы ссылочного графа", level=1)
        top_pr = (pipeline.get("pagerank") or [])[:10]
        if top_pr:
            pr_rows = [[row.get("url", ""), row.get("score", 0)] for row in top_pr]
            self._add_table(doc, ["URL", "PageRank"], pr_rows)
        else:
            doc.add_paragraph("Данные PageRank недоступны.")

        self._add_heading(doc, "4. Кластеры тем", level=1)
        clusters = (pipeline.get("topic_clusters") or [])[:20]
        if clusters:
            cluster_rows = [[c.get("topic", "misc"), c.get("count", 0), ", ".join((c.get("urls") or [])[:3])] for c in clusters]
            self._add_table(doc, ["Тема", "Страницы", "Примеры URL"], cluster_rows)
        else:
            doc.add_paragraph("Кластеры тем недоступны.")

        self._add_heading(doc, "5. Рекомендации", level=1)
        recommendations = [p.get("recommendation") for p in pages if p.get("recommendation")]
        if recommendations:
            for rec in recommendations[:20]:
                doc.add_paragraph(str(rec), style="List Bullet")
        else:
            doc.add_paragraph("Рекомендации недоступны.")

        self._add_full_result_appendix(doc, data)
        filepath = os.path.join(self.reports_dir, f"{task_id}.docx")
        self._save_document(doc, filepath)
        return filepath

    def generate_link_profile_report(self, task_id: str, data: Dict[str, Any]) -> str:
        """Generate DOCX report for link_profile_audit."""
        doc = Document()
        results = data.get("results", {}) or {}
        summary = results.get("summary", {}) or {}
        tables = results.get("tables", {}) or {}
        prompts = results.get("prompts", {}) or {}
        warnings = results.get("warnings", []) or []
        errors = results.get("errors", []) or []
        keywords = results.get("keywords", {}) or {}

        self._add_cover_page(doc, 'Link Profile Report', 'Аудит ссылочного профиля', data.get('url', ''), datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        self._add_header_footer(doc, 'Link Profile Report')

        doc.add_paragraph(f"Домен: {data.get('url', summary.get('our_domain', 'н/д'))}")
        doc.add_paragraph(f"Сформирован: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("Показаны первые 100 строк на таблицу. Все данные в выгрузке XLSX, где обрабатываются строки по лимиту.")

        def _lim(rows: List[Dict[str, Any]], limit: int = 100) -> List[Dict[str, Any]]:
            return (rows or [])[:limit]

        def _add_dict_table(title: str, rows: List[Dict[str, Any]], level: int = 2, limit: int = 100) -> None:
            rows = rows or []
            if not rows:
                return
            self._add_heading(doc, title, level=level)
            limited = _lim(rows, limit)
            cols = list((limited[0] or {}).keys())
            if not cols:
                doc.add_paragraph("Нет данных.")
                return
            values = [[r.get(c, "") for c in cols] for r in limited]
            self._add_table(doc, cols, values)
            if len(limited) < len(rows):
                doc.add_paragraph(
                    f"Показаны первые {len(limited)} строк. Все данные в выгрузке XLSX, где обрабатываются строки по лимиту."
                )

        self._add_heading(doc, "1. Сводка", level=1)
        summary_rows = [
            ["Строк ссылок", summary.get("rows_total", 0)],
            ["Уникальных доноров", summary.get("unique_ref_domains", 0)],
            ["Уникальных конкурентов", summary.get("unique_competitors", 0)],
            ["Ссылок на наш домен", summary.get("our_links", 0)],
            ["Dofollow", summary.get("dofollow", 0)],
            ["Nofollow", summary.get("nofollow", 0)],
            ["Unknown follow", summary.get("unknown_follow", 0)],
            ["Dofollow %", summary.get("dofollow_pct", "н/д")],
            ["Nofollow %", summary.get("nofollow_pct", "н/д")],
            ["Lost links %", summary.get("lost_links_pct", "н/д")],
            ["HTTP 2xx %", summary.get("http_2xx_pct", "н/д")],
            ["Средний DR", summary.get("avg_dr", "н/д")],
        ]
        self._add_table(doc, ["Метрика", "Значение"], summary_rows)

        self._add_heading(doc, "2. Executive", level=1)
        _add_dict_table("2.1 Executive overview", tables.get("executive_overview", []) or [], level=2)
        _add_dict_table("2.2 KPI", tables.get("executive_kpi", []) or [], level=2)
        _add_dict_table("2.3 Приоритеты", tables.get("priority_dashboard", []) or [], level=2)
        _add_dict_table("2.4 Validation checks", tables.get("validation_checks", []) or [], level=2)

        self._add_heading(doc, "3. Competitors", level=1)
        _add_dict_table("3.1 Competitor benchmark", tables.get("competitor_benchmark", []) or [], level=2)
        _add_dict_table("3.2 Рейтинг конкурентов", tables.get("competitor_ranking", []) or [], level=2)
        _add_dict_table("3.3 Качество профиля конкурентов", tables.get("competitor_quality", []) or [], level=2)
        _add_dict_table("3.4 Сырые метрики конкурентов", tables.get("competitor_analysis", []) or [], level=2)

        self._add_heading(doc, "4. Gap-доноры", level=1)
        _add_dict_table("4.1 Gap donors priority", tables.get("gap_donors_priority", []) or [], level=2)
        _add_dict_table("4.2 Donor overlap matrix", tables.get("donor_overlap_matrix", []) or [], level=2)
        _add_dict_table("4.3 Ready-to-buy domains", tables.get("ready_buy_domains", []) or [], level=2)
        _add_dict_table("4.4 Priority score domains", tables.get("priority_score_domains", []) or [], level=2)

        doc.add_page_break()
        self._add_heading(doc, "5. Quality", level=1)
        _add_dict_table("5.1 Link attributes", tables.get("link_attributes", []) or [], level=2)
        _add_dict_table("5.2 HTTP / Type / Lang / Platform", tables.get("http_type_lang_platform", []) or [], level=2)
        _add_dict_table("5.3 Target structure", tables.get("target_structure", []) or [], level=2)
        _add_dict_table("5.4 Follow/Nofollow mix", tables.get("follow_mix_pct", []) or [], level=2)
        _add_dict_table("5.5 Follow domain mix", tables.get("follow_domain_mix_pct", []) or [], level=2)

        self._add_heading(doc, "6. Loss", level=1)
        _add_dict_table("6.1 Loss & recovery", tables.get("loss_recovery", []) or [], level=2)
        _add_dict_table("6.2 Lost status mix", tables.get("lost_status_mix", []) or [], level=2)
        _add_dict_table("6.3 Ссылки с редиректов (sample)", tables.get("raw_redirect_links", []) or [], level=2)

        doc.add_page_break()
        self._add_heading(doc, "7. Anchors", level=1)
        anchor_types = results.get("anchor_breakdown", {}) or {}
        if anchor_types:
            rows = [[k, v] for k, v in anchor_types.items()]
            self._add_table(doc, ["Тип анкоров", "Количество"], rows)
        _add_dict_table("7.1 Anchor summary", tables.get("anchor_analysis", []) or [], level=2)
        _add_dict_table("7.2 Word analysis", tables.get("anchor_word_analysis", []) or [], level=2)
        _add_dict_table("7.3 Anchor mix", tables.get("anchor_mix_pct", []) or [], level=2)

        self._add_heading(doc, "8. Risks", level=1)
        _add_dict_table("8.1 Risk signals", tables.get("risk_signals", []) or [], level=2)
        _add_dict_table("8.2 Дубликаты без нашего сайта", tables.get("raw_duplicates_without_our", []) or [], level=2)

        doc.add_page_break()
        self._add_heading(doc, "9. Plan", level=1)
        _add_dict_table("9.1 План 30/60/90", tables.get("action_queue_90d", []) or [], level=2)
        _add_dict_table("9.2 Очередь действий", tables.get("action_queue", []) or [], level=2)
        for key in ("ourSite", "competitors", "comparison", "plan", "anchorTemplate", "riskTemplate", "outreachTemplate", "rowReviewTemplate"):
            text = prompts.get(key)
            if text:
                doc.add_paragraph(f"{key}: {text}")

        self._add_heading(doc, "10. Рекомендации и предупреждения", level=1)
        if errors:
            self._add_heading(doc, "10.1 Ошибки", level=2)
            for item in errors[:100]:
                doc.add_paragraph(str(item), style="List Bullet")
        if warnings:
            self._add_heading(doc, "10.2 Предупреждения", level=2)
            for item in warnings[:100]:
                doc.add_paragraph(str(item), style="List Bullet")

        derived = keywords.get("derivedBrandKeywords", []) or []
        if derived:
            self._add_heading(doc, "11. Авто-брендовые ключи", level=1)
            doc.add_paragraph(", ".join(map(str, derived[:100])))

        self._add_full_result_appendix(doc, data)
        filepath = os.path.join(self.reports_dir, f"{task_id}.docx")
        self._save_document(doc, filepath)
        return filepath

    def generate_redirect_checker_report(self, task_id: str, data: Dict[str, Any]) -> str:
        """Generate detailed DOCX report for redirect_checker."""
        doc = Document()

        results = data.get("results", {}) or {}
        summary = results.get("summary", {}) or {}
        scenarios = results.get("scenarios", []) or []
        recommendations = results.get("recommendations", []) or []
        selected_ua = results.get("selected_user_agent", {}) or {}
        applied_policy = results.get("applied_policy", {}) or {}
        checked_url = data.get("url") or results.get("checked_url") or "н/д"

        def _status_label(value: Any) -> str:
            status = str(value or "").lower()
            if status == "passed":
                return "Passed"
            if status == "warning":
                return "Warning"
            if status == "error":
                return "Error"
            return "Unknown"

        def _priority_for_status(value: Any) -> str:
            status = str(value or "").lower()
            if status == "error":
                return "P1"
            if status == "warning":
                return "P2"
            return "P3"

        def _sla_for_status(value: Any) -> str:
            status = str(value or "").lower()
            if status == "error":
                return "24-48 часов"
            if status == "warning":
                return "до 7 дней"
            return "планово"

        def _owner_for_key(key: str) -> str:
            mapping = {
                "http_to_https": "DevOps/Backend",
                "www_consistency": "DevOps/Backend",
                "multiple_slashes": "Backend",
                "url_case": "Backend",
                "index_files": "Backend",
                "trailing_slash": "Backend",
                "legacy_extensions": "Backend",
                "canonical_tag": "Frontend+SEO",
                "missing_404": "Backend",
                "redirect_chains": "Backend",
                "user_agent_emulation": "DevOps/Security",
            }
            return mapping.get(str(key or ""), "SEO+Dev")

        def _chain_codes(item: Dict[str, Any]) -> str:
            codes = item.get("response_codes") or []
            if not isinstance(codes, list) or not codes:
                return "-"
            return " -> ".join(str(code) for code in codes)

        def _clip(value: Any, limit: int = 180) -> str:
            text = str(value or "-")
            return text if len(text) <= limit else text[: limit - 1] + "…"

        def _fmt_duration_ms(value: Any) -> str:
            try:
                return f"{int(value or 0)} ms"
            except Exception:
                return "0 ms"

        def _safe_iso_dt(value: Any) -> str:
            raw = str(value or "").strip()
            if not raw:
                return "н/д"
            try:
                parsed = datetime.fromisoformat(raw.replace("Z", "+00:00"))
                return parsed.strftime("%Y-%m-%d %H:%M:%S")
            except Exception:
                return raw

        generated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        self._add_cover_page(doc, 'Redirect Checker Report', 'Проверка редиректов', data.get('url', ''), generated_at)
        self._add_header_footer(doc, 'Redirect Checker Report')

        doc.add_paragraph(f"URL: {checked_url}")
        doc.add_paragraph(f"User-Agent: {selected_ua.get('label') or selected_ua.get('key') or 'н/д'}")
        doc.add_paragraph(
            "Policy: "
            f"host={applied_policy.get('canonical_host_policy', 'auto')}, "
            f"slash={applied_policy.get('trailing_slash_policy', 'auto')}, "
            f"lowercase={'on' if applied_policy.get('enforce_lowercase', True) else 'off'}"
        )
        if applied_policy.get("allowed_query_params"):
            doc.add_paragraph(f"Allowed params: {', '.join(map(str, applied_policy.get('allowed_query_params') or []))}")
        if applied_policy.get("required_query_params"):
            doc.add_paragraph(f"Required params: {', '.join(map(str, applied_policy.get('required_query_params') or []))}")
        doc.add_paragraph(f"Проверено: {_safe_iso_dt(results.get('checked_at'))}")
        doc.add_paragraph(f"Сформирован: {generated_at}")

        total_scenarios = int(summary.get("total_scenarios") or len(scenarios))
        passed_count = int(summary.get("passed") or 0)
        warning_count = int(summary.get("warnings") or 0)
        error_count = int(summary.get("errors") or 0)
        quality_score = summary.get("quality_score", "н/д")
        quality_grade = summary.get("quality_grade", "н/д")
        duration_ms = summary.get("duration_ms", "н/д")

        violations = [
            scenario for scenario in scenarios
            if str(scenario.get("status") or "").lower() in ("warning", "error")
        ]
        violations.sort(
            key=lambda item: (
                0 if str(item.get("status") or "").lower() == "error" else 1,
                int(item.get("id") or 999),
            )
        )

        if error_count > 0:
            executive_status = "Обнаружены критические нарушения редиректов. Нужны исправления в ближайший спринт."
        elif warning_count > 0:
            executive_status = "Критичных ошибок нет, но есть предупреждения. Рекомендуется закрыть технический долг."
        else:
            executive_status = "Нарушения не выявлены. Структура редиректов консистентна."

        self._add_heading(doc, "1. Executive summary", level=1)
        self._add_table(
            doc,
            ["Метрика", "Значение"],
            [
                ["Всего сценариев", total_scenarios],
                ["Passed", passed_count],
                ["Warning", warning_count],
                ["Error", error_count],
                ["Score", f"{quality_score} ({quality_grade})"],
                ["Время проверки", f"{duration_ms} ms"],
            ],
        )
        doc.add_paragraph(executive_status)

        self._add_heading(doc, "2. Детальная матрица сценариев", level=1)
        if scenarios:
            scenario_rows = []
            for item in scenarios:
                scenario_rows.append(
                    [
                        item.get("id", "-"),
                        item.get("key", "-"),
                        item.get("title", "-"),
                        _status_label(item.get("status")),
                        _fmt_duration_ms(item.get("duration_ms")),
                        _chain_codes(item),
                        item.get("hops", 0),
                        _clip(item.get("expected"), 110),
                        _clip(item.get("actual"), 150),
                        _clip(item.get("recommendation"), 140),
                    ]
                )
            self._add_table(
                doc,
                ["#", "Key", "Сценарий", "Статус", "Время", "Коды", "Хопы", "Expected", "Actual", "Рекомендация"],
                scenario_rows,
            )
        else:
            doc.add_paragraph("Сценарии не получены от сервиса Redirect Checker.")

        doc.add_page_break()
        self._add_heading(doc, "3. Нарушения и разбор", level=1)
        if not violations:
            doc.add_paragraph("Все сценарии в статусе Passed. Дополнительные действия не требуются.")
        else:
            for idx, item in enumerate(violations, start=1):
                scenario_title = item.get("title") or item.get("key") or f"Сценарий {idx}"
                self._add_heading(
                    doc,
                    f"3.{idx} [{_status_label(item.get('status')).upper()}] {scenario_title}",
                    level=2,
                )
                doc.add_paragraph(f"Что проверялось: {item.get('what_checked') or '-'}")
                doc.add_paragraph(f"Тестовый URL: {item.get('test_url') or '-'}")
                doc.add_paragraph(f"Время сценария: {_fmt_duration_ms(item.get('duration_ms'))}")
                doc.add_paragraph(f"Expected: {item.get('expected') or '-'}")
                doc.add_paragraph(f"Actual: {item.get('actual') or '-'}")
                doc.add_paragraph(f"Коды ответа: {_chain_codes(item)}")
                doc.add_paragraph(f"Хопы: {item.get('hops', 0)}")
                if item.get("final_url"):
                    doc.add_paragraph(f"Финальный URL: {item.get('final_url')}")
                if item.get("error"):
                    doc.add_paragraph(f"Техническая ошибка: {item.get('error')}")
                if item.get("recommendation"):
                    doc.add_paragraph(f"Рекомендация: {item.get('recommendation')}")

                chain = item.get("chain") or []
                if isinstance(chain, list) and chain:
                    chain_rows = []
                    for pos, hop in enumerate(chain[:12], start=1):
                        chain_rows.append(
                            [
                                pos,
                                hop.get("url", "-"),
                                hop.get("status_code", "-"),
                                hop.get("location", "-") or "-",
                            ]
                        )
                    self._add_table(doc, ["Шаг", "URL", "HTTP", "Location"], chain_rows)

        doc.add_page_break()
        self._add_heading(doc, "4. План действий при нарушениях", level=1)
        if not violations:
            doc.add_paragraph("План действий не требуется: нарушений не найдено.")
        else:
            action_rows = []
            for item in violations:
                key = str(item.get("key") or "")
                action_rows.append(
                    [
                        _priority_for_status(item.get("status")),
                        item.get("title") or key or "-",
                        _owner_for_key(key),
                        _sla_for_status(item.get("status")),
                        item.get("recommendation") or "Свести поведение к каноническому сценарию.",
                    ]
                )
            self._add_table(
                doc,
                ["Приоритет", "Сценарий", "Владелец", "SLA", "Действие"],
                action_rows,
            )

        risk_404_items = [
            item for item in scenarios
            if str(item.get("key") or "").lower() in {"missing_404", "soft_404_detection"}
        ]
        risk_404_items.sort(
            key=lambda item: (
                0 if str(item.get("status") or "").lower() == "error" else 1,
                -int(item.get("duration_ms") or 0),
            )
        )

        self._add_heading(doc, "5. 404 Risks", level=1)
        if not risk_404_items:
            doc.add_paragraph("Риски soft-404 и некорректных 404 не обнаружены.")
        else:
            risk_rows = []
            for item in risk_404_items:
                risk_rows.append(
                    [
                        item.get("key", "-"),
                        item.get("title", "-"),
                        _status_label(item.get("status")),
                        _fmt_duration_ms(item.get("duration_ms")),
                        _clip(item.get("actual"), 140),
                        _clip(item.get("recommendation"), 120),
                    ]
                )
            self._add_table(
                doc,
                ["Key", "Сценарий", "Статус", "Время", "Actual", "Рекомендация"],
                risk_rows,
            )

        self._add_heading(doc, "6. Рекомендации", level=1)
        if recommendations:
            for rec in recommendations[:50]:
                doc.add_paragraph(str(rec), style="List Bullet")
        elif violations:
            for item in violations[:50]:
                rec = str(item.get("recommendation") or "").strip()
                if rec:
                    doc.add_paragraph(rec, style="List Bullet")
        else:
            doc.add_paragraph("Критических рекомендаций нет.")

        doc.add_page_break()
        self._add_heading(doc, "7. Краткие ТЗ на исправление", level=1)
        tz_templates = {
            "http_to_https": {
                "goal": "Все запросы по HTTP должны вести на HTTPS одним постоянным редиректом.",
                "tasks": [
                    "Настроить 301/308 на уровне nginx/apache для всего хоста.",
                    "Проверить отсутствие промежуточных URL в цепочке.",
                    "Обновить внутренние ссылки и sitemap на HTTPS.",
                ],
                "done": "curl -I http://domain.tld возвращает 301/308 на https://domain.tld и далее один финальный 200.",
            },
            "www_consistency": {
                "goal": "Версии с www и без www должны вести на единый канонический хост.",
                "tasks": [
                    "Выбрать канонический host (www или без www).",
                    "Сделать постоянный 301/308 редирект со второй версии.",
                    "Обновить canonical и sitemap на канонический host.",
                ],
                "done": "Оба варианта хоста приводят к одной финальной странице без цепочек 2+.",
            },
            "multiple_slashes": {
                "goal": "URL с // должны нормализоваться до чистого пути.",
                "tasks": [
                    "Добавить rewrite-правило удаления повторных слешей.",
                    "Направлять на нормализованный URL через 301.",
                ],
                "done": "Запросы с // не отдают дубль и ведут на единственный канонический путь.",
            },
            "url_case": {
                "goal": "Единый регистр URL для исключения дублей.",
                "tasks": [
                    "Определить lowercase как стандарт URL.",
                    "Сделать 301 с uppercase-вариантов на lowercase.",
                ],
                "done": "Uppercase URL не индексируются отдельно и ведут на lowercase-версию.",
            },
            "index_files": {
                "goal": "Index-файлы не должны жить как отдельные URL.",
                "tasks": [
                    "Сделать 301 с /index.html (и /index.php при необходимости) на /.",
                    "Проверить внутренние ссылки, чтобы не было ссылок на index-файлы.",
                ],
                "done": "Index URL не отдают 200 как отдельные страницы и не формируют дубли.",
            },
            "trailing_slash": {
                "goal": "Нужна единая политика trailing slash по всему сайту.",
                "tasks": [
                    "Выбрать формат: /page или /page/.",
                    "Настроить 301 с альтернативной формы.",
                    "Привести canonical и внутренние ссылки к выбранному формату.",
                ],
                "done": "Обе версии URL сходятся в одну каноническую без циклов и лишних редиректов.",
            },
            "legacy_extensions": {
                "goal": "Старые .html/.php URL не должны создавать дубли.",
                "tasks": [
                    "Настроить 301 с устаревших расширений на чистые URL.",
                    "Проверить отсутствие 302 и цепочек при миграции.",
                ],
                "done": "Legacy URL корректно переезжают на чистые и не остаются в индексе отдельно.",
            },
            "canonical_tag": {
                "goal": "Каждая страница должна иметь корректный canonical на канонический URL.",
                "tasks": [
                    "Добавить <link rel=\"canonical\"> в <head> шаблона.",
                    "Проверить, что canonical указывает на правильный host/scheme/path.",
                ],
                "done": "Canonical присутствует и соответствует финальному URL страницы.",
            },
            "missing_404": {
                "goal": "Несуществующие страницы должны отдавать 404/410, а не 200 или редирект на главную.",
                "tasks": [
                    "Настроить обработчик 404 в роутинге/веб-сервере.",
                    "Убедиться, что soft-404 отсутствуют.",
                ],
                "done": "Случайный несуществующий URL стабильно возвращает 404 (или осознанный 410).",
            },
            "redirect_chains": {
                "goal": "Сократить цепочки редиректов до одного шага.",
                "tasks": [
                    "Удалить промежуточные узлы A->B->C, оставить A->C.",
                    "Проверить правила на уровне приложения и веб-сервера.",
                ],
                "done": "Максимум 1 редирект до финального URL по ключевым сценариям.",
            },
            "user_agent_emulation": {
                "goal": "Для ключевых ботов должен быть одинаковый канонический результат.",
                "tasks": [
                    "Проверить WAF/rate-limit и исключить блокировку Googlebot/Yandex Bot.",
                    "Синхронизировать редиректы и финальные коды между desktop/mobile/bot.",
                ],
                "done": "Googlebot Desktop, Googlebot Smartphone и Yandex Bot получают консистентные ответы.",
            },
        }

        if not violations:
            doc.add_paragraph("ТЗ не требуется: нарушений не найдено.")
        else:
            unique_violations = {}
            for item in violations:
                key = str(item.get("key") or "")
                if key and key not in unique_violations:
                    unique_violations[key] = item

            for idx, (key, item) in enumerate(unique_violations.items(), start=1):
                template = tz_templates.get(key, {})
                title_tz = item.get("title") or key or f"Сценарий {idx}"
                self._add_heading(doc, f"6.{idx} {title_tz}", level=2)
                doc.add_paragraph(f"Цель: {template.get('goal') or 'Устранить нарушение и привести URL к канонической схеме.'}")
                tasks = template.get("tasks") or []
                if tasks:
                    for task in tasks:
                        doc.add_paragraph(str(task), style="List Bullet")
                else:
                    fallback_task = item.get("recommendation") or "Исправить правило редиректа и проверить повторно."
                    doc.add_paragraph(str(fallback_task), style="List Bullet")
                done_text = template.get("done") or "После исправления сценарий должен перейти в статус Passed."
                doc.add_paragraph(f"Критерий приемки: {done_text}")

        self._add_full_result_appendix(doc, data)
        filepath = os.path.join(self.reports_dir, f"{task_id}.docx")
        self._save_document(doc, filepath)
        return filepath

    def generate_core_web_vitals_report(self, task_id: str, data: Dict[str, Any]) -> str:
        """Generate Core Web Vitals report (single or batch) to DOCX."""
        doc = Document()

        def _as_float(value: Any) -> Any:
            try:
                if value is None:
                    return None
                return float(value)
            except Exception:
                return None

        def _fmt(value: Any, digits: int = 1) -> str:
            num = _as_float(value)
            if num is None:
                return "-"
            if digits <= 0:
                return str(int(round(num)))
            return f"{num:.{digits}f}"

        def _status_label(value: Any) -> str:
            token = str(value or "unknown").strip().lower()
            if token == "good":
                return "GOOD"
            if token == "needs_improvement":
                return "NEEDS IMPROVEMENT"
            if token == "poor":
                return "POOR"
            if token == "error":
                return "ERROR"
            return "UNKNOWN"

        results = data.get("results", {}) or {}
        mode = str(results.get("mode") or "single").strip().lower()
        strategy = str(results.get("strategy") or "").upper() or "-"
        source = str(results.get("source") or "pagespeed_insights_api")
        url = str(data.get("url") or results.get("url") or "-")
        summary = results.get("summary", {}) or {}

        self._add_cover_page(doc, 'Core Web Vitals Report', 'Core Web Vitals', data.get('url', ''), datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        self._add_header_footer(doc, 'Core Web Vitals Report')

        doc.add_paragraph(f"URL: {url}")
        doc.add_paragraph(f"Strategy: {strategy}")
        doc.add_paragraph(f"Mode: {mode}")
        doc.add_paragraph(f"Источник: {source}")
        doc.add_paragraph(f"Сформирован: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        self._add_heading(doc, "1. Executive Summary", level=1)
        if mode == "competitor":
            benchmark = results.get("benchmark", {}) or {}
            self._add_table(
                doc,
                ["Показатель", "Значение"],
                [
                    ["Total URLs", summary.get("total_urls", 0)],
                    ["Success", summary.get("successful_urls", 0)],
                    ["Errors", summary.get("failed_urls", 0)],
                    ["Primary URL", summary.get("primary_url", "-")],
                    ["Primary score", _fmt(summary.get("primary_score"), 1)],
                    ["Primary CWV", _status_label(summary.get("primary_cwv_status"))],
                    ["Primary rank", summary.get("primary_rank", "-")],
                    ["Market leader", summary.get("market_leader_url", "-")],
                    ["Market leader score", _fmt(summary.get("market_leader_score"), 1)],
                    ["Competitor median score", _fmt(benchmark.get("competitor_median_score"), 1)],
                    ["Competitor median LCP (ms)", _fmt(benchmark.get("competitor_median_lcp_ms"), 1)],
                    ["Competitor median INP (ms)", _fmt(benchmark.get("competitor_median_inp_ms"), 1)],
                    ["Competitor median CLS", _fmt(benchmark.get("competitor_median_cls"), 3)],
                ],
            )
        elif mode == "batch" or isinstance(results.get("sites"), list):
            self._add_table(
                doc,
                ["Показатель", "Значение"],
                [
                    ["Total URLs", summary.get("total_urls", 0)],
                    ["Success", summary.get("successful_urls", 0)],
                    ["Errors", summary.get("failed_urls", 0)],
                    ["CWV status", _status_label(summary.get("core_web_vitals_status"))],
                    ["Avg score", _fmt(summary.get("average_performance_score"), 1)],
                    ["Median score", _fmt(summary.get("median_performance_score"), 1)],
                    ["Min score", _fmt(summary.get("min_performance_score"), 1)],
                    ["Max score", _fmt(summary.get("max_performance_score"), 1)],
                ],
            )
        else:
            self._add_table(
                doc,
                ["Показатель", "Значение"],
                [
                    ["Performance score", summary.get("performance_score", "-")],
                    ["CWV status", _status_label(summary.get("core_web_vitals_status"))],
                    ["Health index", summary.get("health_index", "-")],
                    ["Risk level", str(summary.get("risk_level") or "-").upper()],
                    ["Grade", summary.get("grade", "-")],
                ],
            )

        if mode == "competitor":
            primary = results.get("primary", {}) or {}
            primary_summary = primary.get("summary", {}) or {}
            primary_metrics = primary.get("metrics", {}) or {}
            comparison_rows = results.get("comparison_rows", []) or []
            gaps_for_primary = results.get("gaps_for_primary", []) or []
            strengths_of_primary = results.get("strengths_of_primary", []) or []
            common_opportunities = results.get("common_opportunities", []) or []
            recommendations = results.get("recommendations", []) or []
            action_plan = results.get("action_plan", []) or []

            doc.add_page_break()
            self._add_heading(doc, "2. Primary Snapshot", level=1)
            lcp_primary = (primary_metrics.get("lcp") or {}).get("field_value_ms")
            if lcp_primary is None:
                lcp_primary = (primary_metrics.get("lcp") or {}).get("lab_value_ms")
            inp_primary = (primary_metrics.get("inp") or {}).get("field_value_ms")
            if inp_primary is None:
                inp_primary = (primary_metrics.get("inp") or {}).get("lab_value_ms")
            cls_primary = (primary_metrics.get("cls") or {}).get("field_value")
            if cls_primary is None:
                cls_primary = (primary_metrics.get("cls") or {}).get("lab_value")

            self._add_table(
                doc,
                ["Показатель", "Значение"],
                [
                    ["Primary URL", primary.get("url", "-")],
                    ["Primary status", str(primary.get("status") or "-").upper()],
                    ["Primary score", primary_summary.get("performance_score", "-")],
                    ["Primary CWV", _status_label(primary_summary.get("core_web_vitals_status"))],
                    ["Primary LCP (ms)", _fmt(lcp_primary, 1)],
                    ["Primary INP (ms)", _fmt(inp_primary, 1)],
                    ["Primary CLS", _fmt(cls_primary, 3)],
                    ["Primary error", primary.get("error") or "-"],
                ],
            )

            self._add_heading(doc, "3. Competitor Comparison", level=1)
            if comparison_rows:
                rows: List[List[Any]] = []
                for idx, item in enumerate(comparison_rows, start=1):
                    rows.append(
                        [
                            idx,
                            item.get("url", "-"),
                            _status_label(item.get("cwv_status")),
                            _fmt(item.get("score"), 1),
                            _fmt(item.get("lcp_ms"), 1),
                            _fmt(item.get("inp_ms"), 1),
                            _fmt(item.get("cls"), 3),
                            _fmt(item.get("score_delta_vs_primary"), 1),
                            _fmt(item.get("lcp_delta_ms_vs_primary"), 1),
                            _fmt(item.get("inp_delta_ms_vs_primary"), 1),
                            _fmt(item.get("cls_delta_vs_primary"), 3),
                            item.get("top_focus") or item.get("error") or "-",
                        ]
                    )
                self._add_table(
                    doc,
                    [
                        "#",
                        "Competitor URL",
                        "CWV",
                        "Score",
                        "LCP",
                        "INP",
                        "CLS",
                        "Δ Score",
                        "Δ LCP",
                        "Δ INP",
                        "Δ CLS",
                        "Focus",
                    ],
                    rows,
                )
            else:
                doc.add_paragraph("Сравнение с конкурентами отсутствует.")

            doc.add_page_break()
            self._add_heading(doc, "4. Gap & Strength Analysis", level=1)
            doc.add_paragraph("Primary gaps:")
            if gaps_for_primary:
                for rec in gaps_for_primary[:30]:
                    doc.add_paragraph(str(rec), style="List Bullet")
            else:
                doc.add_paragraph("Явных отставаний не обнаружено.")
            doc.add_paragraph("Primary strengths:")
            if strengths_of_primary:
                for rec in strengths_of_primary[:30]:
                    doc.add_paragraph(str(rec), style="List Bullet")
            else:
                doc.add_paragraph("Явные преимущества не зафиксированы.")

            self._add_heading(doc, "5. Common Competitor Issues", level=1)
            if common_opportunities:
                self._add_table(
                    doc,
                    ["Issue", "Group", "Count"],
                    [
                        [
                            item.get("title") or item.get("id") or "-",
                            item.get("group") or "-",
                            item.get("count") or 0,
                        ]
                        for item in common_opportunities[:25]
                    ],
                )
            else:
                doc.add_paragraph("Частые проблемы у конкурентов не выявлены.")

            self._add_heading(doc, "6. Action Plan (Primary)", level=1)
            if action_plan:
                rows: List[List[Any]] = []
                for item in action_plan[:25]:
                    if "affected_urls" in item:
                        rows.append([item.get("priority") or "P2", item.get("action") or "-", item.get("affected_urls") or 0])
                    else:
                        rows.append(
                            [
                                item.get("priority") or "P3",
                                item.get("area") or "-",
                                item.get("owner") or "-",
                                item.get("action") or "-",
                                item.get("expected_impact") or "-",
                            ]
                        )
                if rows and len(rows[0]) == 3:
                    self._add_table(doc, ["Priority", "Action", "Affected URLs"], rows)
                else:
                    self._add_table(doc, ["Priority", "Area", "Owner", "Action", "Expected impact"], rows)
            else:
                doc.add_paragraph("Action plan не сформирован.")

            self._add_heading(doc, "7. Recommendations", level=1)
            if recommendations:
                for rec in recommendations[:50]:
                    doc.add_paragraph(str(rec), style="List Bullet")
            else:
                doc.add_paragraph("Рекомендации не сформированы.")
        elif mode == "batch" or isinstance(results.get("sites"), list):
            sites = results.get("sites", []) or []
            doc.add_page_break()
            self._add_heading(doc, "2. URL Details", level=1)
            detail_rows: List[List[Any]] = []
            for idx, site in enumerate(sites, start=1):
                site_url = str(site.get("url") or "-")
                if str(site.get("status") or "").lower() != "success":
                    detail_rows.append(
                        [
                            idx,
                            site_url,
                            "ERROR",
                            "-",
                            "-",
                            "-",
                            "-",
                            str(site.get("error") or "scan error"),
                        ]
                    )
                    continue
                site_summary = site.get("summary", {}) or {}
                metrics = site.get("metrics", {}) or {}
                lcp = (metrics.get("lcp") or {}).get("field_value_ms")
                if lcp is None:
                    lcp = (metrics.get("lcp") or {}).get("lab_value_ms")
                inp = (metrics.get("inp") or {}).get("field_value_ms")
                if inp is None:
                    inp = (metrics.get("inp") or {}).get("lab_value_ms")
                cls = (metrics.get("cls") or {}).get("field_value")
                if cls is None:
                    cls = (metrics.get("cls") or {}).get("lab_value")
                opportunities = site.get("opportunities", []) or []
                recommendations = site.get("recommendations", []) or []
                focus = "-"
                if opportunities and isinstance(opportunities[0], dict):
                    focus = str(opportunities[0].get("title") or "-")
                elif recommendations:
                    focus = str(recommendations[0] or "-")
                detail_rows.append(
                    [
                        idx,
                        site_url,
                        _status_label(site_summary.get("core_web_vitals_status")),
                        site_summary.get("performance_score", "-"),
                        _fmt(lcp, 0),
                        _fmt(inp, 0),
                        _fmt(cls, 3),
                        focus,
                    ]
                )
            if detail_rows:
                self._add_table(
                    doc,
                    ["#", "URL", "CWV", "Score", "LCP (ms)", "INP (ms)", "CLS", "Priority Focus"],
                    detail_rows,
                )
            else:
                doc.add_paragraph("Детали по URL отсутствуют.")

            common_opportunities = results.get("common_opportunities", []) or []
            self._add_heading(doc, "3. Common Opportunities", level=1)
            if common_opportunities:
                rows = []
                for item in common_opportunities[:20]:
                    rows.append(
                        [
                            item.get("title") or item.get("id") or "-",
                            item.get("group") or "-",
                            item.get("count") or 0,
                            f"{item.get('critical_count') or 0}/{item.get('high_count') or 0}",
                            _fmt(item.get("total_savings_ms"), 1),
                            _fmt(item.get("total_savings_kib"), 1),
                        ]
                    )
                self._add_table(
                    doc,
                    ["Проблема", "Group", "URL count", "Critical/High", "Savings ms", "Savings KiB"],
                    rows,
                )
            else:
                doc.add_paragraph("Частые opportunities не обнаружены.")

            self._add_heading(doc, "4. Batch Action Plan", level=1)
            batch_action_plan = results.get("action_plan", []) or []
            if batch_action_plan:
                self._add_table(
                    doc,
                    ["Priority", "Action", "Affected URLs"],
                    [
                        [
                            item.get("priority") or "P2",
                            item.get("action") or "-",
                            item.get("affected_urls") or 0,
                        ]
                        for item in batch_action_plan[:20]
                    ],
                )
            else:
                doc.add_paragraph("Batch action plan не сформирован.")
        else:
            metrics = results.get("metrics", {}) or {}
            categories = results.get("categories", {}) or {}
            diagnostics = results.get("diagnostics", {}) or {}
            opportunities = results.get("opportunities", []) or []
            action_plan = results.get("action_plan", []) or []
            recommendations = results.get("recommendations", []) or []

            doc.add_page_break()
            self._add_heading(doc, "2. Metrics", level=1)
            metric_rows: List[List[Any]] = []
            for key, label, field_key, lab_key in [
                ("lcp", "LCP", "field_value_ms", "lab_value_ms"),
                ("inp", "INP", "field_value_ms", "lab_value_ms"),
                ("cls", "CLS", "field_value", "lab_value"),
                ("fcp", "FCP", None, "lab_value_ms"),
                ("ttfb", "TTFB", None, "lab_value_ms"),
                ("speed_index", "Speed Index", None, "lab_value_ms"),
                ("tbt", "TBT", None, "lab_value_ms"),
                ("tti", "TTI", None, "lab_value_ms"),
            ]:
                payload = metrics.get(key) or {}
                field_val = payload.get(field_key) if field_key else None
                lab_val = payload.get(lab_key) if lab_key else None
                metric_rows.append(
                    [
                        label,
                        _fmt(field_val, 3 if key == "cls" else 1) if field_val is not None else "-",
                        _fmt(lab_val, 3 if key == "cls" else 1) if lab_val is not None else "-",
                        _status_label(payload.get("status")),
                    ]
                )
            self._add_table(doc, ["Metric", "Field", "Lab", "Status"], metric_rows)

            self._add_heading(doc, "3. Lighthouse Categories", level=1)
            self._add_table(
                doc,
                ["Category", "Score"],
                [
                    ["Performance", categories.get("performance", "-")],
                    ["Accessibility", categories.get("accessibility", "-")],
                    ["Best Practices", categories.get("best_practices", "-")],
                    ["SEO", categories.get("seo", "-")],
                ],
            )

            self._add_heading(doc, "4. Technical Diagnostics", level=1)
            self._add_table(
                doc,
                ["Показатель", "Значение"],
                [
                    ["Requests", diagnostics.get("num_requests", "-")],
                    ["Scripts", diagnostics.get("num_scripts", "-")],
                    ["Stylesheets", diagnostics.get("num_stylesheets", "-")],
                    ["Tasks > 50ms", diagnostics.get("num_tasks_over_50ms", "-")],
                    ["Tasks > 100ms", diagnostics.get("num_tasks_over_100ms", "-")],
                    ["Total byte weight (KiB)", _fmt(diagnostics.get("total_byte_weight_kib"), 1)],
                ],
            )

            doc.add_page_break()
            self._add_heading(doc, "5. Top Opportunities", level=1)
            if opportunities:
                self._add_table(
                    doc,
                    ["Priority", "Title", "Group", "Score", "Savings ms", "Savings KiB"],
                    [
                        [
                            str(item.get("priority") or "medium").upper(),
                            item.get("title") or item.get("id") or "-",
                            item.get("group") or "-",
                            _fmt(item.get("score"), 3),
                            _fmt(item.get("savings_ms"), 1),
                            _fmt(item.get("savings_kib"), 1),
                        ]
                        for item in opportunities[:25]
                    ],
                )
            else:
                doc.add_paragraph("Серьезных opportunities не найдено.")

            self._add_heading(doc, "6. Action Plan", level=1)
            if action_plan:
                self._add_table(
                    doc,
                    ["Priority", "Area", "Owner", "Action", "Expected Impact"],
                    [
                        [
                            item.get("priority") or "P3",
                            item.get("area") or "-",
                            item.get("owner") or "-",
                            item.get("action") or "-",
                            item.get("expected_impact") or "-",
                        ]
                        for item in action_plan[:25]
                    ],
                )
            else:
                doc.add_paragraph("Action plan не сформирован.")

            self._add_heading(doc, "7. Recommendations", level=1)
            if recommendations:
                for rec in recommendations[:40]:
                    doc.add_paragraph(str(rec), style="List Bullet")
            else:
                doc.add_paragraph("Рекомендации не сформированы.")

        self._add_full_result_appendix(doc, data)
        filepath = os.path.join(self.reports_dir, f"{task_id}.docx")
        self._save_document(doc, filepath)
        return filepath

    def generate_report(self, task_id: str, task_type: str, data: Dict[str, Any]) -> str:
        """Генерирует отчет в зависимости от типа задачи."""
        generators = {
            'site_analyze': self.generate_site_analyze_report,
            'robots_check': self.generate_robots_report,
            'sitemap_validate': self.generate_sitemap_report,
            'render_audit': self.generate_render_report,
            'mobile_check': self.generate_mobile_report,
            'bot_check': self.generate_bot_report,
            'site_audit_pro': self.generate_site_audit_pro_report,
            'onpage_audit': self.generate_onpage_report,
            'link_profile_audit': self.generate_link_profile_report,
            'redirect_checker': self.generate_redirect_checker_report,
            'core_web_vitals': self.generate_core_web_vitals_report,
        }
        
        generator = generators.get(task_type, self.generate_site_analyze_report)
        return generator(task_id, data)


# Singleton
docx_generator = DOCXGenerator()

