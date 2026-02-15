"""
SEO Tools API Routes - Full integration with original scripts
"""
from fastapi import APIRouter, BackgroundTasks, HTTPException
from pydantic import BaseModel, field_validator
from typing import Optional, List, Dict, Any
from datetime import datetime
import re
import json
import time
import requests
from urllib.parse import urlparse

router = APIRouter(prefix="/api", tags=["SEO Tools"])

# Redis-based storage for task results
_redis_client = None
_redis_available = True

def get_redis_client():
    global _redis_client, _redis_available
    if _redis_client is None and _redis_available:
        try:
            import redis
            from app.config import settings
            _redis_client = redis.from_url(settings.REDIS_URL, decode_responses=True)
            _redis_client.ping()
            print("[API] Redis connection established for task results")
        except Exception as e:
            print(f"[API] Redis unavailable for task results: {e}")
            _redis_available = False
            _redis_client = None
    return _redis_client

def get_task_result(task_id: str) -> Optional[Dict[str, Any]]:
    """Get task result from Redis or memory fallback"""
    redis_client = get_redis_client()
    if redis_client:
        try:
            data = redis_client.get(f"task:{task_id}")
            if data:
                return json.loads(data)
        except Exception as e:
            print(f"[API] Error getting task from Redis: {e}")
    
    # Fallback to memory (for development without Redis)
    return task_results_memory.get(task_id)

def _save_task_payload(task_id: str, data: Dict[str, Any]) -> None:
    """Persist task payload in Redis (24h TTL) or memory fallback."""
    redis_client = get_redis_client()
    if redis_client:
        try:
            redis_client.setex(f"task:{task_id}", 86400, json.dumps(data))
            return
        except Exception as e:
            print(f"[API] Error saving task in Redis: {e}")
    task_results_memory[task_id] = data

def create_task_result(task_id: str, task_type: str, url: str, result: Dict[str, Any]):
    """Store task result in Redis with 24 hour TTL"""
    data = {
        "task_id": task_id,
        "task_type": task_type,
        "url": url,
        "status": "SUCCESS",
        "progress": 100,
        "status_message": "Completed",
        "error": None,
        "created_at": datetime.utcnow().isoformat(),
        "result": result,
        "completed_at": datetime.utcnow().isoformat()
    }
    
    _save_task_payload(task_id, data)
    print(f"[API] Task {task_id} stored")


def create_task_pending(task_id: str, task_type: str, url: str, status_message: str = "Queued") -> None:
    """Create task record in pending state."""
    now = datetime.utcnow().isoformat()
    data = {
        "task_id": task_id,
        "task_type": task_type,
        "url": url,
        "status": "PENDING",
        "progress": 0,
        "status_message": status_message,
        "error": None,
        "created_at": now,
        "completed_at": None,
        "result": None,
    }
    _save_task_payload(task_id, data)


def update_task_state(
    task_id: str,
    *,
    status: Optional[str] = None,
    progress: Optional[int] = None,
    status_message: Optional[str] = None,
    result: Optional[Dict[str, Any]] = None,
    error: Optional[str] = None,
) -> None:
    """Update task fields while preserving existing payload."""
    task = get_task_result(task_id)
    if not task:
        return
    if status is not None:
        task["status"] = status
    if progress is not None:
        task["progress"] = max(0, min(100, int(progress)))
    if status_message is not None:
        task["status_message"] = status_message
    if result is not None:
        task["result"] = result
    if error is not None:
        task["error"] = error
    if status in ("SUCCESS", "FAILURE"):
        task["completed_at"] = datetime.utcnow().isoformat()
    _save_task_payload(task_id, task)


def append_task_artifact(task_id: str, artifact_path: str, kind: str = "report") -> None:
    """Attach generated artifact path to task payload for future cleanup."""
    task = get_task_result(task_id)
    if not task:
        return
    bucket = task.setdefault("artifacts", {})
    by_kind = bucket.setdefault(kind, [])
    if artifact_path not in by_kind:
        by_kind.append(artifact_path)
    _save_task_payload(task_id, task)


def delete_task_result(task_id: str) -> bool:
    """Delete task result from Redis/memory storage."""
    deleted = False
    redis_client = get_redis_client()
    if redis_client:
        try:
            deleted = bool(redis_client.delete(f"task:{task_id}")) or deleted
        except Exception as e:
            print(f"[API] Error deleting task from Redis: {e}")
    if task_id in task_results_memory:
        del task_results_memory[task_id]
        deleted = True
    return deleted

# Memory fallback storage
task_results_memory = {}

# ============ ORIGINAl ROBOTS AUDIT LOGIC ============
EXPECTED_BOTS = ["googlebot", "yandex", "bingbot"]

SENSITIVE_PATHS = [
    "/admin", "/wp-admin", "/administrator", "/cgi-bin", "/config",
    "/database", "/backup", "/logs", "/phpmyadmin", "/.git", "/.env",
    "/wp-config", "/include", "/tmp", "/temp", "/private", "/secret",
    "/uploads",
]

RECOMMENDATIONS = [
    "Группируйте правила по User-agent для лучшей читаемости",
    "Блокируйте служебные папки: /admin, /tmp, /backup, /.git",
    "Не блокируйте CSS и JS файлы - это мешает сканированию",
    "Используйте Crawl-delay для больших сайтов",
    "Всегда указывайте Sitemap с полным URL",
    "Удаляйте дублирующиеся правила",
    "Избегайте Disallow: / если не хотите полностью заблокировать сайт",
    "Используйте Allow для создания исключений из Disallow",
    "Проверяйте файл в Google Search Console",
    "Учитывайте различия интерпретации директив разными поисковиками",
]

# Google ignores Crawl-delay; keep default recommendations aligned with current search guidance.
RECOMMENDATIONS = [item for item in RECOMMENDATIONS if "Crawl-delay" not in item]


class Rule:
    def __init__(self, user_agent: str, path: str, line: int):
        self.user_agent = user_agent
        self.path = path
        self.line = line


class Group:
    def __init__(self):
        self.user_agents = []
        self.disallow = []
        self.allow = []


class ParseResult:
    def __init__(self):
        self.groups = []
        self.sitemaps = []
        self.crawl_delays = {}
        self.clean_params = []
        self.hosts = []
        self.raw_lines = []
        self.syntax_errors = []
        self.warnings = []
        # For compatibility
        self.all_disallow = []
        self.all_allow = []


from dataclasses import dataclass, field
from typing import Dict, List, Tuple, Any, Optional
from collections import defaultdict
import re


def fetch_robots(url: str, timeout: int = 20) -> Tuple[Optional[str], Optional[int], Optional[str]]:
    """Fetch robots.txt and return (content, status_code, error)"""
    try:
        robots_url = url.rstrip('/') + '/robots.txt'
        resp = requests.get(robots_url, timeout=timeout, headers={"User-Agent": "Mozilla/5.0"})
        return resp.text, resp.status_code, None
    except requests.exceptions.Timeout:
        return None, None, "Timeout"
    except requests.exceptions.ConnectionError:
        return None, None, "Connection Error"
    except Exception as e:
        return None, None, str(e)


def parse_robots(text: str) -> ParseResult:
    """Parse robots.txt content - FULL original implementation"""
    lines = text.splitlines()
    result = ParseResult()
    result.raw_lines = lines
    
    current_group = None
    
    for idx, raw in enumerate(lines, start=1):
        stripped = raw.strip()
        if not stripped or stripped.startswith("#"):
            continue
        if ":" not in stripped:
            err = f"Строка {idx}: Неверный синтаксис - отсутствует ':'"
            result.syntax_errors.append({"line": idx, "error": err, "content": raw})
            result.warnings.append(err)
            continue
        
        key, value = stripped.split(":", 1)
        key = key.strip().lower()
        value = value.strip()
        
        if key == "user-agent":
            if current_group is None or current_group.user_agents:
                current_group = Group()
                result.groups.append(current_group)
            current_group.user_agents.append(value)
            
        elif key == "disallow":
            if not current_group or not current_group.user_agents:
                err = f"Строка {idx}: Disallow без предшествующего User-agent"
                result.syntax_errors.append({"line": idx, "error": err, "content": raw})
                result.warnings.append(err)
                continue
            if value and not value.startswith("/"):
                result.warnings.append(f"Строка {idx}: Путь '{value}' должен начинаться с '/'")
            for ua in current_group.user_agents:
                rule = Rule(ua, value, idx)
                current_group.disallow.append(rule)
                result.all_disallow.append(rule)
                
        elif key == "allow":
            if not current_group or not current_group.user_agents:
                err = f"Строка {idx}: Allow без предшествующего User-agent"
                result.syntax_errors.append({"line": idx, "error": err, "content": raw})
                result.warnings.append(err)
                continue
            if value and not value.startswith("/"):
                result.warnings.append(f"Строка {idx}: Путь '{value}' должен начинаться с '/'")
            for ua in current_group.user_agents:
                rule = Rule(ua, value, idx)
                current_group.allow.append(rule)
                result.all_allow.append(rule)
                
        elif key == "sitemap":
            if value and not value.startswith("http://") and not value.startswith("https://"):
                result.warnings.append(f"Строка {idx}: Sitemap должен содержать полный URL")
            result.sitemaps.append(value)
            
        elif key == "crawl-delay":
            if not current_group or not current_group.user_agents:
                err = f"Строка {idx}: Crawl-delay без предшествующего User-agent"
                result.syntax_errors.append({"line": idx, "error": err, "content": raw})
                result.warnings.append(err)
                continue
            try:
                delay = float(value)
                if delay < 0:
                    raise ValueError("negative")
                for ua in current_group.user_agents:
                    result.crawl_delays[ua] = delay
            except Exception:
                err = f"Строка {idx}: Некорректный Crawl-delay: '{value}'"
                result.syntax_errors.append({"line": idx, "error": err, "content": raw})
                result.warnings.append(err)
                continue
                
        elif key == "clean-param":
            result.clean_params.append(value)
            
        elif key == "host":
            result.hosts.append(value)
            
        else:
            result.warnings.append(f"Строка {idx}: Неизвестная директива '{key}' - будет проигнорирована")
    
    return result


def find_duplicates(all_rules: List[Rule], label: str) -> List[str]:
    """Find duplicate rules"""
    warnings = []
    by_value: Dict[str, List[int]] = defaultdict(list)
    for rule in all_rules:
        key = f"{rule.user_agent.lower()}|{rule.path}"
        by_value[key].append(rule.line)
    for key, line_nos in by_value.items():
        if len(line_nos) > 1:
            ua, path = key.split("|", 1)
            lines = ", ".join(str(n) for n in line_nos)
            warnings.append(f"Дублирующееся правило {label} для {ua}: {path} (строки: {lines})")
    return warnings


def dedupe_keep_order(items: List[str]) -> List[str]:
    """Remove duplicates preserving original order."""
    seen = set()
    out = []
    for item in items:
        if item not in seen:
            seen.add(item)
            out.append(item)
    return out


def build_param_merge_recommendations(result: ParseResult) -> List[str]:
    """Enhanced recommendations for merging parameter rules."""
    recs: List[str] = []
    by_base: Dict[str, List[str]] = defaultdict(list)
    by_query_prefix: Dict[str, List[str]] = defaultdict(list)

    for group in result.groups:
        for rule in group.disallow:
            path = (rule.path or "").strip()
            if not path:
                continue

            if path.endswith("*") and len(path) > 1:
                recs.append(
                    f"Pattern '{path}' can usually be shortened to '{path[:-1]}' (trailing '*' is often redundant)."
                )

            if "?" in path:
                query_prefix = path.split("*", 1)[0]
                if query_prefix.endswith("?"):
                    by_query_prefix[query_prefix].append(path)

            if "?" in path or "=" in path:
                split_idx = min([i for i in [path.find("?"), path.find("=")] if i != -1])
                base = path[:split_idx] if split_idx > 0 else path
                by_base[base].append(path)

    for base, paths in by_base.items():
        uniq = sorted(set(paths))
        if len(uniq) >= 3:
            preview = ", ".join(uniq[:5])
            recs.append(f"Consider merging parameter patterns for '{base}': {preview}")

    for prefix, paths in by_query_prefix.items():
        uniq = sorted(set(paths))
        if len(uniq) >= 3:
            preview = ", ".join(uniq[:6])
            recs.append(
                f"Found {len(uniq)} similar rules: {preview}. "
                f"You can merge them into one rule: {prefix}"
            )

    return dedupe_keep_order(recs)


def validate_sitemaps(sitemaps: List[str], timeout: int = 4, max_checks: int = 5) -> List[Dict[str, Any]]:
    """Validate sitemap URLs declared in robots.txt."""
    checks: List[Dict[str, Any]] = []
    unique_sitemaps = dedupe_keep_order([s for s in sitemaps if isinstance(s, str) and s.strip()])
    for index, sm in enumerate(unique_sitemaps):
        sm = sm.strip()
        if index >= max_checks:
            checks.append({
                "url": sm,
                "ok": None,
                "status_code": None,
                "content_type": None,
                "error": "Skipped (limit reached)"
            })
            continue
        try:
            parsed = urlparse(sm)
            if parsed.scheme not in ("http", "https"):
                checks.append({
                    "url": sm,
                    "ok": False,
                    "status_code": None,
                    "content_type": None,
                    "error": "Invalid sitemap URL scheme"
                })
                continue
            resp = requests.get(sm, timeout=timeout, headers={"User-Agent": "Mozilla/5.0"}, allow_redirects=True)
            content_type = (resp.headers.get("Content-Type") or "").lower()
            looks_xml = any(token in content_type for token in ["xml", "text/plain"]) or resp.text.lstrip().startswith("<?xml")
            ok = resp.status_code == 200 and looks_xml
            checks.append({
                "url": sm,
                "ok": ok,
                "status_code": resp.status_code,
                "content_type": content_type,
                "error": None if ok else "Sitemap not accessible or not XML"
            })
        except Exception as e:
            checks.append({
                "url": sm,
                "ok": False,
                "status_code": None,
                "content_type": None,
                "error": str(e)
            })
    return checks


def build_quality_metrics(
    issues: List[str],
    warnings: List[str],
    syntax_errors: List[Dict[str, Any]],
    missing_bots: List[str],
    sitemap_checks: List[Dict[str, Any]],
    full_block: bool,
    blocked_ext: List[str]
) -> Dict[str, Any]:
    """Build score, grade, production readiness and top fixes."""
    score = 100
    critical_count = len(issues)
    warning_count = len(warnings)
    syntax_count = len(syntax_errors)

    if full_block:
        score -= 70
    if blocked_ext:
        score -= 20
    score -= min(20, syntax_count * 2)
    score -= min(22, warning_count)
    if missing_bots:
        score -= min(8, len(missing_bots) * 2)

    has_sitemap = any(check.get("ok") for check in sitemap_checks) if sitemap_checks else False
    if sitemap_checks and not has_sitemap:
        score -= 10
    if not sitemap_checks:
        score -= 8

    score = max(0, min(100, score))
    if score >= 90:
        grade = "A"
    elif score >= 80:
        grade = "B"
    elif score >= 70:
        grade = "C"
    elif score >= 60:
        grade = "D"
    else:
        grade = "F"

    top_fixes: List[Dict[str, str]] = []
    if full_block:
        top_fixes.append({
            "priority": "critical",
            "title": "Уберите глобальную блокировку сайта",
            "why": "Disallow: / для * блокирует индексацию всего сайта.",
            "action": "Оставьте только точечные запреты для служебных разделов."
        })
    if blocked_ext:
        top_fixes.append({
            "priority": "high",
            "title": "Разблокируйте CSS/JS",
            "why": "Блокировка CSS/JS ухудшает рендеринг для поисковых роботов.",
            "action": "Удалите правила, блокирующие .css и .js."
        })
    if missing_bots:
        top_fixes.append({
            "priority": "medium",
            "title": "Добавьте группы для ключевых ботов",
            "why": "Явные правила для поисковых ботов делают поведение предсказуемым.",
            "action": f"Добавьте User-agent группы для: {', '.join(missing_bots)}."
        })
    if sitemap_checks and not has_sitemap:
        top_fixes.append({
            "priority": "high",
            "title": "Исправьте sitemap URL",
            "why": "Sitemap недоступен или возвращает некорректный контент.",
            "action": "Проверьте URL sitemap в robots.txt и доступность по HTTP 200."
        })
    if syntax_count:
        top_fixes.append({
            "priority": "high",
            "title": "Исправьте синтаксис robots.txt",
            "why": "Синтаксические ошибки могут приводить к игнорированию правил.",
            "action": "Исправьте строки с ошибками в блоке syntax_errors."
        })

    production_ready = score >= 75 and not full_block and not blocked_ext
    return {
        "quality_score": score,
        "quality_grade": grade,
        "production_ready": production_ready,
        "top_fixes": top_fixes[:5],
        "severity_counts": {
            "critical": critical_count,
            "warning": warning_count,
            "info": 0
        }
    }


def build_issues_and_warnings(result: ParseResult) -> Dict[str, Any]:
    """Build issues, warnings, and recommendations - FULL original implementation"""
    issues: List[str] = []
    warnings: List[str] = list(result.warnings)
    
    all_disallow: List[Rule] = []
    all_allow: List[Rule] = []
    for group in result.groups:
        all_disallow.extend(group.disallow)
        all_allow.extend(group.allow)
    
    warnings.extend(find_duplicates(all_disallow, "Disallow"))
    warnings.extend(find_duplicates(all_allow, "Allow"))
    if result.crawl_delays:
        warnings.append("Crawl-delay found: Google ignores this directive.")
    
    if not result.sitemaps:
        warnings.append("Не указана директива Sitemap")
    
    # Check for full site block
    full_block = any(
        rule.user_agent.lower() == "*" and rule.path.strip() == "/"
        for rule in all_disallow
    )
    if full_block:
        issues.append("КРИТИЧНО: Весь сайт заблокирован для всех роботов (Disallow: /)")
    
    # Check blocked extensions
    blocked_ext = []
    for ext in [".css", ".js"]:
        if any(ext in rule.path for rule in all_disallow):
            blocked_ext.append(ext)
    if blocked_ext:
        issues.append(
            "Заблокированы важные ресурсы: "
            + ", ".join(blocked_ext)
            + " - это мешает сканированию"
        )
    
    # Check expected bots
    present_agents = set()
    for group in result.groups:
        for ua in group.user_agents:
            present_agents.add(ua.lower())
    
    missing_bots = [
        bot for bot in EXPECTED_BOTS if not any(bot in ua for ua in present_agents)
    ]
    if missing_bots:
        warnings.append("Рекомендуется добавить правила для: " + ", ".join(missing_bots))
    
    # Check sensitive paths
    unblocked_sensitive = []
    blocked_paths = set(rule.path for rule in all_disallow)
    for path in SENSITIVE_PATHS:
        if path not in blocked_paths:
            unblocked_sensitive.append(path)
    if unblocked_sensitive:
        warnings.append(
            "Рекомендуется заблокировать: "
            + ", ".join(unblocked_sensitive[:8])
        )
    
    # Generate recommendations and quality metrics
    param_recs = build_param_merge_recommendations(result)
    all_recommendations = dedupe_keep_order(RECOMMENDATIONS.copy() + param_recs)
    warnings = dedupe_keep_order(warnings)
    issues = dedupe_keep_order(issues)

    sitemap_checks = validate_sitemaps(result.sitemaps)
    metrics = build_quality_metrics(
        issues=issues,
        warnings=warnings,
        syntax_errors=result.syntax_errors,
        missing_bots=missing_bots,
        sitemap_checks=sitemap_checks,
        full_block=full_block,
        blocked_ext=blocked_ext
    )
    info_issues = [
        f"Обнаружено групп правил: {len(result.groups)}",
        f"Проверено sitemap URL: {len(sitemap_checks)}"
    ]
    if metrics["production_ready"]:
        info_issues.append("Robots.txt готов к продакшн-использованию.")
    else:
        info_issues.append("Требуются правки перед продакшн-использованием.")
    
    return {
        "issues": issues,
        "critical_issues": issues,
        "warnings": warnings,
        "warning_issues": warnings,
        "info_issues": info_issues,
        "recommendations": all_recommendations,
        "param_recommendations": param_recs,
        "present_agents": list(present_agents),
        "missing_bots": missing_bots,
        "sitemaps": result.sitemaps,
        "sitemap_checks": sitemap_checks,
        "crawl_delays": result.crawl_delays,
        "hosts": result.hosts,
        "syntax_errors": result.syntax_errors,
        "quality_score": metrics["quality_score"],
        "quality_grade": metrics["quality_grade"],
        "production_ready": metrics["production_ready"],
        "top_fixes": metrics["top_fixes"],
        "severity_counts": {
            "critical": metrics["severity_counts"]["critical"],
            "warning": metrics["severity_counts"]["warning"],
            "info": len(info_issues)
        },
        "quick_status": "pass" if metrics["production_ready"] else ("fail" if metrics["quality_score"] < 60 else "warn")
    }


def collect_stats(result: ParseResult) -> Dict[str, int]:
    """Collect statistics from parsed result"""
    stats = {
        "user_agents": 0,
        "disallow_rules": 0,
        "allow_rules": 0,
        "sitemaps": len(result.sitemaps),
        "crawl_delays": len(result.crawl_delays),
        "clean_params": len(result.clean_params),
        "hosts": len(result.hosts),
        "lines_count": len(result.raw_lines),
    }
    for group in result.groups:
        stats["user_agents"] += len(group.user_agents)
        stats["disallow_rules"] += len(group.disallow)
        stats["allow_rules"] += len(group.allow)
    return stats


def check_robots_full(url: str) -> Dict[str, Any]:
    """
    FULL robots.txt audit - полная интеграция оригинального скрипта
    Returns complete analysis matching original robots_audit.py
    """
    print(f"[ROBOTS] Starting full audit for: {url}")
    
    raw_text, status_code, error = fetch_robots(url)
    
    if error:
        return {
            "task_type": "robots_check",
            "url": url,
            "completed_at": datetime.utcnow().isoformat(),
            "results": {
                "robots_txt_found": False,
                "status_code": None,
                "content_length": 0,
                "lines_count": 0,
                "user_agents": 0,
                "disallow_rules": 0,
                "allow_rules": 0,
                "sitemaps": [],
                "issues": [f"Ошибка загрузки: {error}"],
                "warnings": [],
                "recommendations": RECOMMENDATIONS,
                "syntax_errors": [],
                "critical_issues": [f"Ошибка загрузки: {error}"],
                "warning_issues": [],
                "info_issues": [],
                "hosts": [],
                "sitemap_checks": [],
                "quality_score": 0,
                "quality_grade": "F",
                "production_ready": False,
                "top_fixes": [{
                    "priority": "critical",
                    "title": "Обеспечьте доступность robots.txt",
                    "why": "Файл robots.txt недоступен, анализ и управление индексацией невозможны.",
                    "action": "Проверьте DNS/SSL/доступность сайта и путь /robots.txt."
                }],
                "severity_counts": {"critical": 1, "warning": 0, "info": 0},
                "error": error,
                "can_continue": False,
                "raw_content": "",
            }
        }
    
    if status_code != 200:
        return {
            "task_type": "robots_check",
            "url": url,
            "completed_at": datetime.utcnow().isoformat(),
            "results": {
                "robots_txt_found": False,
                "status_code": status_code,
                "content_length": len(raw_text) if raw_text else 0,
                "lines_count": len(raw_text.splitlines()) if raw_text else 0,
                "user_agents": 0,
                "disallow_rules": 0,
                "allow_rules": 0,
                "sitemaps": [],
                "issues": [f"Robots.txt не найден (статус: {status_code})"],
                "warnings": ["Создайте файл robots.txt для управления индексацией"],
                "recommendations": RECOMMENDATIONS,
                "syntax_errors": [],
                "critical_issues": [f"Robots.txt не найден (статус: {status_code})"],
                "warning_issues": ["Создайте файл robots.txt для управления индексацией"],
                "info_issues": [],
                "hosts": [],
                "sitemap_checks": [],
                "quality_score": 20 if status_code else 0,
                "quality_grade": "F",
                "production_ready": False,
                "top_fixes": [{
                    "priority": "high",
                    "title": "Создайте robots.txt",
                    "why": "Поисковые системы не получили правила индексации.",
                    "action": "Добавьте файл /robots.txt и укажите sitemap."
                }],
                "severity_counts": {"critical": 1, "warning": 1, "info": 0},
                "error": None,
                "can_continue": True,
                "raw_content": raw_text or "",
            }
        }
    
    # Full parsing and analysis
    result = parse_robots(raw_text)
    stats = collect_stats(result)
    analysis = build_issues_and_warnings(result)
    if len(raw_text.encode("utf-8")) > 512000:
        analysis["warnings"] = dedupe_keep_order(
            analysis["warnings"] + ["robots.txt is larger than 500 KiB; Google ignores content after this limit."]
        )
        analysis["warning_issues"] = analysis["warnings"]
    
    # Build detailed response
    response = {
        "task_type": "robots_check",
        "url": url,
        "completed_at": datetime.utcnow().isoformat(),
        "results": {
            "robots_txt_found": True,
            "status_code": status_code,
            "content_length": len(raw_text),
            "lines_count": stats["lines_count"],
            "user_agents": stats["user_agents"],
            "disallow_rules": stats["disallow_rules"],
            "allow_rules": stats["allow_rules"],
            "sitemaps": analysis["sitemaps"],
            "issues": analysis["issues"],
            "critical_issues": analysis.get("critical_issues", analysis["issues"]),
            "warnings": analysis["warnings"],
            "warning_issues": analysis.get("warning_issues", analysis["warnings"]),
            "info_issues": analysis.get("info_issues", []),
            "recommendations": analysis["recommendations"],
            "syntax_errors": analysis["syntax_errors"],
            "hosts": analysis["hosts"],
            "crawl_delays": analysis["crawl_delays"],
            "clean_params": result.clean_params,
            "param_recommendations": analysis.get("param_recommendations", []),
            "present_agents": analysis["present_agents"],
            "missing_bots": analysis.get("missing_bots", []),
            "sitemap_checks": analysis.get("sitemap_checks", []),
            "quality_score": analysis.get("quality_score", 0),
            "quality_grade": analysis.get("quality_grade", "F"),
            "production_ready": analysis.get("production_ready", False),
            "top_fixes": analysis.get("top_fixes", []),
            "severity_counts": analysis.get("severity_counts", {"critical": len(analysis["issues"]), "warning": len(analysis["warnings"]), "info": 0}),
            "quick_status": analysis.get("quick_status", "warn"),
            "machine_summary": {
                "user_agents_count": stats["user_agents"],
                "disallow_count": stats["disallow_rules"],
                "allow_count": stats["allow_rules"],
                "sitemap_count": len(analysis.get("sitemaps", [])),
                "critical_count": len(analysis.get("critical_issues", analysis["issues"])),
                "warning_count": len(analysis.get("warning_issues", analysis["warnings"])),
                "score": analysis.get("quality_score", 0),
                "grade": analysis.get("quality_grade", "F"),
                "production_ready": analysis.get("production_ready", False)
            },
            "error": None,
            "can_continue": True,
            "raw_content": raw_text,
            # Detailed groups for UI
            "groups_detail": [
                {
                    "user_agents": group.user_agents,
                    "disallow": [{"path": r.path, "line": r.line} for r in group.disallow],
                    "allow": [{"path": r.path, "line": r.line} for r in group.allow],
                }
                for group in result.groups
            ],
        }
    }
    
    print(f"[ROBOTS] Audit completed: {stats['user_agents']} UAs, {stats['disallow_rules']} disallow rules")
    
    return response


def check_robots_simple(url: str) -> Dict[str, Any]:
    """Simplified robots.txt analysis (legacy fallback)."""
    import requests
    
    robots_url = url.rstrip('/') + '/robots.txt'
    
    try:
        response = requests.get(robots_url, timeout=10, headers={"User-Agent": "Mozilla/5.0"})
        raw_text = response.text
        
        result = parse_robots(raw_text)
        analysis = build_issues_and_warnings(result)
        
        return {
            "task_type": "robots_check",
            "url": url,
            "completed_at": datetime.utcnow().isoformat(),
            "results": {
                "robots_txt_found": response.status_code == 200,
                "status_code": response.status_code,
                "content_length": len(raw_text),
                "lines_count": len(raw_text.splitlines()),
                "user_agents": len(analysis["present_agents"]),
                "disallow_rules": len(result.all_disallow),
                "allow_rules": len(result.all_allow),
                "sitemaps": analysis["sitemaps"],
                "issues": analysis["issues"],
                "critical_issues": analysis.get("critical_issues", analysis["issues"]),
                "warnings": analysis["warnings"],
                "warning_issues": analysis.get("warning_issues", analysis["warnings"]),
                "info_issues": analysis.get("info_issues", []),
                "recommendations": analysis["recommendations"],
                "syntax_errors": analysis["syntax_errors"],
                "hosts": analysis["hosts"],
                "sitemap_checks": analysis.get("sitemap_checks", []),
                "quality_score": analysis.get("quality_score", 0),
                "quality_grade": analysis.get("quality_grade", "F"),
                "production_ready": analysis.get("production_ready", False),
                "top_fixes": analysis.get("top_fixes", []),
                "severity_counts": analysis.get("severity_counts", {"critical": len(analysis["issues"]), "warning": len(analysis["warnings"]), "info": 0}),
                "quick_status": analysis.get("quick_status", "warn"),
                "raw_content": raw_text,
            }
        }
    except Exception as e:
        return {
            "task_type": "robots_check",
            "url": url,
            "completed_at": datetime.utcnow().isoformat(),
            "results": {
                "error": str(e),
                "robots_txt_found": False,
                "quality_score": 0,
                "quality_grade": "F",
                "production_ready": False,
                "critical_issues": [str(e)],
                "warning_issues": [],
                "info_issues": [],
                "top_fixes": [],
                "severity_counts": {"critical": 1, "warning": 0, "info": 0}
            }
        }


def check_sitemap_full(url: str) -> Dict[str, Any]:
    """Full sitemap validation with sitemap index traversal and URL export."""
    import xml.etree.ElementTree as ET

    def local_name(tag: str) -> str:
        if not tag:
            return ""
        return tag.split("}", 1)[1] if "}" in tag else tag

    def find_child_text(node: ET.Element, child_name: str) -> str:
        for child in list(node):
            if local_name(child.tag).lower() == child_name.lower():
                return (child.text or "").strip()
        return ""

    def is_http_url(value: str) -> bool:
        try:
            p = urlparse(value)
            return p.scheme in ("http", "https") and bool(p.netloc)
        except Exception:
            return False

    def is_valid_lastmod(value: str) -> bool:
        if not value:
            return True
        date_only = re.fullmatch(r"\d{4}-\d{2}-\d{2}", value)
        dt_utc = re.fullmatch(r"\d{4}-\d{2}-\d{2}T\d{2}:\d{2}:\d{2}Z", value)
        dt_tz = re.fullmatch(r"\d{4}-\d{2}-\d{2}T\d{2}:\d{2}:\d{2}(?:\+|-)\d{2}:\d{2}", value)
        return bool(date_only or dt_utc or dt_tz)

    max_sitemaps = 30
    max_export_urls = 100000
    export_chunk_size = 25000
    max_urls_preview_per_sitemap = 2000
    max_file_size = 52428800
    max_urls_per_sitemap = 50000
    queue: List[str] = [url]
    visited: set = set()
    sitemap_files: List[Dict[str, Any]] = []
    all_urls: List[str] = []
    seen_urls: set = set()
    url_first_seen_in: Dict[str, str] = {}
    duplicate_urls_count = 0
    duplicate_details: List[Dict[str, str]] = []
    duplicate_details_truncated = False
    max_duplicate_details = 500
    invalid_urls_count = 0
    invalid_lastmod_count = 0
    invalid_changefreq_count = 0
    invalid_priority_count = 0
    warnings: List[str] = []
    errors: List[str] = []
    allowed_changefreq = {"always", "hourly", "daily", "weekly", "monthly", "yearly", "never"}
    root_status_code = None

    try:
        session = requests.Session()
        session.headers.update({"User-Agent": "Mozilla/5.0"})

        while queue and len(visited) < max_sitemaps:
            sitemap_url = queue.pop(0).strip()
            if not sitemap_url or sitemap_url in visited:
                continue
            visited.add(sitemap_url)

            file_report: Dict[str, Any] = {
                "sitemap_url": sitemap_url,
                "ok": False,
                "status_code": None,
                "type": "unknown",
                "size_bytes": 0,
                "urls_count": 0,
                "duplicate_count": 0,
                "duplicate_urls": [],
                "urls_omitted": 0,
                "errors": [],
                "warnings": [],
                "urls": [],
            }

            try:
                response = session.get(sitemap_url, timeout=20, allow_redirects=True)
                file_report["status_code"] = response.status_code
                file_report["size_bytes"] = len(response.content or b"")
                if root_status_code is None:
                    root_status_code = response.status_code

                if response.status_code != 200:
                    file_report["errors"].append(f"HTTP {response.status_code}")
                    sitemap_files.append(file_report)
                    continue

                if file_report["size_bytes"] > max_file_size:
                    file_report["warnings"].append("File is larger than 50 MiB.")

                try:
                    root = ET.fromstring(response.content)
                except ET.ParseError as parse_error:
                    file_report["errors"].append(f"XML parse error: {parse_error}")
                    sitemap_files.append(file_report)
                    continue

                root_tag = local_name(root.tag).lower()
                file_report["type"] = root_tag

                if root_tag == "sitemapindex":
                    child_count = 0
                    for sm_node in root.iter():
                        if local_name(sm_node.tag).lower() != "sitemap":
                            continue
                        loc = find_child_text(sm_node, "loc")
                        if not loc:
                            file_report["warnings"].append("sitemap entry without <loc>.")
                            continue
                        if not is_http_url(loc):
                            file_report["warnings"].append(f"Invalid child sitemap URL: {loc}")
                            continue
                        child_count += 1
                        if loc not in visited and loc not in queue and (len(visited) + len(queue) < max_sitemaps):
                            queue.append(loc)
                    if child_count == 0:
                        file_report["warnings"].append("Sitemap index has no child sitemaps.")
                    file_report["ok"] = len(file_report["errors"]) == 0

                elif root_tag == "urlset":
                    file_urls: List[str] = []
                    file_duplicate_urls: List[str] = []
                    file_duplicate_occurrences = 0
                    for url_node in root.iter():
                        if local_name(url_node.tag).lower() != "url":
                            continue
                        loc = find_child_text(url_node, "loc")
                        if not loc:
                            file_report["warnings"].append("url entry without <loc>.")
                            continue
                        if not is_http_url(loc):
                            invalid_urls_count += 1
                            file_report["warnings"].append(f"Invalid URL in <loc>: {loc}")
                            continue

                        lastmod = find_child_text(url_node, "lastmod")
                        if lastmod and not is_valid_lastmod(lastmod):
                            invalid_lastmod_count += 1

                        changefreq = find_child_text(url_node, "changefreq").lower()
                        if changefreq and changefreq not in allowed_changefreq:
                            invalid_changefreq_count += 1

                        priority_raw = find_child_text(url_node, "priority")
                        if priority_raw:
                            try:
                                priority_value = float(priority_raw)
                                if priority_value < 0 or priority_value > 1:
                                    invalid_priority_count += 1
                            except Exception:
                                invalid_priority_count += 1

                        file_urls.append(loc)
                        if loc in seen_urls:
                            duplicate_urls_count += 1
                            file_duplicate_occurrences += 1
                            file_duplicate_urls.append(loc)
                            first_sitemap = url_first_seen_in.get(loc, "")
                            if len(duplicate_details) < max_duplicate_details:
                                duplicate_details.append({
                                    "url": loc,
                                    "first_sitemap": first_sitemap,
                                    "duplicate_sitemap": sitemap_url
                                })
                            else:
                                duplicate_details_truncated = True
                        else:
                            seen_urls.add(loc)
                            url_first_seen_in[loc] = sitemap_url
                            if len(all_urls) < max_export_urls:
                                all_urls.append(loc)

                    file_report["urls_count"] = len(file_urls)
                    file_report["urls"] = file_urls[:max_urls_preview_per_sitemap]
                    file_report["urls_omitted"] = max(0, len(file_urls) - max_urls_preview_per_sitemap)
                    file_report["duplicate_count"] = file_duplicate_occurrences
                    file_report["duplicate_urls"] = dedupe_keep_order(file_duplicate_urls)[:200]
                    if len(file_urls) > max_urls_per_sitemap:
                        file_report["warnings"].append("More than 50,000 URLs in one sitemap file.")
                    if file_report["urls_omitted"] > 0:
                        file_report["warnings"].append(
                            f"URLs preview truncated: {file_report['urls_omitted']} omitted in API response."
                        )
                    file_report["ok"] = len(file_report["errors"]) == 0

                else:
                    file_report["errors"].append(f"Unsupported XML root tag: {root_tag}")

                sitemap_files.append(file_report)

            except Exception as fetch_error:
                file_report["errors"].append(str(fetch_error))
                sitemap_files.append(file_report)

        if queue:
            warnings.append(f"Sitemap traversal limit reached: {max_sitemaps} files.")

        errors.extend([f"{item['sitemap_url']}: {err}" for item in sitemap_files for err in item.get("errors", [])])
        warnings.extend([f"{item['sitemap_url']}: {warn}" for item in sitemap_files for warn in item.get("warnings", [])])

        if invalid_lastmod_count:
            warnings.append(f"Invalid lastmod format found: {invalid_lastmod_count}.")
        if invalid_changefreq_count:
            warnings.append(f"Invalid changefreq values found: {invalid_changefreq_count}.")
        if invalid_priority_count:
            warnings.append(f"Invalid priority values found: {invalid_priority_count}.")
        if invalid_urls_count:
            warnings.append(f"Invalid URLs found in sitemap: {invalid_urls_count}.")
        if duplicate_details_truncated:
            warnings.append(f"Duplicate details were truncated to {max_duplicate_details} entries.")

        recommendations = [
            "Use only absolute HTTP/HTTPS URLs in <loc>.",
            "Keep each sitemap under 50,000 URLs and 50 MiB uncompressed.",
            "Use W3C date format for <lastmod>.",
        ]
        if duplicate_urls_count > 0:
            recommendations.append("Remove duplicate URLs across sitemap files.")

        valid_files = sum(1 for item in sitemap_files if item.get("ok"))
        total_urls_discovered = sum(item.get("urls_count", 0) for item in sitemap_files if item.get("type") == "urlset")
        urls_export_truncated = total_urls_discovered > len(all_urls)
        export_parts_count = (len(all_urls) + export_chunk_size - 1) // export_chunk_size if all_urls else 0

        return {
            "task_type": "sitemap_validate",
            "url": url,
            "completed_at": datetime.utcnow().isoformat(),
            "results": {
                "valid": len(errors) == 0 and len(sitemap_files) > 0,
                "status_code": root_status_code,
                "urls_count": total_urls_discovered,
                "unique_urls_count": len(seen_urls),
                "duplicate_urls_count": duplicate_urls_count,
                "duplicate_details": duplicate_details,
                "duplicate_details_truncated": duplicate_details_truncated,
                "sitemaps_scanned": len(sitemap_files),
                "sitemaps_valid": valid_files,
                "errors": dedupe_keep_order(errors),
                "warnings": dedupe_keep_order(warnings),
                "recommendations": recommendations,
                "sitemap_files": sitemap_files,
                "export_urls": all_urls,
                "urls_export_truncated": urls_export_truncated,
                "max_export_urls": max_export_urls,
                "export_chunk_size": export_chunk_size,
                "export_parts_count": export_parts_count,
                "invalid_lastmod_count": invalid_lastmod_count,
                "invalid_changefreq_count": invalid_changefreq_count,
                "invalid_priority_count": invalid_priority_count,
                "invalid_urls_count": invalid_urls_count,
                "size": sum(item.get("size_bytes", 0) for item in sitemap_files),
            }
        }
    except Exception as e:
        return {
            "task_type": "sitemap_validate",
            "url": url,
            "completed_at": datetime.utcnow().isoformat(),
            "results": {
                "valid": False,
                "error": str(e),
                "urls_count": 0,
                "export_urls": [],
                "sitemap_files": [],
            }
        }


def check_bots_full(
    url: str,
    selected_bots: Optional[List[str]] = None,
    bot_groups: Optional[List[str]] = None,
) -> Dict[str, Any]:
    """Bot accessibility check with feature-flagged v2 engine."""
    from app.config import settings

    engine = (getattr(settings, "BOT_CHECK_ENGINE", "legacy") or "legacy").lower()
    has_custom_selection = bool(selected_bots or bot_groups)
    if has_custom_selection:
        engine = "v2"
    if engine == "v2":
        try:
            from app.tools.bots.service_v2 import BotAccessibilityServiceV2

            checker = BotAccessibilityServiceV2(
                timeout=getattr(settings, "BOT_CHECK_TIMEOUT", 15),
                max_workers=getattr(settings, "BOT_CHECK_MAX_WORKERS", 10),
            )
            return checker.run(url, selected_bots=selected_bots, bot_groups=bot_groups)
        except Exception as e:
            print(f"[API] bot v2 failed, fallback to legacy: {e}")
            legacy = _check_bots_legacy(url)
            legacy_results = legacy.get("results", {})
            legacy_results["engine"] = "legacy-fallback"
            legacy_results["engine_error"] = str(e)
            legacy_results["selected_bots_ignored"] = selected_bots or []
            legacy_results["selected_groups_ignored"] = bot_groups or []
            return legacy

    return _check_bots_legacy(url)


def _check_bots_legacy(url: str) -> Dict[str, Any]:
    import requests

    bots = [
        ("Googlebot", "Mozilla/5.0 (compatible; Googlebot/2.1; +http://www.google.com/bot.html)"),
        ("YandexBot", "Mozilla/5.0 (compatible; YandexBot/3.0; +http://yandex.com/bots)"),
        ("Bingbot", "Mozilla/5.0 (compatible; bingbot/2.0; +http://www.bing.com/bingbot.htm)"),
        ("DuckDuckBot", "DuckDuckBot/1.0; (+https://duckduckgo.com/duckbot)"),
        ("GPTBot", "Mozilla/5.0 AppleWebKit/537.36 (KHTML, like Gecko; compatible; GPTBot/1.0; +https://openai.com/gptbot)"),
    ]
    
    results = {}
    for bot_name, user_agent in bots:
        try:
            resp = requests.get(url, headers={"User-Agent": user_agent}, timeout=10)
            results[bot_name] = {
                "status": resp.status_code,
                "accessible": resp.status_code == 200,
                "response_time": resp.elapsed.total_seconds() if hasattr(resp, 'elapsed') else None
            }
        except Exception as e:
            results[bot_name] = {"error": str(e), "accessible": False}
    
    return {
        "task_type": "bot_check",
        "url": url,
        "completed_at": datetime.utcnow().isoformat(),
        "results": {
            "engine": "legacy",
            "bots_checked": [b[0] for b in bots],
            "bot_results": results,
            "summary": {
                "total": len(bots),
                "accessible": sum(1 for r in results.values() if r.get("accessible")),
            }
        }
    }


# ============ REQUEST MODELS ============
class RobotsCheckRequest(BaseModel):
    url: str

class SitemapValidateRequest(BaseModel):
    url: str

class BotCheckRequest(BaseModel):
    url: str
    selected_bots: Optional[List[str]] = None
    bot_groups: Optional[List[str]] = None

    @field_validator("selected_bots", "bot_groups", mode="before")
    @classmethod
    def _normalize_list_fields(cls, value):
        if value is None:
            return None
        if isinstance(value, str):
            v = value.strip()
            return [v] if v else []
        if isinstance(value, list):
            return [str(item).strip() for item in value if str(item).strip()]
        return value


# ============ API ENDPOINTS ============

@router.post("/tasks/robots-check")
async def create_robots_check(data: RobotsCheckRequest):
    """Full robots.txt analysis"""
    url = data.url
    
    print(f"[API] Full robots.txt analysis for: {url}")
    
    result = check_robots_full(url)
    task_id = f"robots-{datetime.now().timestamp()}"
    create_task_result(task_id, "robots_check", url, result)
    
    return {
        "task_id": task_id,
        "status": "SUCCESS",
        "message": "Robots.txt analysis completed"
    }


class ExportRequest(BaseModel):
    task_id: str


@router.post("/export/robots")
async def export_robots_word(data: ExportRequest):
    """Export robots.txt analysis to Word document - pass task_id in JSON body"""
    import re
    
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io
        
        task_id = data.task_id
        task = get_task_result(task_id)
        
        if not task:
            return {"error": "Задача не найдена", "task_id": task_id}
        
        task_result = task.get("result", {})
        result = task_result.get("results", task_result)
        url = task.get("url", "") or task_result.get("url", "")
        
        # Debug log
        import logging
        logging.info(f"DOCX export for task {task_id}, result keys: {result.keys()}")
        
        doc = Document()
        doc.add_heading('Отчет анализа Robots.txt', level=0)
        
        # Info section
        doc.add_heading('Информация о проверке', level=1)
        now = datetime.now().strftime("%d.%m.%Y, %H:%M:%S")
        
        info_table = doc.add_table(rows=10, cols=2)
        info_table.style = 'Table Grid'
        
        # Handle different data formats for user_agents
        user_agents_str = "Не найдено"
        if isinstance(result.get("user_agents"), list):
            user_agents_str = ", ".join(result.get("user_agents", []))
        elif isinstance(result.get("user_agents"), (int, float)):
            user_agents_str = f"{result.get('user_agents', 0)} шт."
        else:
            user_agents_str = str(result.get("user_agents", "Не найдено"))
        
        info_data = [
            ("URL сайта:", url),
            ("Дата проверки:", now),
            ("Robots.txt найден:", "Да" if result.get("robots_txt_found") else "Нет"),
            ("Статус HTTP:", str(result.get("status_code", "н/д"))),
            ("Размер файла:", f"{result.get('content_length', 0)} байт"),
            ("Кол-во строк:", str(result.get("lines_count", 0))),
            ("User-Agents:", user_agents_str),
        ]
        
        info_data.extend([
            ("SEO-оценка:", str(result.get("quality_score", "н/д"))),
            ("Класс оценки:", str(result.get("quality_grade", "н/д"))),
            ("Готовность к продакшену:", "Да" if result.get("production_ready") else "Нет"),
        ])

        for i, (label, value) in enumerate(info_data):
            info_table.rows[i].cells[0].text = label
            info_table.rows[i].cells[1].text = str(value)
        
        # Stats
        doc.add_heading('Статистика', level=1)
        stats_table = doc.add_table(rows=7, cols=2)
        stats_table.style = 'Table Grid'
        
        stats_data = [
            ("User-Agents:", user_agents_str),
            ("Правил Disallow:", str(result.get("disallow_rules", result.get("disallow_count", result.get("disallow", 0))))),
            ("Правил Allow:", str(result.get("allow_rules", result.get("allow_count", result.get("allow", 0))))),
            ("Sitemaps:", str(len(result.get("sitemaps", [])))),
        ]
        
        stats_data.extend([
            ("Critical:", str((result.get("severity_counts") or {}).get("critical", len(result.get("critical_issues", result.get("issues", [])))))),
            ("Warning:", str((result.get("severity_counts") or {}).get("warning", len(result.get("warning_issues", result.get("warnings", [])))))),
            ("Info:", str((result.get("severity_counts") or {}).get("info", len(result.get("info_issues", []))))),
        ])

        for i, (label, value) in enumerate(stats_data):
            stats_table.rows[i].cells[0].text = label
            stats_table.rows[i].cells[1].text = str(value)
        
        # Issues
        issues = result.get("critical_issues", result.get("issues", []))
        if issues:
            doc.add_heading('Проблемы', level=1)
            for issue in issues:
                p = doc.add_paragraph()
                run = p.add_run("⚠️ " + issue)
                run.font.color.rgb = RGBColor(255, 0, 0)
        
        # Warnings
        warnings = result.get("warning_issues", result.get("warnings", []))
        if warnings:
            doc.add_heading('Предупреждения', level=1)
            for warning in warnings:
                doc.add_paragraph("• " + warning, style='List Bullet')
        
        # Sitemaps
        sitemaps = result.get("sitemaps", [])
        if sitemaps:
            doc.add_heading('Sitemaps', level=1)
            for sm in sitemaps:
                doc.add_paragraph("• " + sm, style='List Bullet')

        sitemap_checks = result.get("sitemap_checks", [])
        if sitemap_checks:
            doc.add_heading('Sitemap URL Checks', level=1)
            for check in sitemap_checks:
                status = "OK" if check.get("ok") else ("SKIPPED" if check.get("ok") is None else "FAIL")
                suffix = f" [HTTP {check.get('status_code')}]" if check.get("status_code") else ""
                line = f"{status}: {check.get('url', '')}{suffix}"
                if check.get("error"):
                    line += f" - {check.get('error')}"
                doc.add_paragraph("• " + line, style='List Bullet')
        
        # Groups
        groups = result.get("groups_detail", [])
        if groups:
            doc.add_heading('Группы правил', level=1)
            for i, group in enumerate(groups, 1):
                doc.add_heading(f'Группа {i}', level=2)
                doc.add_paragraph(f"User-Agents: {', '.join(group.get('user_agents', []))}")
                
                disallow = group.get("disallow", [])
                if disallow:
                    doc.add_paragraph(f"Disallow ({len(disallow)}):", style='List Bullet')
                    for d in disallow[:20]:  # Limit to 20
                        doc.add_paragraph(f"  • {d.get('path', '')} (строка {d.get('line', '')})", style='List Bullet')
        
        # Recommendations
        doc.add_heading('Рекомендации', level=1)
        recommendations = result.get("recommendations", [])
        for rec in recommendations:
            doc.add_paragraph("• " + rec, style='List Bullet')
        
        top_fixes = result.get("top_fixes", [])
        if top_fixes:
            doc.add_heading('Top Fixes', level=1)
            for fix in top_fixes:
                title = fix.get("title", "Fix")
                priority = fix.get("priority", "medium").upper()
                why = fix.get("why", "")
                action = fix.get("action", "")
                doc.add_paragraph(f"[{priority}] {title}", style='List Bullet')
                if why:
                    doc.add_paragraph(f"Why: {why}")
                if action:
                    doc.add_paragraph(f"Action: {action}")
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        from fastapi.responses import Response
        return Response(
            content=buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=robots_report_{task_id}.docx"}
        )
        
    except ImportError:
        return {"error": "python-docx not installed"}
    except Exception as e:
        return {"error": str(e)}


@router.post("/export/sitemap-xlsx")
async def export_sitemap_xlsx(data: ExportRequest):
    """Export sitemap validation report to XLSX."""
    import os
    import re
    from fastapi.responses import Response
    from app.reports.xlsx_generator import xlsx_generator

    try:
        task_id = data.task_id
        task = get_task_result(task_id)
        if not task:
            return {"error": "Задача не найдена", "task_id": task_id}

        task_type = task.get("task_type")
        if task_type != "sitemap_validate":
            return {"error": f"Неподдерживаемый тип задачи для экспорта sitemap XLSX: {task_type}"}

        task_result = task.get("result", {})
        url = task.get("url", "") or task_result.get("url", "")
        report_payload = {
            "url": url,
            "results": task_result.get("results", task_result)
        }
        filepath = xlsx_generator.generate_sitemap_report(task_id, report_payload)
        if not filepath or not os.path.exists(filepath):
            return {"error": "Не удалось сформировать отчет"}

        with open(filepath, "rb") as f:
            content = f.read()

        domain = re.sub(r"[^a-zA-Z0-9._-]+", "_", (urlparse(url).netloc or "site"))
        timestamp = datetime.utcnow().strftime("%Y-%m-%d_%H-%M")
        filename = f"sitemap_report_{domain}_{timestamp}.xlsx"
        return Response(
            content=content,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        return {"error": str(e)}


@router.post("/export/bot-xlsx")
async def export_bot_xlsx(data: ExportRequest):
    """Export bot check report to XLSX."""
    import os
    import re
    from fastapi.responses import Response
    from app.reports.xlsx_generator import xlsx_generator

    try:
        task_id = data.task_id
        task = get_task_result(task_id)
        if not task:
            return {"error": "Задача не найдена", "task_id": task_id}

        task_type = task.get("task_type")
        if task_type != "bot_check":
            return {"error": f"Неподдерживаемый тип задачи для экспорта bot XLSX: {task_type}"}

        task_result = task.get("result", {})
        url = task.get("url", "") or task_result.get("url", "")
        report_payload = {
            "url": url,
            "results": task_result.get("results", task_result),
        }
        filepath = xlsx_generator.generate_bot_report(task_id, report_payload)
        if not filepath or not os.path.exists(filepath):
            return {"error": "Не удалось сформировать отчет"}

        with open(filepath, "rb") as f:
            content = f.read()

        domain = re.sub(r"[^a-zA-Z0-9._-]+", "_", (urlparse(url).netloc or "site"))
        timestamp = datetime.utcnow().strftime("%Y-%m-%d_%H-%M")
        filename = f"bot_report_{domain}_{timestamp}.xlsx"
        return Response(
            content=content,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )
    except Exception as e:
        return {"error": str(e)}


@router.post("/export/mobile-docx")
async def export_mobile_docx(data: ExportRequest):
    """Export mobile check report to DOCX."""
    import os
    import re
    from fastapi.responses import Response
    from app.reports.docx_generator import docx_generator

    try:
        task_id = data.task_id
        task = get_task_result(task_id)
        if not task:
            return {"error": "Задача не найдена", "task_id": task_id}

        task_type = task.get("task_type")
        if task_type != "mobile_check":
            return {"error": f"Неподдерживаемый тип задачи для экспорта mobile DOCX: {task_type}"}

        task_result = task.get("result", {})
        url = task.get("url", "") or task_result.get("url", "")
        report_payload = {
            "url": url,
            "results": task_result.get("results", task_result),
        }
        filepath = docx_generator.generate_mobile_report(task_id, report_payload)
        if not filepath or not os.path.exists(filepath):
            return {"error": "Не удалось сформировать отчет"}
        append_task_artifact(task_id, filepath, kind="export")

        with open(filepath, "rb") as f:
            content = f.read()

        domain = re.sub(r"[^a-zA-Z0-9._-]+", "_", (urlparse(url).netloc or "site"))
        timestamp = datetime.utcnow().strftime("%Y-%m-%d_%H-%M")
        filename = f"mobile_report_{domain}_{timestamp}.docx"
        return Response(
            content=content,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )
    except Exception as e:
        return {"error": str(e)}


@router.post("/export/render-docx")
async def export_render_docx(data: ExportRequest):
    """Export render audit report to DOCX."""
    import os
    import re
    from fastapi.responses import Response
    from app.reports.docx_generator import docx_generator

    try:
        task_id = data.task_id
        task = get_task_result(task_id)
        if not task:
            return {"error": "Задача не найдена", "task_id": task_id}

        task_type = task.get("task_type")
        if task_type != "render_audit":
            return {"error": f"Неподдерживаемый тип задачи для экспорта отчета рендер-аудита (DOCX): {task_type}"}

        task_result = task.get("result", {})
        url = task.get("url", "") or task_result.get("url", "")
        report_payload = {
            "url": url,
            "results": task_result.get("results", task_result),
        }
        filepath = docx_generator.generate_render_report(task_id, report_payload)
        if not filepath or not os.path.exists(filepath):
            return {"error": "Не удалось сформировать отчет"}
        append_task_artifact(task_id, filepath, kind="export")

        with open(filepath, "rb") as f:
            content = f.read()

        domain = re.sub(r"[^a-zA-Z0-9._-]+", "_", (urlparse(url).netloc or "site"))
        timestamp = datetime.utcnow().strftime("%Y-%m-%d_%H-%M")
        filename = f"render_report_{domain}_{timestamp}.docx"
        return Response(
            content=content,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )
    except Exception as e:
        return {"error": str(e)}


@router.post("/export/site-audit-pro-docx")
async def export_site_audit_pro_docx(data: ExportRequest):
    """Export Site Audit Pro report to DOCX."""
    import os
    import re
    from fastapi.responses import Response
    from app.reports.docx_generator import docx_generator
    from app.config import settings

    try:
        if not getattr(settings, "SITE_AUDIT_PRO_ENABLED", True):
            return {"error": "Site Audit Pro is disabled by feature flag"}

        task_id = data.task_id
        task = get_task_result(task_id)
        if not task:
            return {"error": "Задача не найдена", "task_id": task_id}

        task_type = task.get("task_type")
        if task_type != "site_audit_pro":
            return {"error": f"Неподдерживаемый тип задачи для экспорта site_audit_pro DOCX: {task_type}"}

        task_result = task.get("result", {})
        url = task.get("url", "") or task_result.get("url", "")
        report_payload = {
            "url": url,
            "results": task_result.get("results", task_result),
        }
        filepath = docx_generator.generate_site_audit_pro_report(task_id, report_payload)
        if not filepath or not os.path.exists(filepath):
            return {"error": "Не удалось сформировать отчет"}
        append_task_artifact(task_id, filepath, kind="export")

        with open(filepath, "rb") as f:
            content = f.read()

        domain = re.sub(r"[^a-zA-Z0-9._-]+", "_", (urlparse(url).netloc or "site"))
        timestamp = datetime.utcnow().strftime("%Y-%m-%d_%H-%M")
        filename = f"site_audit_pro_{domain}_{timestamp}.docx"
        return Response(
            content=content,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )
    except Exception as e:
        return {"error": str(e)}


@router.post("/export/render-xlsx")
async def export_render_xlsx(data: ExportRequest):
    """Export render issues to XLSX only if issues exist."""
    import os
    import re
    from fastapi.responses import Response
    from app.reports.xlsx_generator import xlsx_generator

    try:
        task_id = data.task_id
        task = get_task_result(task_id)
        if not task:
            return {"error": "Задача не найдена", "task_id": task_id}

        task_type = task.get("task_type")
        if task_type != "render_audit":
            return {"error": f"Неподдерживаемый тип задачи для экспорта отчета рендер-аудита (XLSX): {task_type}"}

        task_result = task.get("result", {})
        url = task.get("url", "") or task_result.get("url", "")
        results = task_result.get("results", task_result) or {}
        issues_count = results.get("issues_count")
        if issues_count is None:
            issues_count = len(results.get("issues", []) or [])
        if issues_count <= 0:
            return {"error": "Проблемы не найдены, XLSX-отчет не формируется"}

        report_payload = {"url": url, "results": results}
        filepath = xlsx_generator.generate_render_report(task_id, report_payload)
        if not filepath or not os.path.exists(filepath):
            return {"error": "Не удалось сформировать отчет"}
        append_task_artifact(task_id, filepath, kind="export")

        with open(filepath, "rb") as f:
            content = f.read()

        domain = re.sub(r"[^a-zA-Z0-9._-]+", "_", (urlparse(url).netloc or "site"))
        timestamp = datetime.utcnow().strftime("%Y-%m-%d_%H-%M")
        filename = f"render_issues_{domain}_{timestamp}.xlsx"
        return Response(
            content=content,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )
    except Exception as e:
        return {"error": str(e)}


@router.post("/export/mobile-xlsx")
async def export_mobile_xlsx(data: ExportRequest):
    """Export mobile issues report to XLSX only if issues exist."""
    import os
    import re
    from fastapi.responses import Response
    from app.reports.xlsx_generator import xlsx_generator

    try:
        task_id = data.task_id
        task = get_task_result(task_id)
        if not task:
            return {"error": "Задача не найдена", "task_id": task_id}

        task_type = task.get("task_type")
        if task_type != "mobile_check":
            return {"error": f"Неподдерживаемый тип задачи для экспорта mobile XLSX: {task_type}"}

        task_result = task.get("result", {})
        url = task.get("url", "") or task_result.get("url", "")
        results = task_result.get("results", task_result) or {}
        issues_count = results.get("issues_count")
        if issues_count is None:
            issues_count = len(results.get("issues", []) or [])
        if issues_count <= 0:
            return {"error": "Проблемы не найдены, XLSX-отчет не формируется"}

        report_payload = {"url": url, "results": results}
        filepath = xlsx_generator.generate_mobile_report(task_id, report_payload)
        if not filepath or not os.path.exists(filepath):
            return {"error": "Не удалось сформировать отчет"}
        append_task_artifact(task_id, filepath, kind="export")

        with open(filepath, "rb") as f:
            content = f.read()

        domain = re.sub(r"[^a-zA-Z0-9._-]+", "_", (urlparse(url).netloc or "site"))
        timestamp = datetime.utcnow().strftime("%Y-%m-%d_%H-%M")
        filename = f"mobile_issues_{domain}_{timestamp}.xlsx"
        return Response(
            content=content,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )
    except Exception as e:
        return {"error": str(e)}


@router.post("/export/site-audit-pro-xlsx")
async def export_site_audit_pro_xlsx(data: ExportRequest):
    """Export Site Audit Pro report to XLSX."""
    import os
    import re
    from fastapi.responses import Response
    from app.reports.xlsx_generator import xlsx_generator
    from app.config import settings

    try:
        if not getattr(settings, "SITE_AUDIT_PRO_ENABLED", True):
            return {"error": "Site Audit Pro is disabled by feature flag"}

        task_id = data.task_id
        task = get_task_result(task_id)
        if not task:
            return {"error": "Задача не найдена", "task_id": task_id}

        task_type = task.get("task_type")
        if task_type != "site_audit_pro":
            return {"error": f"Неподдерживаемый тип задачи для экспорта site_audit_pro XLSX: {task_type}"}

        task_result = task.get("result", {})
        url = task.get("url", "") or task_result.get("url", "")
        results = task_result.get("results", task_result) or {}
        report_payload = {"url": url, "results": results}

        filepath = xlsx_generator.generate_site_audit_pro_report(task_id, report_payload)
        if not filepath or not os.path.exists(filepath):
            return {"error": "Не удалось сформировать отчет"}
        append_task_artifact(task_id, filepath, kind="export")

        with open(filepath, "rb") as f:
            content = f.read()

        domain = re.sub(r"[^a-zA-Z0-9._-]+", "_", (urlparse(url).netloc or "site"))
        timestamp = datetime.utcnow().strftime("%Y-%m-%d_%H-%M")
        filename = f"site_audit_pro_{domain}_{timestamp}.xlsx"
        return Response(
            content=content,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )
    except Exception as e:
        return {"error": str(e)}


@router.get("/mobile-artifacts/{task_id}/{filename}")
async def get_mobile_artifact(task_id: str, filename: str):
    """Serve mobile screenshot artifact for UI gallery."""
    from pathlib import Path
    from fastapi.responses import FileResponse

    try:
        task = get_task_result(task_id)
        if not task:
            return {"error": "Задача не найдена", "task_id": task_id}
        task_result = task.get("result", {})
        results = task_result.get("results", task_result) or {}
        for item in results.get("device_results", []) or []:
            if item.get("screenshot_name") == filename:
                shot_path = item.get("screenshot_path")
                if shot_path and Path(shot_path).exists():
                    return FileResponse(shot_path)
        return {"error": "Артефакт не найден"}
    except Exception as e:
        return {"error": str(e)}


@router.get("/render-artifacts/{task_id}/{filename}")
async def get_render_artifact(task_id: str, filename: str):
    """Serve render audit screenshot artifact for UI gallery."""
    from pathlib import Path
    from fastapi.responses import FileResponse

    try:
        task = get_task_result(task_id)
        if not task:
            return {"error": "Задача не найдена", "task_id": task_id}
        task_result = task.get("result", {})
        results = task_result.get("results", task_result) or {}
        variants = results.get("variants", []) or []
        for variant in variants:
            for shot in (variant.get("screenshots", {}) or {}).values():
                if isinstance(shot, dict) and shot.get("name") == filename:
                    shot_path = shot.get("path")
                    if shot_path and Path(shot_path).exists():
                        return FileResponse(shot_path)
        return {"error": "Артефакт не найден"}
    except Exception as e:
        return {"error": str(e)}


@router.get("/site-pro-artifacts/{task_id}/{filename}")
async def get_site_pro_artifact(task_id: str, filename: str):
    """Serve Site Audit Pro chunk artifact files."""
    from pathlib import Path
    from fastapi.responses import FileResponse
    from app.config import settings

    try:
        task = get_task_result(task_id)
        if not task:
            return {"error": "Задача не найдена", "task_id": task_id}
        task_result = task.get("result", {})
        results = task_result.get("results", task_result) or {}
        artifacts = results.get("artifacts", {}) or {}
        chunk_manifest = artifacts.get("chunk_manifest", {}) or {}
        chunks = chunk_manifest.get("chunks", []) or []
        for chunk in chunks:
            for file_meta in (chunk.get("files") or []):
                if file_meta.get("filename") != filename:
                    continue
                file_path = file_meta.get("path")
                if not file_path:
                    continue
                p = Path(file_path)
                if not p.exists():
                    continue
                reports_root = Path(settings.REPORTS_DIR).resolve()
                resolved = p.resolve()
                if reports_root not in resolved.parents and resolved != reports_root:
                    continue
                return FileResponse(str(resolved), media_type="application/x-ndjson", filename=filename)
        return {"error": "Артефакт не найден"}
    except Exception as e:
        return {"error": str(e)}


@router.get("/site-pro-artifacts/{task_id}/manifest")
async def get_site_pro_artifact_manifest(task_id: str):
    """Return Site Audit Pro chunk manifest and compact payload meta."""
    try:
        task = get_task_result(task_id)
        if not task:
            return {"error": "Задача не найдена", "task_id": task_id}
        task_result = task.get("result", {})
        results = task_result.get("results", task_result) or {}
        artifacts = results.get("artifacts", {}) or {}
        return {
            "task_id": task_id,
            "payload_compacted": bool(artifacts.get("payload_compacted", False)),
            "inline_limits": artifacts.get("inline_limits", {}),
            "omitted_counts": artifacts.get("omitted_counts", {}),
            "chunk_manifest": artifacts.get("chunk_manifest", {}),
        }
    except Exception as e:
        return {"error": str(e)}


@router.post("/tasks/sitemap-validate")
async def create_sitemap_validate(data: SitemapValidateRequest):
    """Full sitemap validation"""
    url = data.url
    
    print(f"[API] Full sitemap validation for: {url}")
    
    result = check_sitemap_full(url)
    task_id = f"sitemap-{datetime.now().timestamp()}"
    create_task_result(task_id, "sitemap_validate", url, result)
    
    return {
        "task_id": task_id,
        "status": "SUCCESS",
        "message": "Sitemap validation completed"
    }


@router.post("/tasks/bot-check")
async def create_bot_check(data: BotCheckRequest):
    """Full bot accessibility check"""
    url = data.url
    
    print(f"[API] Full bot check for: {url}")
    
    result = check_bots_full(url, selected_bots=data.selected_bots, bot_groups=data.bot_groups)
    task_id = f"bots-{datetime.now().timestamp()}"
    create_task_result(task_id, "bot_check", url, result)
    
    return {
        "task_id": task_id,
        "status": "SUCCESS",
        "message": "Bot check completed"
    }


@router.get("/tasks/{task_id}")
async def get_task_status(task_id: str):
    """Get task result"""
    print(f"[API] Getting status for: {task_id}")
    
    data = get_task_result(task_id)
    if data:
        return {
            "task_id": task_id,
            "status": data.get("status", "SUCCESS"),
            "progress": data.get("progress", 100),
            "status_message": data.get("status_message", ""),
            "task_type": data.get("task_type"),
            "url": data.get("url", ""),
            "created_at": data.get("created_at"),
            "completed_at": data.get("completed_at"),
            "result": data.get("result"),
            "error": data.get("error"),
            "can_continue": False
        }
    
    return {
        "task_id": task_id,
        "status": "PENDING",
        "progress": 0,
        "status_message": "Задача пока не найдена",
        "task_type": "site_analyze",
        "url": "",
        "created_at": datetime.utcnow(),
        "completed_at": None,
        "result": None,
        "error": "Задача не найдена",
        "can_continue": False
    }


@router.delete("/tasks/{task_id}")
async def delete_task(task_id: str):
    """Delete task result and linked artifact files."""
    from app.core.task_cleanup import delete_task_artifacts

    task = get_task_result(task_id)
    if not task:
        return {"task_id": task_id, "deleted": False, "error": "Задача не найдена"}

    cleanup = delete_task_artifacts(task)
    removed = delete_task_result(task_id)
    return {
        "task_id": task_id,
        "deleted": bool(removed),
        "artifacts_cleanup": cleanup,
    }


@router.post("/tasks/cleanup-stale-artifacts")
async def cleanup_stale_artifacts(days: Optional[int] = None):
    """Trigger stale report artifacts cleanup under REPORTS_DIR."""
    from app.core.task_cleanup import prune_stale_report_artifacts

    summary = prune_stale_report_artifacts(max_age_days=days)
    return {
        "status": "SUCCESS",
        "cleanup": summary,
    }


@router.get("/rate-limit")
async def get_rate_limit():
    """Rate limit info"""
    return {
        "allowed": True,
        "remaining": 999,
        "reset_in": 3600,
        "limit": 10
    }


@router.get("/celery-status")
async def celery_status():
    """Check celery status"""
    return {"celery_available": False, "mode": "synchronous"}


# ============ ADDITIONAL TOOLS ============

def check_site_full(input_url: str, max_pages: int = 20) -> Dict[str, Any]:
    """
    Full site analysis - максимально приближено к оригинальному seopro.py
    Краулинг, анализ контента, технологии, ссылки, редиректы
    """
    import requests
    from bs4 import BeautifulSoup
    from urllib.parse import urljoin, urlparse
    from collections import defaultdict, Counter
    import re
    import time
    
    url = input_url  # Use input_url as url for compatibility
    
    session = requests.Session()
    session.headers.update({
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
        'Accept-Language': 'ru-RU,ru;q=0.9,en-US;q=0.8,en;q=0.7',
    })
    session.verify = False
    
    parsed_base = urlparse(url)
    base_domain = parsed_base.netloc.lower()
    base_normalized = base_domain[4:] if base_domain.startswith('www.') else base_domain
    
    visited = set()
    to_visit = [(url, 0)]
    all_pages_data = []
    broken_links = []
    external_links = []
    internal_links = []
    all_urls = set()
    redirects = []
    sitemap_urls = set()
    
    # Для подсчёта внутренних ссылок (упрощённый page authority)
    internal_links_count = defaultdict(int)
    
    def is_internal(url_to_check: str) -> bool:
        try:
            parsed = urlparse(url_to_check)
            netloc = parsed.netloc.lower()
            netloc_norm = netloc[4:] if netloc.startswith('www.') else netloc
            return netloc_norm == base_normalized or netloc_norm.endswith('.' + base_normalized)
        except:
            return False
    
    def normalize_url(url: str, drop_query: bool = True) -> str:
        if not url:
            return ''
        try:
            parsed = urlparse(url)
            scheme = (parsed.scheme or '').lower()
            netloc = parsed.netloc or ''
            if netloc.endswith(':80'):
                netloc = netloc[:-3]
            if netloc.endswith(':443'):
                netloc = netloc[:-4]
            path = parsed.path or '/'
            if path != '/':
                path = path.rstrip('/')
            query = '' if drop_query else parsed.query
            return f"{scheme}://{netloc}{path}" + (f"?{query}" if query else '')
        except:
            return url.strip()
    
    def is_valid_page(url_to_check: str) -> bool:
        try:
            parsed = urlparse(url_to_check)
            path = parsed.path.lower()
            exts = ['.jpg', '.jpeg', '.png', '.gif', '.svg', '.webp', '.bmp', '.ico', '.pdf', 
                    '.doc', '.docx', '.xls', '.xlsx', '.zip', '.rar', '.mp3', '.mp4', '.css', 
                    '.js', '.json', '.xml', '.rss', '.woff', '.woff2', '.ttf', '.eot', '.otf']
            for ext in exts:
                if path.endswith(ext):
                    return False
            admin_paths = ['/admin', '/wp-admin', '/administrator', '/manage', '/login', 
                          '/logout', '/signup', '/register', '/api/', '/static/', '/media/', 
                          '/uploads/', '/cdn/', '/captcha', '/recaptcha']
            for admin in admin_paths:
                if admin in path:
                    return False
            return True
        except:
            return False
    
    def looks_blocked(response) -> bool:
        try:
            if response is None:
                return True
            if response.status_code in (403, 429, 503):
                return True
            body = (response.text or '').lower()
            blocked_words = ['captcha', 'access denied', 'forbidden', 'cloudflare', 'ddos-guard', 
                           'ваш доступ заблокирован', 'доступ запрещён']
            if any(x in body for x in blocked_words):
                return True
            if len(body.strip()) < 800:
                return True
        except:
            return True
        return False
    
    def extract_tech(html_text: str) -> list:
        tech = []
        text = html_text.lower()
        if 'react' in text and 'data-react' not in text:
            tech.append('React')
        if 'vue' in text or 'vue.js' in text:
            tech.append('Vue.js')
        if 'angular' in text:
            tech.append('Angular')
        if 'jquery' in text:
            tech.append('jQuery')
        if 'bootstrap' in text:
            tech.append('Bootstrap')
        if 'wordpress' in text:
            tech.append('WordPress')
        if 'bitrix' in text or '1c-bitrix' in text:
            tech.append('1C-Bitrix')
        if 'django' in text:
            tech.append('Django')
        if 'flask' in text:
            tech.append('Flask')
        if 'laravel' in text:
            tech.append('Laravel')
        if 'node.js' in text or 'nodejs' in text:
            tech.append('Node.js')
        if 'php' in text and 'php' not in tech:
            tech.append('PHP')
        if 'asp.net' in text or 'asp.net' in text.lower():
            tech.append('ASP.NET')
        if 'gatsby' in text:
            tech.append('Gatsby')
        if 'next.js' in text or 'nextjs' in text:
            tech.append('Next.js')
        if 'webpack' in text:
            tech.append('Webpack')
        if 'gulp' in text:
            tech.append('Gulp')
        if 'grunt' in text:
            tech.append('Grunt')
        if 'google tag manager' in text or 'gtm.js' in html_text:
            tech.append('Google Tag Manager')
        if 'yandex.metrika' in text or 'яндекс.метрика' in text:
            tech.append('Yandex.Metrika')
        if 'google analytics' in text:
            tech.append('Google Analytics')
        return list(set(tech))
    
    def analyze_page(page_url: str, depth: int, html: str, response) -> Dict:
        soup = BeautifulSoup(html, 'html.parser')
        
        title = soup.find('title')
        meta_desc = soup.find('meta', attrs={'name': 'description'})
        meta_keywords = soup.find('meta', attrs={'name': 'keywords'})
        meta_robots = soup.find('meta', attrs={'name': 'robots'})
        canonical = soup.find('link', rel='canonical')
        og_title = soup.find('meta', property='og:title')
        og_desc = soup.find('meta', property='og:description')
        twitter_card = soup.find('meta', attrs={'name': 'twitter:card'})
        
        h1_tags = soup.find_all('h1')
        h2_tags = soup.find_all('h2')
        h3_tags = soup.find_all('h3')
        
        images = soup.find_all('img')
        images_without_alt = [img for img in images if not img.get('alt') or not img.get('alt').strip()]
        images_empty_alt = [img for img in images if img.get('alt', '').strip() == '']
        
        links = soup.find_all('a')
        page_links = []
        dofollow = 0
        nofollow = 0
        
        for link in links:
            href = link.get('href', '')
            if href and not href.startswith('#') and not href.startswith('javascript'):
                rel = link.get('rel', [])
                if 'nofollow' in rel:
                    nofollow += 1
                else:
                    dofollow += 1
                page_links.append(href)
        
        scripts = soup.find_all('script')
        styles = soup.find_all('style')
        iframes = soup.find_all('iframe')
        
        body = soup.find('body')
        text_content = body.get_text(strip=True) if body else ''
        word_count = len(text_content.split())
        char_count = len(text_content)
        
        has_json_ld = len(soup.find_all('script', {'type': 'application/ld+json'}))
        has_microdata = len(soup.find_all(attrs={'itemscope': True}))
        has_rdfa = len(soup.find_all(attrs={'typeof': True}))
        has_hreflang = len(soup.find_all('link', {'rel': 'alternate', 'hreflang': True}))
        
        breadcrumbs = soup.find_all(class_=re.compile(r'breadcrumb', re.I))
        has_breadcrumbs = 1 if breadcrumbs else 0
        
        page_issues = []
        
        if not title:
            page_issues.append({'type': 'critical', 'issue': 'Отсутствует тег Title', 'url': page_url})
        else:
            if len(title.text) < 30:
                page_issues.append({'type': 'warning', 'issue': f'Title слишком короткий ({len(title.text)} символов)', 'url': page_url})
            elif len(title.text) > 70:
                page_issues.append({'type': 'warning', 'issue': f'Title слишком длинный ({len(title.text)} символов)', 'url': page_url})
        
        if not meta_desc:
            page_issues.append({'type': 'critical', 'issue': 'Отсутствует Meta Description', 'url': page_url})
        elif len(meta_desc.get('content', '')) < 120:
            page_issues.append({'type': 'warning', 'issue': 'Meta Description слишком короткий', 'url': page_url})
        
        if len(h1_tags) == 0:
            page_issues.append({'type': 'warning', 'issue': 'Отсутствует тег H1', 'url': page_url})
        elif len(h1_tags) > 1:
            page_issues.append({'type': 'warning', 'issue': f'Много тегов H1 ({len(h1_tags)})', 'url': page_url})
        
        if len(images_without_alt) > 0:
            page_issues.append({'type': 'warning', 'issue': f'{len(images_without_alt)} изображений без alt-текста', 'url': page_url})
        
        if not has_json_ld and not has_microdata:
            page_issues.append({'type': 'info', 'issue': 'Отсутствует Schema.org разметка', 'url': page_url})
        
        if not og_title or not og_desc:
            page_issues.append({'type': 'info', 'issue': 'Отсутствуют Open Graph теги', 'url': page_url})
        
        if not twitter_card:
            page_issues.append({'type': 'info', 'issue': 'Отсутствует Twitter Cards', 'url': page_url})
        
        if has_breadcrumbs == 0:
            page_issues.append({'type': 'info', 'issue': 'Отсутствуют хлебные крошки', 'url': page_url})
        
        return {
            'url': page_url,
            'depth': depth,
            'status': 'success' if response and response.status_code < 400 else 'error',
            'status_code': response.status_code if response else None,
            'title': title.text if title else None,
            'title_length': len(title.text) if title else 0,
            'meta_description': meta_desc.get('content') if meta_desc else None,
            'meta_keywords': meta_keywords.get('content') if meta_keywords else None,
            'meta_robots': meta_robots.get('content') if meta_robots else None,
            'canonical': canonical.get('href') if canonical else None,
            'og_tags': bool(og_title and og_desc),
            'twitter_card': bool(twitter_card),
            'h1_count': len(h1_tags),
            'h2_count': len(h2_tags),
            'h3_count': len(h3_tags),
            'images_total': len(images),
            'images_without_alt': len(images_without_alt),
            'images_empty_alt': len(images_empty_alt),
            'links_total': len(links),
            'links_found': page_links,
            'dofollow': dofollow,
            'nofollow': nofollow,
            'scripts_count': len(scripts),
            'styles_count': len(styles),
            'iframes_count': len(iframes),
            'word_count': word_count,
            'char_count': char_count,
            'has_json_ld': has_json_ld,
            'has_microdata': has_microdata,
            'has_rdfa': has_rdfa,
            'has_hreflang': has_hreflang,
            'has_breadcrumbs': has_breadcrumbs,
            'tech_stack': extract_tech(html),
            'issues': page_issues,
            'load_time': response.elapsed.total_seconds() if response and hasattr(response, 'elapsed') else None
        }
    
    def load_sitemap() -> bool:
        """Загрузка sitemap.xml из robots.txt или стандартных путей"""
        try:
            robots_url = urljoin(url, '/robots.txt')
            try:
                response = session.get(robots_url, timeout=10)
                if response.status_code == 200:
                    for line in response.text.split('\n'):
                        line = line.strip()
                        if line.lower().startswith('sitemap:'):
                            sitemap_url = line.split(':', 1)[1].strip()
                            try:
                                sm_response = session.get(sitemap_url, timeout=15)
                                if sm_response.status_code == 200:
                                    soup = BeautifulSoup(sm_response.content, 'xml')
                                    for url_tag in soup.find_all('url'):
                                        loc = url_tag.find('loc')
                                        if loc and loc.text:
                                            sitemap_urls.add(loc.text.strip())
                            except:
                                pass
            except:
                pass
        except:
            pass
        
        standard_paths = ['/sitemap.xml', '/sitemap_index.xml', '/sitemaps/sitemap.xml', 
                          '/wp-sitemap.xml', '/sitemap/sitemap.xml']
        for path in standard_paths:
            sitemap_url = urljoin(url, path)
            try:
                response = session.get(sitemap_url, timeout=15)
                if response.status_code == 200:
                    soup = BeautifulSoup(response.content, 'xml')
                    for url_tag in soup.find_all('url'):
                        loc = url_tag.find('loc')
                        if loc and loc.text:
                            sitemap_urls.add(loc.text.strip())
                    return True
            except:
                continue
        return False
    
    # Загружаем sitemap
    load_sitemap()
    
    start_time = time.time()
    pages_crawled = 0
    
    try:
        while to_visit and pages_crawled < max_pages:
            current_url, depth = to_visit.pop(0)
            
            if current_url in visited:
                continue
            if not is_valid_page(current_url):
                continue
            
            visited.add(current_url)
            all_urls.add(current_url)
            
            try:
                response = session.get(current_url, timeout=10, allow_redirects=True)
                
                if response.history:
                    for r in response.history:
                        redirects.append({
                            'from': normalize_url(r.url),
                            'to': normalize_url(response.url),
                            'status': r.status_code
                        })
                    current_url = response.url
                
                if response.status_code >= 400:
                    broken_links.append({
                        'url': normalize_url(current_url),
                        'status': response.status_code,
                        'error': response.reason if hasattr(response, 'reason') else 'Error'
                    })
                    pages_crawled += 1
                    continue
                
                if looks_blocked(response):
                    pages_crawled += 1
                    continue
                
                page_data = analyze_page(current_url, depth, response.text, response)
                all_pages_data.append(page_data)
                
                if depth < 2:
                    for link in page_data.get('links_found', [])[:30]:
                        if link:
                            full_link = urljoin(current_url, link)
                            norm_link = normalize_url(full_link)
                            if is_internal(full_link) and is_valid_page(full_link) and norm_link not in visited:
                                to_visit.append((full_link, depth + 1))
                                all_urls.add(full_link)
                                internal_links.append(norm_link)
                            elif not is_internal(full_link):
                                external_links.append({'url': norm_link, 'text': '', 'follow': 'dofollow'})
                
                pages_crawled += 1
                
            except requests.exceptions.Timeout:
                broken_links.append({'url': normalize_url(current_url), 'status': 'timeout', 'error': 'Connection timeout'})
                pages_crawled += 1
            except requests.exceptions.ConnectionError:
                broken_links.append({'url': normalize_url(current_url), 'status': 'connection_error', 'error': 'Connection error'})
                pages_crawled += 1
            except Exception as e:
                broken_links.append({'url': normalize_url(current_url), 'status': 'error', 'error': str(e)[:100]})
                pages_crawled += 1
        
        elapsed_time = time.time() - start_time
        
        # Агрегация результатов
        all_tech = []
        for page in all_pages_data:
            all_tech.extend(page.get('tech_stack', []))
        tech_counts = Counter(all_tech)
        top_technologies = [{'tech': t, 'count': c} for t, c in tech_counts.most_common(15)]
        
        all_issues = []
        for page in all_pages_data:
            all_issues.extend(page.get('issues', []))
        
        critical_issues = [i for i in all_issues if i.get('type') == 'critical']
        warning_issues = [i for i in all_issues if i.get('type') == 'warning']
        info_issues = [i for i in all_issues if i.get('type') == 'info']
        
        total_images = sum(p.get('images_total', 0) for p in all_pages_data)
        total_images_without_alt = sum(p.get('images_without_alt', 0) for p in all_pages_data)
        total_links = sum(p.get('links_total', 0) for p in all_pages_data)
        
        pages_with_title = len([p for p in all_pages_data if p.get('title')])
        pages_with_desc = len([p for p in all_pages_data if p.get('meta_description')])
        pages_with_h1 = len([p for p in all_pages_data if p.get('h1_count', 0) > 0])
        pages_with_schema = len([p for p in all_pages_data if p.get('has_json_ld') or p.get('has_microdata')])
        pages_with_og = len([p for p in all_pages_data if p.get('og_tags')])
        pages_with_breadcrumbs = len([p for p in all_pages_data if p.get('has_breadcrumbs')])
        
        avg_word_count = sum(p.get('word_count', 0) for p in all_pages_data) / max(1, len(all_pages_data))
        avg_load_time = sum(p.get('load_time', 0) for p in all_pages_data if p.get('load_time')) / max(1, len([p for p in all_pages_data if p.get('load_time')]))
        
        # SEO Score расчёт
        seo_score = 100
        seo_score -= min(30, len(critical_issues) * 10)
        seo_score -= min(20, len(warning_issues) * 3)
        seo_score -= min(10, total_images_without_alt * 0.5)
        seo_score -= min(10, len(broken_links) * 2)
        seo_score -= min(5, len(redirects) * 1)
        seo_score -= min(5, len(info_issues) * 0.5)
        seo_score = max(0, min(100, seo_score))
        
        # Категоризация по страницам
        good_pages = len([p for p in all_pages_data if len([i for i in p.get('issues', []) if i.get('type') in ['critical', 'warning']]) == 0])
        average_pages = len([p for p in all_pages_data if len([i for i in p.get('issues', []) if i.get('type') == 'critical']) == 0 and len([i for i in p.get('issues', []) if i.get('type') == 'warning']) <= 2])
        bad_pages = len(all_pages_data) - good_pages - average_pages
        
        return {
            "task_type": "site_analyze",
            "url": url,
            "completed_at": datetime.utcnow().isoformat(),
            "results": {
                "summary": {
                    "pages_crawled": pages_crawled,
                    "pages_total_found": len(all_urls),
                    "crawl_time_seconds": round(elapsed_time, 2),
                    "sitemap_urls_found": len(sitemap_urls),
                    "internal_links_count": len(set(internal_links)),
                    "external_links_count": len(set([e.get('url') for e in external_links])),
                    "broken_links_count": len(broken_links),
                    "redirects_count": len(redirects),
                    "seo_score": seo_score,
                    "critical_issues": len(critical_issues),
                    "warning_issues": len(warning_issues),
                    "info_issues": len(info_issues),
                    "good_pages": good_pages,
                    "average_pages": average_pages,
                    "bad_pages": max(0, bad_pages)
                },
                "content_analysis": {
                    "total_images": total_images,
                    "images_without_alt": total_images_without_alt,
                    "total_links": total_links,
                    "average_word_count": round(avg_word_count, 0),
                    "average_load_time": round(avg_load_time, 3),
                    "pages_with_title": pages_with_title,
                    "pages_with_meta_desc": pages_with_desc,
                    "pages_with_h1": pages_with_h1,
                    "pages_with_schema_org": pages_with_schema,
                    "pages_with_og_tags": pages_with_og,
                    "pages_with_breadcrumbs": pages_with_breadcrumbs
                },
                "technology_stack": top_technologies,
                "pages_detail": all_pages_data[:30],
                "broken_links": broken_links[:30],
                "redirects": redirects[:20],
                "sitemap_urls": list(sitemap_urls)[:50],
                "all_issues": all_issues[:100],
                "critical_issues": critical_issues[:20],
                "warning_issues": warning_issues[:50],
                "info_issues": info_issues[:30],
                "external_links_sample": external_links[:30],
                "recommendations": generate_recommendations_full(url, critical_issues, warning_issues, info_issues, tech_counts, seo_score, pages_crawled)
            }
        }
    except Exception as e:
        return {
            "task_type": "site_analyze",
            "url": url,
            "completed_at": datetime.utcnow().isoformat(),
            "results": {
                "error": str(e)
            }
        }
    
    def analyze_page(page_url: str, depth: int, html: str) -> Dict:
        soup = BeautifulSoup(html, 'html.parser')
        
        title = soup.find('title')
        meta_desc = soup.find('meta', attrs={'name': 'description'})
        meta_keywords = soup.find('meta', attrs={'name': 'keywords'})
        meta_robots = soup.find('meta', attrs={'name': 'robots'})
        canonical = soup.find('link', rel='canonical')
        h1_tags = soup.find_all('h1')
        h2_tags = soup.find_all('h2')
        
        images = soup.find_all('img')
        images_without_alt = [img for img in images if not img.get('alt') or not img.get('alt').strip()]
        images_with_empty_alt = [img for img in images if img.get('alt', '').strip() == '']
        
        links = soup.find_all('a')
        page_links = []
        for link in links:
            href = link.get('href', '')
            if href and not href.startswith('#') and not href.startswith('javascript'):
                page_links.append(href)
        
        scripts = soup.find_all('script')
        styles = soup.find_all('style')
        iframes = soup.find_all('iframe')
        
        body = soup.find('body')
        text_content = body.get_text(strip=True) if body else ''
        word_count = len(text_content.split())
        
        has_og_tags = bool(soup.find('meta', property='og:title'))
        has_twitter_cards = bool(soup.find('meta', property='twitter:card'))
        has_schema = bool(soup.find('script', type='application/ld+json'))
        
        page_issues = []
        if not title:
            page_issues.append({'type': 'critical', 'issue': 'Отсутствует тег title', 'url': page_url})
        if title and len(title.text) < 30:
            page_issues.append({'type': 'warning', 'issue': 'Слишком короткий title (< 30 символов)', 'url': page_url})
        if title and len(title.text) > 70:
            page_issues.append({'type': 'warning', 'issue': 'Слишком длинный title (> 70 символов)', 'url': page_url})
        if not meta_desc:
            page_issues.append({'type': 'critical', 'issue': 'Отсутствует meta description', 'url': page_url})
        if meta_desc and len(meta_desc.get('content', '')) < 120:
            page_issues.append({'type': 'warning', 'issue': 'Слишком короткий meta description', 'url': page_url})
        if len(h1_tags) == 0:
            page_issues.append({'type': 'warning', 'issue': 'Не найден H1', 'url': page_url})
        elif len(h1_tags) > 1:
            page_issues.append({'type': 'warning', 'issue': f'Несколько H1 ({len(h1_tags)})', 'url': page_url})
        if len(images_without_alt) > 0:
            page_issues.append({'type': 'warning', 'issue': f'Изображений без alt: {len(images_without_alt)}', 'url': page_url})
        
        return {
            'url': page_url,
            'depth': depth,
            'status': 'success',
            'title': title.text if title else None,
            'title_length': len(title.text) if title else 0,
            'meta_description': meta_desc.get('content') if meta_desc else None,
            'meta_keywords': meta_keywords.get('content') if meta_keywords else None,
            'meta_robots': meta_robots.get('content') if meta_robots else None,
            'canonical': canonical.get('href') if canonical else None,
            'h1_count': len(h1_tags),
            'h2_count': len(h2_tags),
            'images_total': len(images),
            'images_without_alt': len(images_without_alt),
            'images_empty_alt': len(images_with_empty_alt),
            'links_total': len(links),
            'links_found': page_links,
            'scripts_count': len(scripts),
            'styles_count': len(styles),
            'iframes_count': len(iframes),
            'word_count': word_count,
            'has_og_tags': has_og_tags,
            'has_twitter_cards': has_twitter_cards,
            'has_schema_org': has_schema,
            'issues': page_issues,
            'tech_stack': extract_tech(html)
        }
    
    start_time = time.time()
    pages_crawled = 0
    
    try:
        while to_visit and pages_crawled < max_pages:
            current_url, depth = to_visit.pop(0)
            
            if current_url in visited:
                continue
            if not is_valid_page(current_url):
                continue
            
            visited.add(current_url)
            all_urls.add(current_url)
            
            try:
                response = session.get(current_url, timeout=10, allow_redirects=True)
                
                if response.history:
                    for r in response.history:
                        redirects.append({
                            'from': r.url,
                            'to': response.url,
                            'status': r.status_code
                        })
                    current_url = response.url
                
                if response.status_code >= 400:
                    broken_links.append({
                        'url': current_url,
                        'status': response.status_code
                    })
                    pages_crawled += 1
                    continue
                
                page_data = analyze_page(current_url, depth, response.text)
                all_pages_data.append(page_data)
                internal_links.extend([l for l in page_data['links_found'] if is_internal(l)])
                
                if depth < 2:
                    for link in page_data['links_found'][:50]:
                        if link not in visited and len(all_urls) < max_pages * 2:
                            full_link = urljoin(current_url, link)
                            if is_internal(full_link) and is_valid_page(full_link):
                                to_visit.append((full_link, depth + 1))
                                all_urls.add(full_link)
                
                external_links.extend([l for l in page_data['links_found'] if not is_internal(l)])
                
                pages_crawled += 1
                
            except requests.exceptions.Timeout:
                broken_links.append({'url': current_url, 'status': 'timeout'})
                pages_crawled += 1
            except requests.exceptions.ConnectionError:
                broken_links.append({'url': current_url, 'status': 'connection_error'})
                pages_crawled += 1
            except Exception as e:
                broken_links.append({'url': current_url, 'status': str(e)})
                pages_crawled += 1
        
        elapsed_time = time.time() - start_time
        
        all_tech = []
        for page in all_pages_data:
            all_tech.extend(page.get('tech_stack', []))
        tech_counts = Counter(all_tech)
        top_technologies = [{'tech': t, 'count': c} for t, c in tech_counts.most_common(10)]
        
        all_issues = []
        for page in all_pages_data:
            all_issues.extend(page.get('issues', []))
        
        critical_issues = [i for i in all_issues if i.get('type') == 'critical']
        warning_issues = [i for i in all_issues if i.get('type') == 'warning']
        
        total_images = sum(p.get('images_total', 0) for p in all_pages_data)
        total_images_without_alt = sum(p.get('images_without_alt', 0) for p in all_pages_data)
        total_links = sum(p.get('links_total', 0) for p in all_pages_data)
        
        pages_with_title = len([p for p in all_pages_data if p.get('title')])
        pages_with_desc = len([p for p in all_pages_data if p.get('meta_description')])
        pages_with_h1 = len([p for p in all_pages_data if p.get('h1_count', 0) > 0])
        pages_with_schema = len([p for p in all_pages_data if p.get('has_schema_org')])
        
        avg_word_count = sum(p.get('word_count', 0) for p in all_pages_data) / max(1, len(all_pages_data))
        
        redirect_count = len(redirects)
        broken_count = len(broken_links)
        
        seo_score = 100
        seo_score -= min(30, len(critical_issues) * 10)
        seo_score -= min(20, len(warning_issues) * 3)
        seo_score -= min(20, total_images_without_alt * 0.5)
        seo_score -= min(10, redirect_count * 2)
        seo_score -= min(10, broken_count * 2)
        seo_score = max(0, seo_score)
        
        return {
            "task_type": "site_analyze",
            "url": url,
            "completed_at": datetime.utcnow().isoformat(),
            "results": {
                "summary": {
                    "pages_crawled": pages_crawled,
                    "pages_total_found": len(all_urls),
                    "crawl_time_seconds": round(elapsed_time, 2),
                    "internal_links_count": len(set(internal_links)),
                    "external_links_count": len(set(external_links)),
                    "broken_links_count": broken_count,
                    "redirects_count": redirect_count,
                    "seo_score": seo_score,
                    "critical_issues": len(critical_issues),
                    "warning_issues": len(warning_issues)
                },
                "content_analysis": {
                    "total_images": total_images,
                    "images_without_alt": total_images_without_alt,
                    "total_links": total_links,
                    "average_word_count": round(avg_word_count, 0),
                    "pages_with_title": pages_with_title,
                    "pages_with_meta_desc": pages_with_desc,
                    "pages_with_h1": pages_with_h1,
                    "pages_with_schema_org": pages_with_schema
                },
                "technology_stack": top_technologies,
                "pages_detail": all_pages_data[:50],
                "broken_links": broken_links[:50],
                "redirects": redirects[:20],
                "all_issues": all_issues[:100],
                "critical_issues": critical_issues[:20],
                "warning_issues": warning_issues[:50],
                "recommendations": generate_recommendations(critical_issues, warning_issues, tech_counts, seo_score)
            }
        }
    except Exception as e:
        return {
            "task_type": "site_analyze",
            "url": url,
            "completed_at": datetime.utcnow().isoformat(),
            "results": {
                "error": str(e)
            }
        }


def generate_recommendations_full(site_url: str, critical_issues: list, warning_issues: list, info_issues: list, tech_counts: dict, score: int, pages_crawled: int) -> list:
    """Генерация рекомендаций на основе анализа - приближено к оригинальному скрипту"""
    recommendations = []
    
    if score < 30:
        recommendations.append({"priority": "critical", "text": "КРИТИЧЕСКИ НИЗКИЙ SEO SCORE! Требуется немедленный аудит и исправления всех критических проблем."})
    elif score < 50:
        recommendations.append({"priority": "high", "text": "Низкая SEO-оценка. Необходимо исправить критические проблемы в первую очередь."})
    elif score < 70:
        recommendations.append({"priority": "medium", "text": "Средняя SEO-оценка. Есть возможности для улучшения."})
    
    # Критические проблемы
    if any('title' in i.get('issue', '').lower() for i in critical_issues):
        count = len([i for i in critical_issues if 'title' in i.get('issue', '').lower()])
        recommendations.append({"priority": "high", "text": f"Добавьте Title теги на {count} страниц (оптимальная длина 60-70 символов)"})
    
    if any('description' in i.get('issue', '').lower() for i in critical_issues):
        count = len([i for i in critical_issues if 'description' in i.get('issue', '').lower()])
        recommendations.append({"priority": "high", "text": f"Добавьте Meta Description на {count} страниц (оптимальная длина 120-160 символов)"})
    
    if any('h1' in i.get('issue', '').lower() for i in warning_issues):
        count = len([i for i in warning_issues if 'h1' in i.get('issue', '').lower()])
        recommendations.append({"priority": "medium", "text": f"Исправьте теги H1 на {count} страниц (должен быть 1 тег H1 на страницу)"})
    
    # Изображения
    img_issues = [i for i in warning_issues if 'image' in i.get('issue', '').lower() or 'alt' in i.get('issue', '').lower()]
    if img_issues:
        total_img = sum(1 for i in img_issues)
        recommendations.append({"priority": "medium", "text": f"Добавьте alt-текст к {total_img} изображениям для улучшения SEO и доступности"})
    
    # Технологии
    if 'WordPress' in tech_counts:
        recommendations.append({"priority": "medium", "text": "Для WordPress установите SEO-плагин: Yoast SEO или Rank Math"})
    
    if 'jQuery' in tech_counts and 'React' not in tech_counts and 'Vue.js' not in tech_counts:
        recommendations.append({"priority": "low", "text": "Сайт использует jQuery. Рассмотрите переход на современный JS-фреймворк"})
    
    # Schema.org
    schema_issues = [i for i in info_issues if 'schema' in i.get('issue', '').lower()]
    if schema_issues:
        recommendations.append({"priority": "medium", "text": "Добавьте Schema.org разметку (Organization, Breadcrumb, Product, Article) для улучшения сниппетов в поиске"})
    
    # Open Graph
    og_issues = [i for i in info_issues if 'open graph' in i.get('issue', '').lower()]
    if og_issues:
        recommendations.append({"priority": "medium", "text": "Добавьте Open Graph теги (og:title, og:description, og:image) для улучшения шаринга в соцсетях"})
    
    # Twitter Cards
    twitter_issues = [i for i in info_issues if 'twitter' in i.get('issue', '').lower()]
    if twitter_issues:
        recommendations.append({"priority": "low", "text": "Добавьте Twitter Cards мета-теги для красивого отображения ссылок в Twitter"})
    
    # Хлебные крошки
    breadcrumb_issues = [i for i in info_issues if 'хлеб' in i.get('issue', '').lower() or 'breadcrumb' in i.get('issue', '').lower()]
    if breadcrumb_issues:
        recommendations.append({"priority": "medium", "text": "Добавьте хлебные крошки на страницы для улучшения навигации и SEO"})
    
    # Аналитика
    has_analytics = any('Analytics' in t or 'Metrika' in t for t in tech_counts.keys())
    if not has_analytics:
        recommendations.append({"priority": "medium", "text": "Подключите системы аналитики: Google Analytics и/или Yandex.Metrika для отслеживания посещаемости"})
    
    # Скорость
    recommendations.append({"priority": "high", "text": "Оптимизируйте скорость загрузки страниц (сжатие изображений, минификация CSS/JS, кэширование)"})
    
    # Битые ссылки
    broken_count = len([i for i in critical_issues if 'broken' in i.get('issue', '').lower() or any(x in i.get('issue', '').lower() for x in ['404', 'ошибка', 'error'])])
    if broken_count > 0:
        recommendations.append({"priority": "high", "text": f"Найдено {broken_count} битых ссылок. Проверьте и исправьте или удалите их"})
    
    # Канонические URL
    canon_issues = [i for i in warning_issues if 'canonical' in i.get('issue', '').lower()]
    if canon_issues:
        recommendations.append({"priority": "medium", "text": "Настройте канонические URL для всех страниц во избежание дублирования контента"})
    
    # Мобильная адаптация
    mobile_issues = [i for i in critical_issues if 'mobile' in i.get('issue', '').lower() or 'viewport' in i.get('issue', '').lower()]
    if mobile_issues:
        recommendations.append({"priority": "high", "text": "Адаптируйте сайт для мобильных устройств (Viewport, Responsive Design)"})
    
    # HTTPS
    if not site_url.startswith('https://'):
        recommendations.append({"priority": "high", "text": "Перенесите сайт на HTTPS для безопасности и SEO"})
    
    return recommendations


def check_render_simple(url: str) -> Dict[str, Any]:
    """Simple render audit - basic HTML vs JS comparison"""
    import requests
    
    try:
        # Get page without JS
        response_no_js = requests.get(url, timeout=10, headers={"User-Agent": "Mozilla/5.0"})
        
        # Basic comparison
        content_length = len(response_no_js.text)
        word_count = len(response_no_js.text.split())
        
        # Check for common JavaScript frameworks
        js_frameworks = []
        text_lower = response_no_js.text.lower()
        if 'react' in text_lower:
            js_frameworks.append("React")
        if 'vue' in text_lower:
            js_frameworks.append("Vue.js")
        if 'angular' in text_lower:
            js_frameworks.append("Angular")
        if 'jquery' in text_lower:
            js_frameworks.append("jQuery")
        
        # Check for SSR indicators
        is_ssr = bool(response_no_js.text.strip())
        
        return {
            "task_type": "render_audit",
            "url": url,
            "completed_at": datetime.utcnow().isoformat(),
            "results": {
                "status_code": response_no_js.status_code,
                "content_length": content_length,
                "word_count": word_count,
                "is_ssr": is_ssr,
                "js_frameworks": js_frameworks,
                "recommendations": [
                    "Ensure critical content is available without JavaScript",
                    "Use server-side rendering for better SEO",
                    "Test with tools like Google Search Console"
                ]
            }
        }
    except Exception as e:
        return {
            "task_type": "render_audit",
            "url": url,
            "completed_at": datetime.utcnow().isoformat(),
            "results": {
                "error": str(e)
            }
        }


def check_render_full(
    url: str,
    task_id: str,
    progress_callback=None,
) -> Dict[str, Any]:
    """Feature-flagged render audit with v2 service fallback."""
    from app.config import settings
    debug_render = bool(getattr(settings, "RENDER_AUDIT_DEBUG", False) or getattr(settings, "DEBUG", False))

    def _ensure_render_profiles(payload: Dict[str, Any]) -> Dict[str, Any]:
        if not isinstance(payload, dict):
            return payload
        results = payload.get("results")
        if not isinstance(results, dict):
            return payload

        variants = results.get("variants")
        if not isinstance(variants, list):
            variants = []
            results["variants"] = variants
        if debug_render:
            print(f"[RENDER-DEBUG][{task_id}] pre-normalize variants={len(variants)}")

        # Normalize profile identity/labels so mobile cannot be shown as desktop.
        normalized: List[Dict[str, Any]] = []
        for v in variants:
            if not isinstance(v, dict):
                continue
            variant_id = str(v.get("variant_id", "")).lower().strip()
            if not variant_id:
                profile_type = str(v.get("profile_type", "")).lower().strip()
                if profile_type == "mobile" or bool(v.get("mobile")):
                    variant_id = "googlebot_mobile"
                else:
                    variant_id = "googlebot_desktop"
                v["variant_id"] = variant_id

            if variant_id == "googlebot_mobile":
                v["variant_label"] = "Googlebot (мобильный)"
                v["mobile"] = True
                v["profile_type"] = "mobile"
            elif variant_id == "googlebot_desktop":
                v["variant_label"] = "Googlebot (ПК)"
                v["mobile"] = False
                v["profile_type"] = "desktop"
            normalized.append(v)
        variants = normalized
        results["variants"] = variants
        if debug_render:
            print(
                f"[RENDER-DEBUG][{task_id}] post-normalize variants="
                + ", ".join([f"{v.get('variant_id')}:{v.get('variant_label')}:{v.get('profile_type')}" for v in variants])
            )

        required = [
            ("googlebot_desktop", "Googlebot (ПК)", False),
            ("googlebot_mobile", "Googlebot (мобильный)", True),
        ]
        existing_ids = {str(v.get("variant_id", "")).lower() for v in variants if isinstance(v, dict)}

        for variant_id, label, is_mobile in required:
            if variant_id in existing_ids:
                continue
            fail_issue = {
                "severity": "critical",
                "code": "variant_missing_in_result",
                "title": "Профиль рендеринга отсутствует в результате",
                "details": f"Профиль {label} не вернулся из движка и добавлен как ошибка.",
            }
            variants.append(
                {
                    "variant_id": variant_id,
                    "variant_label": label,
                    "mobile": is_mobile,
                    "profile_type": "mobile" if is_mobile else "desktop",
                    "raw": {},
                    "rendered": {},
                    "missing": {"visible_text": [], "headings": [], "links": [], "images": [], "structured_data": []},
                    "meta_non_seo": {"raw": {}, "rendered": {}, "comparison": {"total": 0, "same": 0, "changed": 0, "only_rendered": 0, "only_raw": 0, "items": []}},
                    "seo_required": {"total": 0, "pass": 0, "warn": 0, "fail": 0, "items": []},
                    "metrics": {"total_missing": 0.0, "rendered_total": 0.0, "missing_pct": 0.0, "score": 0.0},
                    "timings": {"raw_s": 0.0, "rendered_s": 0.0},
                    "timing_nojs_ms": {},
                    "timing_js_ms": {},
                    "issues": [fail_issue],
                    "recommendations": ["Проверьте логи рендер-движка и окружение Playwright для этого профиля."],
                    "screenshots": {},
                }
            )
            issues = results.get("issues")
            if not isinstance(issues, list):
                issues = []
                results["issues"] = issues
            issues.append({**fail_issue, "variant": label})

        summary = results.get("summary")
        if not isinstance(summary, dict):
            summary = {}
            results["summary"] = summary
        summary["variants_total"] = len(results.get("variants", []))
        results["issues_count"] = len(results.get("issues", []) or [])
        if debug_render:
            print(
                f"[RENDER-DEBUG][{task_id}] ensured variants_total={summary['variants_total']} "
                f"issues_count={results['issues_count']}"
            )
        return payload

    engine = (getattr(settings, "RENDER_AUDIT_ENGINE", "v2") or "v2").lower()
    if engine == "v2":
        try:
            from app.tools.render.service_v2 import RenderAuditServiceV2

            checker = RenderAuditServiceV2(timeout=getattr(settings, "RENDER_AUDIT_TIMEOUT", 35))
            result = checker.run(url=url, task_id=task_id, progress_callback=progress_callback)
            ensured = _ensure_render_profiles(result)
            if debug_render:
                ensured_results = ensured.get("results", {}) if isinstance(ensured, dict) else {}
                ensured_variants = ensured_results.get("variants", []) if isinstance(ensured_results, dict) else []
                print(
                    f"[RENDER-DEBUG][{task_id}] return variants="
                    + ", ".join([f"{v.get('variant_id')}:{v.get('variant_label')}:{v.get('profile_type')}" for v in ensured_variants if isinstance(v, dict)])
                )
            return ensured
        except Exception as e:
            print(f"[API] render v2 failed, fallback to simple: {e}")
            fallback = check_render_simple(url)
            fallback_results = fallback.get("results", {}) if isinstance(fallback, dict) else {}
            fallback_results["engine"] = "legacy-fallback"
            fallback_results["engine_error"] = str(e)
            fallback_results["issues"] = [
                {
                    "severity": "critical",
                    "code": "render_engine_error",
                    "title": "Ошибка движка render v2",
                    "details": str(e),
                }
            ]
            fallback_results["issues_count"] = 1
            fallback_results["summary"] = {
                "variants_total": 0,
                "critical_issues": 1,
                "warning_issues": 0,
                "info_issues": 0,
                "score": None,
                "missing_total": 0,
                "avg_missing_pct": 0,
                "avg_raw_load_ms": 0,
                "avg_js_load_ms": 0,
            }
            fallback_results["variants"] = []
            fallback_results["recommendations"] = ["Движок v2 недоступен, проверьте окружение Playwright."]
            return fallback

    legacy = check_render_simple(url)
    legacy_results = legacy.get("results", {}) if isinstance(legacy, dict) else {}
    legacy_results["engine"] = "legacy"
    return legacy


def check_mobile_simple(url: str) -> Dict[str, Any]:
    """Simple mobile check - viewport and responsive indicators"""
    import requests
    from bs4 import BeautifulSoup
    
    try:
        response = requests.get(url, timeout=10, headers={"User-Agent": "Mozilla/5.0 (iPhone; CPU iPhone OS 14_0 like Mac OS X)"})
        soup = BeautifulSoup(response.text, 'html.parser')
        
        issues = []
        
        # Check viewport meta tag
        viewport = soup.find('meta', attrs={'name': 'viewport'})
        if not viewport:
            issues.append("Отсутствует мета-тег viewport")
        else:
            content = viewport.get('content', '')
            if 'width=device-width' not in content:
                issues.append("В viewport не задано width=device-width")
        
        # Check for responsive images
        images = soup.find_all('img')
        large_images = []
        for img in images:
            src = img.get('src', '')
            if any(size in src for size in ['large', 'big', 'full', 'original']):
                large_images.append(src)
        
        if large_images:
            issues.append(f"Found {len(large_images)} potentially large images")
        
        # Check tap targets
        buttons = soup.find_all(['button', 'a'])
        small_buttons = []
        for btn in buttons:
            text = btn.get_text(strip=True)
            if text and len(text) > 20:
                small_buttons.append(text[:20])
        
        # Mobile-friendly indicators
        mobile_friendly = len(issues) == 0
        
        return {
            "task_type": "mobile_check",
            "url": url,
            "completed_at": datetime.utcnow().isoformat(),
            "results": {
                "status_code": response.status_code,
                "viewport_found": bool(viewport),
                "viewport_content": viewport.get('content') if viewport else None,
                "total_images": len(images),
                "issues": issues,
                "issues_count": len(issues),
                "mobile_friendly": mobile_friendly,
                "score": max(0, 100 - len(issues) * 20),
                "recommendations": [
                    "Use responsive design with flexible layouts",
                    "Ensure tap targets are at least 48x48 pixels",
                    "Use appropriate font sizes (16px minimum)",
                    "Test with Google Mobile-Friendly Test"
                ]
            }
        }
    except Exception as e:
        return {
            "task_type": "mobile_check",
            "url": url,
            "completed_at": datetime.utcnow().isoformat(),
            "results": {
                "error": str(e)
            }
        }


def check_mobile_full(
    url: str,
    task_id: str,
    mode: str = "full",
    devices: Optional[List[str]] = None,
    progress_callback=None,
) -> Dict[str, Any]:
    """Feature-flagged mobile check with v2 service fallback."""
    from app.config import settings

    engine = (getattr(settings, "MOBILE_CHECK_ENGINE", "v2") or "v2").lower()
    if engine == "v2":
        try:
            from app.tools.mobile.service_v2 import MobileCheckServiceV2

            checker = MobileCheckServiceV2(timeout=getattr(settings, "MOBILE_CHECK_TIMEOUT", 20))
            selected_mode = (mode or getattr(settings, "MOBILE_CHECK_MODE", "quick") or "quick").lower()
            if selected_mode not in ("quick", "full"):
                selected_mode = "full"
            return checker.run(
                url=url,
                task_id=task_id,
                mode=selected_mode,
                selected_devices=devices,
                progress_callback=progress_callback,
            )
        except Exception as e:
            print(f"[API] mobile v2 failed, fallback to simple: {e}")
            fallback = check_mobile_simple(url)
            fallback_results = fallback.get("results", {})
            fallback_results["engine"] = "legacy-fallback"
            fallback_results["engine_error"] = str(e)
            fallback_results["mobile_friendly"] = False
            fallback_results["score"] = None
            fallback_results["issues"] = [
                {
                    "severity": "critical",
                    "code": "mobile_engine_error",
                    "title": "Ошибка движка mobile v2",
                    "details": str(e),
                }
            ]
            fallback_results["issues_count"] = 1
            return fallback

    legacy = check_mobile_simple(url)
    legacy_results = legacy.get("results", {})
    legacy_results["engine"] = "legacy"
    return legacy


def check_site_audit_pro(
    url: str,
    task_id: str,
    mode: str = "quick",
    max_pages: int = 5,
    batch_mode: bool = False,
    batch_urls: Optional[List[str]] = None,
    extended_hreflang_checks: bool = False,
    progress_callback=None,
) -> Dict[str, Any]:
    """Feature-flagged Site Audit Pro entrypoint."""
    from app.tools.site_pro.service import SiteAuditProService

    service = SiteAuditProService()
    return service.run(
        url=url,
        task_id=task_id,
        mode=mode,
        max_pages=max_pages,
        batch_mode=batch_mode,
        batch_urls=batch_urls or [],
        extended_hreflang_checks=extended_hreflang_checks,
        progress_callback=progress_callback,
    )


class SiteAnalyzeRequest(BaseModel):
    url: str
    max_pages: int = 20

class RenderAuditRequest(BaseModel):
    url: str

class MobileCheckRequest(BaseModel):
    url: str
    mode: Optional[str] = "quick"
    devices: Optional[List[str]] = None


class SiteAuditProRequest(BaseModel):
    url: Optional[str] = None
    mode: Optional[str] = "quick"
    max_pages: int = 5
    batch_mode: bool = False
    batch_urls: Optional[List[str]] = None
    extended_hreflang_checks: bool = False

    @field_validator("batch_urls", mode="before")
    @classmethod
    def _normalize_batch_urls(cls, value):
        if value is None:
            return []
        if isinstance(value, str):
            return [value]
        if isinstance(value, list):
            return [str(v).strip() for v in value if str(v).strip()]
        return []


@router.post("/tasks/site-analyze")
async def create_site_analyze(data: SiteAnalyzeRequest):
    """Full site analysis with multi-page crawling"""
    url = data.url
    max_pages = min(data.max_pages, 50)
    
    print(f"[API] Full site analysis for: {url} (max_pages={max_pages})")
    
    result = check_site_full(url, max_pages)
    task_id = f"site-{datetime.now().timestamp()}"
    create_task_result(task_id, "site_analyze", url, result)
    
    return {
        "task_id": task_id,
        "status": "SUCCESS",
        "message": "Site analysis completed"
    }


@router.post("/tasks/render-audit")
async def create_render_audit(data: RenderAuditRequest, background_tasks: BackgroundTasks):
    """Render audit with background progress updates."""
    url = data.url
    from app.config import settings
    debug_render = bool(getattr(settings, "RENDER_AUDIT_DEBUG", False) or getattr(settings, "DEBUG", False))

    print(f"[API] Render audit queued for: {url}")
    task_id = f"render-{datetime.now().timestamp()}"
    create_task_pending(task_id, "render_audit", url, status_message="Задача поставлена в очередь")

    def _run_render_task() -> None:
        try:
            update_task_state(task_id, status="RUNNING", progress=5, status_message="Подготовка рендер-аудита")

            def _progress(progress: int, message: str) -> None:
                update_task_state(task_id, status="RUNNING", progress=progress, status_message=message)

            result = check_render_full(url, task_id=task_id, progress_callback=_progress)
            results = result.get("results", {}) if isinstance(result, dict) else {}
            if debug_render:
                variants = results.get("variants", []) if isinstance(results, dict) else []
                print(
                    f"[RENDER-DEBUG][{task_id}] background-result variants="
                    + ", ".join(
                        [
                            f"{v.get('variant_id')}:{v.get('variant_label')}:{v.get('profile_type')}:{v.get('mobile')}"
                            for v in variants
                            if isinstance(v, dict)
                        ]
                    )
                )
            engine = (results.get("engine") or "").lower()
            has_engine_error = bool(results.get("engine_error")) or engine in ("legacy-fallback", "")

            if has_engine_error:
                error_message = results.get("engine_error") or "Ошибка движка render"
                update_task_state(
                    task_id,
                    status="FAILURE",
                    progress=100,
                    status_message="Ошибка рендер-аудита",
                    result=result,
                    error=error_message,
                )
                return

            update_task_state(
                task_id,
                status="SUCCESS",
                progress=100,
                status_message="Рендер-аудит завершен",
                result=result,
            )
        except Exception as exc:
            update_task_state(
                task_id,
                status="FAILURE",
                progress=100,
                status_message="Ошибка рендер-аудита",
                error=str(exc),
            )

    background_tasks.add_task(_run_render_task)
    return {
        "task_id": task_id,
        "status": "PENDING",
        "message": "Render audit queued",
    }


@router.post("/tasks/mobile-check")
async def create_mobile_check(data: MobileCheckRequest, background_tasks: BackgroundTasks):
    """Mobile check with background progress updates."""
    url = data.url
    from app.config import settings
    mode = data.mode or getattr(settings, "MOBILE_CHECK_MODE", "quick") or "quick"
    devices = data.devices

    print(f"[API] Mobile check queued for: {url}")
    task_id = f"mobile-{datetime.now().timestamp()}"
    create_task_pending(task_id, "mobile_check", url, status_message="Задача поставлена в очередь")

    def _run_mobile_task() -> None:
        try:
            update_task_state(task_id, status="RUNNING", progress=5, status_message="Подготовка мобильного аудита")

            def _progress(progress: int, message: str) -> None:
                update_task_state(task_id, status="RUNNING", progress=progress, status_message=message)

            result = check_mobile_full(
                url,
                task_id=task_id,
                mode=mode,
                devices=devices,
                progress_callback=_progress,
            )

            results = result.get("results", {}) if isinstance(result, dict) else {}
            engine = (results.get("engine") or "").lower()
            has_engine_error = bool(results.get("engine_error")) or engine in ("legacy-fallback", "")

            if has_engine_error:
                error_message = results.get("engine_error") or "Ошибка движка mobile"
                update_task_state(
                    task_id,
                    status="FAILURE",
                    progress=100,
                    status_message="Мобильный аудит завершился с ошибкой",
                    result=result,
                    error=error_message,
                )
                return

            update_task_state(
                task_id,
                status="SUCCESS",
                progress=100,
                status_message="Мобильный аудит завершен",
                result=result,
                error=None,
            )
        except Exception as e:
            update_task_state(
                task_id,
                status="FAILURE",
                progress=100,
                status_message="Мобильный аудит завершился с ошибкой",
                error=str(e),
            )

    background_tasks.add_task(_run_mobile_task)
    return {"task_id": task_id, "status": "PENDING", "message": "Проверка мобильной версии запущена"}




@router.post("/tasks/site-audit-pro")
async def create_site_audit_pro(data: SiteAuditProRequest, background_tasks: BackgroundTasks):
    """Site Audit Pro queued as isolated background task."""
    from app.config import settings
    if not getattr(settings, "SITE_AUDIT_PRO_ENABLED", True):
        return {"error": "Site Audit Pro is disabled by feature flag"}

    raw_url = (data.url or "").strip()
    default_mode = (getattr(settings, "SITE_AUDIT_PRO_DEFAULT_MODE", "quick") or "quick").lower()
    mode = (data.mode or default_mode).lower()
    if mode not in ("quick", "full"):
        mode = "quick"
    batch_mode = bool(getattr(data, "batch_mode", False))
    extended_hreflang_checks = bool(getattr(data, "extended_hreflang_checks", False))
    if batch_mode:
        mode = "full"
    max_pages_limit = int(getattr(settings, "SITE_AUDIT_PRO_MAX_PAGES_LIMIT", 5) or 5)
    effective_max_pages_limit = 500 if batch_mode else max_pages_limit
    max_pages = max(1, min(int(data.max_pages or 5), effective_max_pages_limit))

    raw_batch_urls = list(getattr(data, "batch_urls", []) or [])
    normalized_batch_urls: List[str] = []
    seen_batch_urls: set[str] = set()
    for item in raw_batch_urls:
        candidate = str(item or "").strip()
        if not candidate:
            continue
        if not candidate.startswith(("http://", "https://")):
            candidate = f"https://{candidate}"
        parsed = urlparse(candidate)
        if not parsed.scheme or not parsed.netloc:
            continue
        if candidate in seen_batch_urls:
            continue
        seen_batch_urls.add(candidate)
        normalized_batch_urls.append(candidate)
        if len(normalized_batch_urls) >= 500:
            break

    if batch_mode and not normalized_batch_urls:
        raise HTTPException(status_code=422, detail="Batch mode requires at least one valid URL")

    if batch_mode:
        url = normalized_batch_urls[0]
    else:
        url = raw_url
        if url and not url.startswith(("http://", "https://")):
            url = f"https://{url}"
        parsed_url = urlparse(url)
        if not parsed_url.scheme or not parsed_url.netloc:
            raise HTTPException(status_code=422, detail="A valid site URL is required in crawl mode")

    print(
        f"[API] Site Audit Pro queued for: {url} "
        f"(mode={mode}, max_pages={max_pages}, batch_mode={batch_mode}, batch_urls={len(normalized_batch_urls)}, "
        f"extended_hreflang_checks={extended_hreflang_checks})"
    )
    task_id = f"sitepro-{datetime.now().timestamp()}"
    create_task_pending(task_id, "site_audit_pro", url, status_message="Site Audit Pro queued")

    def _run_site_audit_pro_task() -> None:
        t0 = time.perf_counter()
        print(
            "[SITE_PRO] "
            + json.dumps(
                {
                    "event": "task_started",
                    "task_id": task_id,
                    "tool": "site_audit_pro",
                    "url": url,
                    "mode": mode,
                    "max_pages": max_pages,
                    "batch_mode": batch_mode,
                    "batch_urls_count": len(normalized_batch_urls),
                    "extended_hreflang_checks": extended_hreflang_checks,
                },
                ensure_ascii=False,
            )
        )
        try:
            update_task_state(task_id, status="RUNNING", progress=5, status_message="Preparing Site Audit Pro")

            def _progress(progress: int, message: str) -> None:
                update_task_state(task_id, status="RUNNING", progress=progress, status_message=message)

            result = check_site_audit_pro(
                url=url,
                task_id=task_id,
                mode=mode,
                max_pages=max_pages,
                batch_mode=batch_mode,
                batch_urls=normalized_batch_urls,
                extended_hreflang_checks=extended_hreflang_checks,
                progress_callback=_progress,
            )
            chunk_manifest = (((result or {}).get("results") or {}).get("artifacts") or {}).get("chunk_manifest", {})
            for chunk in (chunk_manifest.get("chunks") or []):
                for file_meta in (chunk.get("files") or []):
                    file_path = file_meta.get("path")
                    if file_path:
                        append_task_artifact(task_id, file_path, kind="site_pro_chunk")
            update_task_state(
                task_id,
                status="SUCCESS",
                progress=100,
                status_message="Site Audit Pro completed",
                result=result,
                error=None,
            )
            duration_ms = int((time.perf_counter() - t0) * 1000)
            summary = ((result or {}).get("results") or {}).get("summary", {})
            print(
                "[SITE_PRO] "
                + json.dumps(
                    {
                        "event": "task_completed",
                        "task_id": task_id,
                        "tool": "site_audit_pro",
                        "status": "SUCCESS",
                        "duration_ms": duration_ms,
                        "pages": summary.get("total_pages", 0),
                        "issues_total": summary.get("issues_total", 0),
                    },
                    ensure_ascii=False,
                )
            )
        except Exception as exc:
            update_task_state(
                task_id,
                status="FAILURE",
                progress=100,
                status_message="Site Audit Pro failed",
                error=str(exc),
            )
            duration_ms = int((time.perf_counter() - t0) * 1000)
            print(
                "[SITE_PRO] "
                + json.dumps(
                    {
                        "event": "task_completed",
                        "task_id": task_id,
                        "tool": "site_audit_pro",
                        "status": "FAILURE",
                        "duration_ms": duration_ms,
                        "error": str(exc),
                    },
                    ensure_ascii=False,
                )
            )

    background_tasks.add_task(_run_site_audit_pro_task)
    return {"task_id": task_id, "status": "PENDING", "message": "Site Audit Pro started"}
