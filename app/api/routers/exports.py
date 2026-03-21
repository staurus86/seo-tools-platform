"""
Export endpoints for all SEO tool reports (DOCX / XLSX).

Every handler reads the task from the store, delegates to the matching
report generator, and streams the file as an attachment.
"""
import os
import re
from datetime import datetime, timezone
from urllib.parse import urlparse

from fastapi import APIRouter, Request
from fastapi.responses import Response, JSONResponse
from pydantic import BaseModel, Field

from app.api.routers._task_store import append_task_artifact, get_task_result

router = APIRouter(tags=["Exports"])


class ExportRequest(BaseModel):
    task_id: str = Field(..., min_length=1, max_length=256, pattern=r"^[a-zA-Z0-9_\-\.]+$")


# ─── helpers ───────────────────────────────────────────────────────────────


def _safe_domain(url: str) -> str:
    parsed = urlparse(str(url or ""))
    candidate = parsed.netloc or parsed.path or "site"
    return re.sub(r"[^a-zA-Z0-9._-]+", "_", candidate)


def _ts() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%d_%H-%M-%S")


def _export_filename(prefix: str, url: str, extension: str) -> str:
    return f"{prefix}_{_safe_domain(url)}_{_ts()}.{extension}"


def _file_response(filepath: str, media_type: str, filename: str) -> Response:
    with open(filepath, "rb") as f:
        content = f.read()
    return Response(
        content=content,
        media_type=media_type,
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


def _error_response(message: str, status_code: int = 400, **extra) -> JSONResponse:
    payload = {"error": message}
    payload.update(extra)
    return JSONResponse(status_code=status_code, content=payload)


_DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_XLSX = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


# ─── robots ────────────────────────────────────────────────────────────────


@router.post("/export/robots")
async def export_robots_word(data: ExportRequest):
    """Export robots.txt analysis to DOCX."""
    from app.reports.docx_generator import docx_generator

    try:
        task = get_task_result(data.task_id)
        if not task:
            return _error_response("Задача не найдена", status_code=404, task_id=data.task_id)
        if task.get("task_type") != "robots_check":
            return _error_response(f"Неподдерживаемый тип задачи: {task.get('task_type')}", status_code=400)

        task_result = task.get("result", {})
        url = task.get("url", "") or task_result.get("url", "")
        payload = {"url": url, "results": task_result.get("results", task_result)}

        filepath = docx_generator.generate_robots_report(data.task_id, payload)
        if not filepath or not os.path.exists(filepath):
            return _error_response("Не удалось сформировать отчет", status_code=500)

        return _file_response(filepath, _DOCX, _export_filename("robots_report", url, "docx"))
    except Exception as e:
        return _error_response(str(e), status_code=500)


# ─── sitemap ───────────────────────────────────────────────────────────────


@router.post("/export/sitemap-xlsx")
async def export_sitemap_xlsx(data: ExportRequest):
    """Export sitemap validation report to XLSX."""
    from app.reports.xlsx_generator import xlsx_generator

    try:
        task = get_task_result(data.task_id)
        if not task:
            return _error_response("Задача не найдена", status_code=404, task_id=data.task_id)
        if task.get("task_type") != "sitemap_validate":
            return _error_response(f"Неподдерживаемый тип задачи: {task.get('task_type')}", status_code=400)

        task_result = task.get("result", {})
        url = task.get("url", "") or task_result.get("url", "")
        payload = {"url": url, "results": task_result.get("results", task_result)}

        filepath = xlsx_generator.generate_sitemap_report(data.task_id, payload)
        if not filepath or not os.path.exists(filepath):
            return _error_response("Не удалось сформировать отчет", status_code=500)

        return _file_response(filepath, _XLSX, _export_filename("sitemap_report", url, "xlsx"))
    except Exception as e:
        return _error_response(str(e), status_code=500)


@router.post("/export/sitemap-docx")
async def export_sitemap_docx(data: ExportRequest):
    """Export sitemap validation report to DOCX."""
    from app.reports.docx_generator import docx_generator

    try:
        task = get_task_result(data.task_id)
        if not task:
            return _error_response("Задача не найдена", status_code=404, task_id=data.task_id)
        if task.get("task_type") != "sitemap_validate":
            return _error_response(f"Неподдерживаемый тип задачи: {task.get('task_type')}", status_code=400)

        task_result = task.get("result", {})
        url = task.get("url", "") or task_result.get("url", "")
        payload = {"url": url, "results": task_result.get("results", task_result)}

        filepath = docx_generator.generate_sitemap_report(data.task_id, payload)
        if not filepath or not os.path.exists(filepath):
            return _error_response("Не удалось сформировать отчет", status_code=500)

        return _file_response(filepath, _DOCX, _export_filename("sitemap_report", url, "docx"))
    except Exception as e:
        return _error_response(str(e), status_code=500)


# ─── bot checker ───────────────────────────────────────────────────────────


@router.post("/export/bot-xlsx")
async def export_bot_xlsx(data: ExportRequest):
    """Export bot check report to XLSX."""
    from app.reports.xlsx_generator import xlsx_generator

    try:
        task = get_task_result(data.task_id)
        if not task:
            return _error_response("Задача не найдена", status_code=404, task_id=data.task_id)
        if task.get("task_type") != "bot_check":
            return _error_response(f"Неподдерживаемый тип задачи: {task.get('task_type')}", status_code=400)

        task_result = task.get("result", {})
        url = task.get("url", "") or task_result.get("url", "")
        payload = {"url": url, "results": task_result.get("results", task_result)}

        filepath = xlsx_generator.generate_bot_report(data.task_id, payload)
        if not filepath or not os.path.exists(filepath):
            return _error_response("Не удалось сформировать отчет", status_code=500)

        return _file_response(filepath, _XLSX, _export_filename("bot_report", url, "xlsx"))
    except Exception as e:
        return _error_response(str(e), status_code=500)


@router.post("/export/bot-docx")
async def export_bot_docx(data: ExportRequest):
    """Export bot check report to DOCX."""
    from app.reports.docx_generator import docx_generator

    try:
        task = get_task_result(data.task_id)
        if not task:
            return _error_response("Задача не найдена", status_code=404, task_id=data.task_id)
        if task.get("task_type") != "bot_check":
            return _error_response(f"Неподдерживаемый тип задачи: {task.get('task_type')}", status_code=400)

        task_result = task.get("result", {})
        url = task.get("url", "") or task_result.get("url", "")
        payload = {"url": url, "results": task_result.get("results", task_result)}

        filepath = docx_generator.generate_bot_report(data.task_id, payload)
        if not filepath or not os.path.exists(filepath):
            return _error_response("Не удалось сформировать отчет", status_code=500)

        return _file_response(filepath, _DOCX, _export_filename("bot_report", url, "docx"))
    except Exception as e:
        return _error_response(str(e), status_code=500)


# ─── mobile check ──────────────────────────────────────────────────────────


@router.post("/export/mobile-docx")
async def export_mobile_docx(data: ExportRequest, request: Request):
    """Export mobile check report to DOCX."""
    import traceback

    try:
        from app.reports.docx_generator import docx_generator

        task = get_task_result(data.task_id)
        if not task:
            return _error_response("Задача не найдена", status_code=404, task_id=data.task_id)
        if task.get("task_type") != "mobile_check":
            return _error_response(f"Неподдерживаемый тип задачи: {task.get('task_type')}", status_code=400)

        task_result = task.get("result", {})
        url = task.get("url", "") or task_result.get("url", "")
        payload = {
            "url": url,
            "results": task_result.get("results", task_result),
            "server_base_url": str(request.base_url),
        }

        filepath = docx_generator.generate_mobile_report(data.task_id, payload)
        if not filepath or not os.path.exists(filepath):
            return _error_response("Не удалось сформировать отчет", status_code=500)
        append_task_artifact(data.task_id, filepath, kind="export")

        return _file_response(filepath, _DOCX, _export_filename("mobile_report", url, "docx"))
    except Exception as e:
        print(f"[mobile-docx] export failed: {e}\n{traceback.format_exc()}")
        return JSONResponse(status_code=500, content={"error": str(e)})


@router.post("/export/mobile-xlsx")
async def export_mobile_xlsx(data: ExportRequest, request: Request):
    """Export mobile check report to XLSX."""
    from app.reports.xlsx_generator import xlsx_generator

    try:
        task = get_task_result(data.task_id)
        if not task:
            return _error_response("Задача не найдена", status_code=404, task_id=data.task_id)
        if task.get("task_type") != "mobile_check":
            return _error_response(f"Неподдерживаемый тип задачи: {task.get('task_type')}", status_code=400)

        task_result = task.get("result", {})
        url = task.get("url", "") or task_result.get("url", "")
        results = task_result.get("results", task_result) or {}
        issues_count = results.get("issues_count") or len(results.get("issues", []) or [])
        if issues_count <= 0:
            return _error_response("Проблемы не найдены, XLSX-отчет не формируется", status_code=400)

        payload = {"url": url, "results": results, "server_base_url": str(request.base_url)}
        filepath = xlsx_generator.generate_mobile_report(data.task_id, payload)
        if not filepath or not os.path.exists(filepath):
            return _error_response("Не удалось сформировать отчет", status_code=500)
        append_task_artifact(data.task_id, filepath, kind="export")

        return _file_response(filepath, _XLSX, _export_filename("mobile_issues", url, "xlsx"))
    except Exception as e:
        return _error_response(str(e), status_code=500)


# ─── render audit ──────────────────────────────────────────────────────────


@router.post("/export/render-docx")
async def export_render_docx(data: ExportRequest, request: Request):
    """Export render audit report to DOCX."""
    from app.reports.docx_generator import docx_generator

    try:
        task = get_task_result(data.task_id)
        if not task:
            return _error_response("Задача не найдена", status_code=404, task_id=data.task_id)
        if task.get("task_type") != "render_audit":
            return _error_response(f"Неподдерживаемый тип задачи: {task.get('task_type')}", status_code=400)

        task_result = task.get("result", {})
        url = task.get("url", "") or task_result.get("url", "")
        payload = {
            "url": url,
            "results": task_result.get("results", task_result),
            "server_base_url": str(request.base_url),
        }

        filepath = docx_generator.generate_render_report(data.task_id, payload)
        if not filepath or not os.path.exists(filepath):
            return _error_response("Не удалось сформировать отчет", status_code=500)
        append_task_artifact(data.task_id, filepath, kind="export")

        return _file_response(filepath, _DOCX, _export_filename("render_report", url, "docx"))
    except Exception as e:
        return _error_response(str(e), status_code=500)


@router.post("/export/render-xlsx")
async def export_render_xlsx(data: ExportRequest):
    """Export render issues to XLSX."""
    from app.reports.xlsx_generator import xlsx_generator

    try:
        task = get_task_result(data.task_id)
        if not task:
            return _error_response("Задача не найдена", status_code=404, task_id=data.task_id)
        if task.get("task_type") != "render_audit":
            return _error_response(f"Неподдерживаемый тип задачи: {task.get('task_type')}", status_code=400)

        task_result = task.get("result", {})
        url = task.get("url", "") or task_result.get("url", "")
        results = task_result.get("results", task_result) or {}
        issues_count = results.get("issues_count") or len(results.get("issues", []) or [])
        if issues_count <= 0:
            return _error_response("Проблемы не найдены, XLSX-отчет не формируется", status_code=400)

        payload = {"url": url, "results": results}
        filepath = xlsx_generator.generate_render_report(data.task_id, payload)
        if not filepath or not os.path.exists(filepath):
            return _error_response("Не удалось сформировать отчет", status_code=500)
        append_task_artifact(data.task_id, filepath, kind="export")

        return _file_response(filepath, _XLSX, _export_filename("render_issues", url, "xlsx"))
    except Exception as e:
        return _error_response(str(e), status_code=500)


# ─── onpage audit ──────────────────────────────────────────────────────────


@router.post("/export/onpage-docx")
async def export_onpage_docx(data: ExportRequest):
    """Export OnPage audit report to DOCX."""
    from app.reports.docx_generator import docx_generator

    try:
        task = get_task_result(data.task_id)
        if not task:
            return _error_response("Задача не найдена", status_code=404, task_id=data.task_id)
        if task.get("task_type") != "onpage_audit":
            return _error_response(f"Неподдерживаемый тип задачи: {task.get('task_type')}", status_code=400)

        task_result = task.get("result", {})
        url = task.get("url", "") or task_result.get("url", "")
        payload = {"url": url, "results": task_result.get("results", task_result)}

        filepath = docx_generator.generate_onpage_report(data.task_id, payload)
        if not filepath or not os.path.exists(filepath):
            return _error_response("Не удалось сформировать отчет", status_code=500)
        append_task_artifact(data.task_id, filepath, kind="export")

        return _file_response(filepath, _DOCX, _export_filename("onpage_report", url, "docx"))
    except Exception as e:
        return _error_response(str(e), status_code=500)


@router.post("/export/onpage-xlsx")
async def export_onpage_xlsx(data: ExportRequest):
    """Export OnPage audit report to XLSX."""
    from app.reports.xlsx_generator import xlsx_generator

    try:
        task = get_task_result(data.task_id)
        if not task:
            return _error_response("Задача не найдена", status_code=404, task_id=data.task_id)
        if task.get("task_type") != "onpage_audit":
            return _error_response(f"Неподдерживаемый тип задачи: {task.get('task_type')}", status_code=400)

        task_result = task.get("result", {})
        url = task.get("url", "") or task_result.get("url", "")
        payload = {"url": url, "results": task_result.get("results", task_result)}

        filepath = xlsx_generator.generate_onpage_report(data.task_id, payload)
        if not filepath or not os.path.exists(filepath):
            return _error_response("Не удалось сформировать отчет", status_code=500)
        append_task_artifact(data.task_id, filepath, kind="export")

        return _file_response(filepath, _XLSX, _export_filename("onpage_report", url, "xlsx"))
    except Exception as e:
        return _error_response(str(e), status_code=500)


# ─── redirect checker ──────────────────────────────────────────────────────


@router.post("/export/redirect-checker-docx")
async def export_redirect_checker_docx(data: ExportRequest):
    """Export Redirect Checker report to DOCX."""
    from app.reports.docx_generator import docx_generator

    try:
        task = get_task_result(data.task_id)
        if not task:
            return _error_response("Задача не найдена", status_code=404, task_id=data.task_id)
        if task.get("task_type") != "redirect_checker":
            return _error_response(f"Неподдерживаемый тип задачи: {task.get('task_type')}", status_code=400)

        task_result = task.get("result", {}) or {}
        url = task.get("url", "") or task_result.get("url", "")
        payload = {"url": url, "results": task_result.get("results", task_result)}

        filepath = docx_generator.generate_redirect_checker_report(data.task_id, payload)
        if not filepath or not os.path.exists(filepath):
            return _error_response("Не удалось сформировать отчет", status_code=500)
        append_task_artifact(data.task_id, filepath, kind="export")

        return _file_response(filepath, _DOCX, _export_filename("redirect_checker_report", url, "docx"))
    except Exception as e:
        return _error_response(str(e), status_code=500)


@router.post("/export/redirect-checker-xlsx")
async def export_redirect_checker_xlsx(data: ExportRequest):
    """Export Redirect Checker report to XLSX."""
    from app.reports.xlsx_generator import xlsx_generator

    try:
        task = get_task_result(data.task_id)
        if not task:
            return _error_response("Задача не найдена", status_code=404, task_id=data.task_id)
        if task.get("task_type") != "redirect_checker":
            return _error_response(f"Неподдерживаемый тип задачи: {task.get('task_type')}", status_code=400)

        task_result = task.get("result", {}) or {}
        url = task.get("url", "") or task_result.get("url", "")
        payload = {"url": url, "results": task_result.get("results", task_result)}

        filepath = xlsx_generator.generate_redirect_checker_report(data.task_id, payload)
        if not filepath or not os.path.exists(filepath):
            return _error_response("Не удалось сформировать отчет", status_code=500)
        append_task_artifact(data.task_id, filepath, kind="export")

        return _file_response(filepath, _XLSX, _export_filename("redirect_checker_report", url, "xlsx"))
    except Exception as e:
        return _error_response(str(e), status_code=500)


# ─── core web vitals ───────────────────────────────────────────────────────


@router.post("/export/core-web-vitals-docx")
async def export_core_web_vitals_docx(data: ExportRequest):
    """Export Core Web Vitals report to DOCX."""
    from app.reports.docx_generator import docx_generator

    try:
        task = get_task_result(data.task_id)
        if not task:
            return _error_response("Задача не найдена", status_code=404, task_id=data.task_id)
        if task.get("task_type") != "core_web_vitals":
            return _error_response(f"Неподдерживаемый тип задачи: {task.get('task_type')}", status_code=400)

        task_result = task.get("result", {}) or {}
        url = task.get("url", "") or task_result.get("url", "")
        payload = {"url": url, "results": task_result.get("results", task_result)}

        filepath = docx_generator.generate_core_web_vitals_report(data.task_id, payload)
        if not filepath or not os.path.exists(filepath):
            return _error_response("Не удалось сформировать отчет", status_code=500)
        append_task_artifact(data.task_id, filepath, kind="export")

        return _file_response(filepath, _DOCX, _export_filename("core_web_vitals_report", url, "docx"))
    except Exception as e:
        return _error_response(str(e), status_code=500)


@router.post("/export/core-web-vitals-xlsx")
async def export_core_web_vitals_xlsx(data: ExportRequest):
    """Export Core Web Vitals report to XLSX."""
    from app.reports.xlsx_generator import xlsx_generator

    try:
        task = get_task_result(data.task_id)
        if not task:
            return _error_response("Задача не найдена", status_code=404, task_id=data.task_id)
        if task.get("task_type") != "core_web_vitals":
            return _error_response(f"Неподдерживаемый тип задачи: {task.get('task_type')}", status_code=400)

        task_result = task.get("result", {}) or {}
        url = task.get("url", "") or task_result.get("url", "")
        payload = {"url": url, "results": task_result.get("results", task_result)}

        filepath = xlsx_generator.generate_core_web_vitals_report(data.task_id, payload)
        if not filepath or not os.path.exists(filepath):
            return _error_response("Не удалось сформировать отчет", status_code=500)
        append_task_artifact(data.task_id, filepath, kind="export")

        return _file_response(filepath, _XLSX, _export_filename("core_web_vitals_report", url, "xlsx"))
    except Exception as e:
        return _error_response(str(e), status_code=500)


# ─── site audit pro ────────────────────────────────────────────────────────


@router.post("/export/site-audit-pro-docx")
async def export_site_audit_pro_docx(data: ExportRequest):
    """Export Site Audit Pro report to DOCX."""
    from app.config import settings
    from app.reports.docx_generator import docx_generator

    try:
        if not getattr(settings, "SITE_AUDIT_PRO_ENABLED", True):
            return _error_response("Site Audit Pro отключён через feature flag", status_code=403)

        task = get_task_result(data.task_id)
        if not task:
            return _error_response("Задача не найдена", status_code=404, task_id=data.task_id)
        if task.get("task_type") != "site_audit_pro":
            return _error_response(f"Неподдерживаемый тип задачи: {task.get('task_type')}", status_code=400)

        task_result = task.get("result", {})
        url = task.get("url", "") or task_result.get("url", "")
        payload = {"url": url, "results": task_result.get("results", task_result)}

        filepath = docx_generator.generate_site_audit_pro_report(data.task_id, payload)
        if not filepath or not os.path.exists(filepath):
            return _error_response("Не удалось сформировать отчет", status_code=500)
        append_task_artifact(data.task_id, filepath, kind="export")

        return _file_response(filepath, _DOCX, _export_filename("site_audit_pro", url, "docx"))
    except Exception as e:
        return _error_response(str(e), status_code=500)


@router.post("/export/site-audit-pro-xlsx")
async def export_site_audit_pro_xlsx(data: ExportRequest):
    """Export Site Audit Pro report to XLSX."""
    from app.config import settings
    from app.reports.xlsx_generator import xlsx_generator

    try:
        if not getattr(settings, "SITE_AUDIT_PRO_ENABLED", True):
            return _error_response("Site Audit Pro отключён через feature flag", status_code=403)

        task = get_task_result(data.task_id)
        if not task:
            return _error_response("Задача не найдена", status_code=404, task_id=data.task_id)
        if task.get("task_type") != "site_audit_pro":
            return _error_response(f"Неподдерживаемый тип задачи: {task.get('task_type')}", status_code=400)

        task_result = task.get("result", {})
        url = task.get("url", "") or task_result.get("url", "")
        payload = {"url": url, "results": task_result.get("results", task_result) or {}}

        filepath = xlsx_generator.generate_site_audit_pro_report(data.task_id, payload)
        if not filepath or not os.path.exists(filepath):
            return _error_response("Не удалось сформировать отчет", status_code=500)
        append_task_artifact(data.task_id, filepath, kind="export")

        return _file_response(filepath, _XLSX, _export_filename("site_audit_pro", url, "xlsx"))
    except Exception as e:
        return _error_response(str(e), status_code=500)


# ─── clusterizer ───────────────────────────────────────────────────────────


@router.post("/export/clusterizer-xlsx")
async def export_clusterizer_xlsx(data: ExportRequest):
    """Export keyword clusterizer report to XLSX."""
    from app.reports.xlsx_generator import xlsx_generator

    try:
        task = get_task_result(data.task_id)
        if not task:
            return _error_response("Задача не найдена", status_code=404, task_id=data.task_id)
        if task.get("task_type") != "clusterizer":
            return _error_response(f"Неподдерживаемый тип задачи: {task.get('task_type')}", status_code=400)

        task_result = task.get("result", {})
        payload = {
            "url": task.get("url", "") or task_result.get("url", ""),
            "results": task_result.get("results", task_result),
        }

        filepath = xlsx_generator.generate_clusterizer_report(data.task_id, payload)
        if not filepath or not os.path.exists(filepath):
            return _error_response("Не удалось сформировать отчет", status_code=500)
        append_task_artifact(data.task_id, filepath, kind="export")

        report_url = payload.get("url", "")
        return _file_response(filepath, _XLSX, _export_filename("clusterizer_report", report_url, "xlsx"))
    except Exception as e:
        return _error_response(str(e), status_code=500)


# ─── link profile ──────────────────────────────────────────────────────────


@router.post("/export/link-profile-docx")
async def export_link_profile_docx(data: ExportRequest):
    """Export link profile audit report to DOCX."""
    from app.reports.docx_generator import docx_generator

    try:
        task = get_task_result(data.task_id)
        if not task:
            return _error_response("Задача не найдена", status_code=404, task_id=data.task_id)
        if task.get("task_type") != "link_profile_audit":
            return _error_response(f"Неподдерживаемый тип задачи: {task.get('task_type')}", status_code=400)

        task_result = task.get("result", {})
        url = task.get("url", "") or task_result.get("url", "")
        payload = {"url": url, "results": task_result.get("results", task_result)}

        filepath = docx_generator.generate_link_profile_report(data.task_id, payload)
        if not filepath or not os.path.exists(filepath):
            return _error_response("Не удалось сформировать отчет", status_code=500)
        append_task_artifact(data.task_id, filepath, kind="export")

        return _file_response(filepath, _DOCX, _export_filename("link_profile_report", url, "docx"))
    except Exception as e:
        return _error_response(str(e), status_code=500)


@router.post("/export/link-profile-xlsx")
async def export_link_profile_xlsx(data: ExportRequest):
    """Export link profile audit report to XLSX."""
    from app.reports.xlsx_generator import xlsx_generator

    try:
        task = get_task_result(data.task_id)
        if not task:
            return _error_response("Задача не найдена", status_code=404, task_id=data.task_id)
        if task.get("task_type") != "link_profile_audit":
            return _error_response(f"Неподдерживаемый тип задачи: {task.get('task_type')}", status_code=400)

        task_result = task.get("result", {})
        url = task.get("url", "") or task_result.get("url", "")
        payload = {"url": url, "results": task_result.get("results", task_result)}

        filepath = xlsx_generator.generate_link_profile_report(data.task_id, payload)
        if not filepath or not os.path.exists(filepath):
            return _error_response("Не удалось сформировать отчет", status_code=500)
        append_task_artifact(data.task_id, filepath, kind="export")

        return _file_response(filepath, _XLSX, _export_filename("link_profile_report", url, "xlsx"))
    except Exception as e:
        return _error_response(str(e), status_code=500)


# ═══════════════════════════════════════════════════════════════════════════
# GET endpoints — /tasks/{tool}/{task_id}/export/{format}
# ═══════════════════════════════════════════════════════════════════════════

# Mapping: path tool slug → (task_type, report_prefix, generator_method_suffix)
_TOOL_META = {
    "bot-check": ("bot_check", "bot_report", "bot_report"),
    "mobile-check": ("mobile_check", "mobile_report", "mobile_report"),
    "render-audit": ("render_audit", "render_report", "render_report"),
    "onpage-audit": ("onpage_audit", "onpage_report", "onpage_report"),
    "redirect-checker": ("redirect_checker", "redirect_checker_report", "redirect_checker_report"),
    "core-web-vitals": ("core_web_vitals", "core_web_vitals_report", "core_web_vitals_report"),
}


def _get_task_or_error(task_id: str, expected_type: str):
    """Return (task, task_result, url, error_response).

    If error_response is not None the caller should return it immediately.
    """
    task = get_task_result(task_id)
    if not task:
        return None, None, None, _error_response("Task not found", status_code=404, task_id=task_id)
    if task.get("task_type") != expected_type:
        return None, None, None, _error_response(
            f"Unsupported task type: {task.get('task_type')}", status_code=400
        )
    task_result = task.get("result", {}) or {}
    url = task.get("url", "") or task_result.get("url", "")
    return task, task_result, url, None


# ─── bot-check GET ────────────────────────────────────────────────────────


@router.get("/tasks/bot-check/{task_id}/export/xlsx")
async def get_bot_check_xlsx(task_id: str):
    """GET export: Bot Checker → XLSX."""
    from app.reports.xlsx_generator import xlsx_generator

    try:
        task, task_result, url, err = _get_task_or_error(task_id, "bot_check")
        if err:
            return err
        payload = {"url": url, "results": task_result.get("results", task_result)}
        filepath = xlsx_generator.generate_bot_report(task_id, payload)
        if not filepath or not os.path.exists(filepath):
            return _error_response("Failed to generate report", status_code=500)
        append_task_artifact(task_id, filepath, kind="export")
        return _file_response(filepath, _XLSX, _export_filename("bot_report", url, "xlsx"))
    except Exception as e:
        return _error_response(str(e), status_code=500)


@router.get("/tasks/bot-check/{task_id}/export/docx")
async def get_bot_check_docx(task_id: str):
    """GET export: Bot Checker → DOCX."""
    from app.reports.docx_generator import docx_generator

    try:
        task, task_result, url, err = _get_task_or_error(task_id, "bot_check")
        if err:
            return err
        payload = {"url": url, "results": task_result.get("results", task_result)}
        filepath = docx_generator.generate_bot_report(task_id, payload)
        if not filepath or not os.path.exists(filepath):
            return _error_response("Failed to generate report", status_code=500)
        append_task_artifact(task_id, filepath, kind="export")
        return _file_response(filepath, _DOCX, _export_filename("bot_report", url, "docx"))
    except Exception as e:
        return _error_response(str(e), status_code=500)


# ─── mobile-check GET ────────────────────────────────────────────────────


@router.get("/tasks/mobile-check/{task_id}/export/xlsx")
async def get_mobile_check_xlsx(task_id: str, request: Request):
    """GET export: Mobile Audit → XLSX."""
    from app.reports.xlsx_generator import xlsx_generator

    try:
        task, task_result, url, err = _get_task_or_error(task_id, "mobile_check")
        if err:
            return err
        results = task_result.get("results", task_result) or {}
        payload = {"url": url, "results": results, "server_base_url": str(request.base_url)}
        filepath = xlsx_generator.generate_mobile_report(task_id, payload)
        if not filepath or not os.path.exists(filepath):
            return _error_response("Failed to generate report", status_code=500)
        append_task_artifact(task_id, filepath, kind="export")
        return _file_response(filepath, _XLSX, _export_filename("mobile_report", url, "xlsx"))
    except Exception as e:
        return _error_response(str(e), status_code=500)


@router.get("/tasks/mobile-check/{task_id}/export/docx")
async def get_mobile_check_docx(task_id: str, request: Request):
    """GET export: Mobile Audit → DOCX."""
    from app.reports.docx_generator import docx_generator

    try:
        task, task_result, url, err = _get_task_or_error(task_id, "mobile_check")
        if err:
            return err
        payload = {
            "url": url,
            "results": task_result.get("results", task_result),
            "server_base_url": str(request.base_url),
        }
        filepath = docx_generator.generate_mobile_report(task_id, payload)
        if not filepath or not os.path.exists(filepath):
            return _error_response("Failed to generate report", status_code=500)
        append_task_artifact(task_id, filepath, kind="export")
        return _file_response(filepath, _DOCX, _export_filename("mobile_report", url, "docx"))
    except Exception as e:
        return _error_response(str(e), status_code=500)


# ─── render-audit GET ─────────────────────────────────────────────────────


@router.get("/tasks/render-audit/{task_id}/export/xlsx")
async def get_render_audit_xlsx(task_id: str):
    """GET export: Render Audit → XLSX."""
    from app.reports.xlsx_generator import xlsx_generator

    try:
        task, task_result, url, err = _get_task_or_error(task_id, "render_audit")
        if err:
            return err
        results = task_result.get("results", task_result) or {}
        payload = {"url": url, "results": results}
        filepath = xlsx_generator.generate_render_report(task_id, payload)
        if not filepath or not os.path.exists(filepath):
            return _error_response("Failed to generate report", status_code=500)
        append_task_artifact(task_id, filepath, kind="export")
        return _file_response(filepath, _XLSX, _export_filename("render_report", url, "xlsx"))
    except Exception as e:
        return _error_response(str(e), status_code=500)


@router.get("/tasks/render-audit/{task_id}/export/docx")
async def get_render_audit_docx(task_id: str, request: Request):
    """GET export: Render Audit → DOCX."""
    from app.reports.docx_generator import docx_generator

    try:
        task, task_result, url, err = _get_task_or_error(task_id, "render_audit")
        if err:
            return err
        payload = {
            "url": url,
            "results": task_result.get("results", task_result),
            "server_base_url": str(request.base_url),
        }
        filepath = docx_generator.generate_render_report(task_id, payload)
        if not filepath or not os.path.exists(filepath):
            return _error_response("Failed to generate report", status_code=500)
        append_task_artifact(task_id, filepath, kind="export")
        return _file_response(filepath, _DOCX, _export_filename("render_report", url, "docx"))
    except Exception as e:
        return _error_response(str(e), status_code=500)


# ─── onpage-audit GET ─────────────────────────────────────────────────────


@router.get("/tasks/onpage-audit/{task_id}/export/xlsx")
async def get_onpage_audit_xlsx(task_id: str):
    """GET export: OnPage Audit → XLSX."""
    from app.reports.xlsx_generator import xlsx_generator

    try:
        task, task_result, url, err = _get_task_or_error(task_id, "onpage_audit")
        if err:
            return err
        payload = {"url": url, "results": task_result.get("results", task_result)}
        filepath = xlsx_generator.generate_onpage_report(task_id, payload)
        if not filepath or not os.path.exists(filepath):
            return _error_response("Failed to generate report", status_code=500)
        append_task_artifact(task_id, filepath, kind="export")
        return _file_response(filepath, _XLSX, _export_filename("onpage_report", url, "xlsx"))
    except Exception as e:
        return _error_response(str(e), status_code=500)


@router.get("/tasks/onpage-audit/{task_id}/export/docx")
async def get_onpage_audit_docx(task_id: str):
    """GET export: OnPage Audit → DOCX."""
    from app.reports.docx_generator import docx_generator

    try:
        task, task_result, url, err = _get_task_or_error(task_id, "onpage_audit")
        if err:
            return err
        payload = {"url": url, "results": task_result.get("results", task_result)}
        filepath = docx_generator.generate_onpage_report(task_id, payload)
        if not filepath or not os.path.exists(filepath):
            return _error_response("Failed to generate report", status_code=500)
        append_task_artifact(task_id, filepath, kind="export")
        return _file_response(filepath, _DOCX, _export_filename("onpage_report", url, "docx"))
    except Exception as e:
        return _error_response(str(e), status_code=500)


# ─── redirect-checker GET ─────────────────────────────────────────────────


@router.get("/tasks/redirect-checker/{task_id}/export/xlsx")
async def get_redirect_checker_xlsx(task_id: str):
    """GET export: Redirect Checker → XLSX."""
    from app.reports.xlsx_generator import xlsx_generator

    try:
        task, task_result, url, err = _get_task_or_error(task_id, "redirect_checker")
        if err:
            return err
        payload = {"url": url, "results": task_result.get("results", task_result)}
        filepath = xlsx_generator.generate_redirect_checker_report(task_id, payload)
        if not filepath or not os.path.exists(filepath):
            return _error_response("Failed to generate report", status_code=500)
        append_task_artifact(task_id, filepath, kind="export")
        return _file_response(filepath, _XLSX, _export_filename("redirect_checker_report", url, "xlsx"))
    except Exception as e:
        return _error_response(str(e), status_code=500)


@router.get("/tasks/redirect-checker/{task_id}/export/docx")
async def get_redirect_checker_docx(task_id: str):
    """GET export: Redirect Checker → DOCX."""
    from app.reports.docx_generator import docx_generator

    try:
        task, task_result, url, err = _get_task_or_error(task_id, "redirect_checker")
        if err:
            return err
        payload = {"url": url, "results": task_result.get("results", task_result)}
        filepath = docx_generator.generate_redirect_checker_report(task_id, payload)
        if not filepath or not os.path.exists(filepath):
            return _error_response("Failed to generate report", status_code=500)
        append_task_artifact(task_id, filepath, kind="export")
        return _file_response(filepath, _DOCX, _export_filename("redirect_checker_report", url, "docx"))
    except Exception as e:
        return _error_response(str(e), status_code=500)


# ─── core-web-vitals GET ──────────────────────────────────────────────────


@router.get("/tasks/core-web-vitals/{task_id}/export/xlsx")
async def get_core_web_vitals_xlsx(task_id: str):
    """GET export: Core Web Vitals → XLSX."""
    from app.reports.xlsx_generator import xlsx_generator

    try:
        task, task_result, url, err = _get_task_or_error(task_id, "core_web_vitals")
        if err:
            return err
        payload = {"url": url, "results": task_result.get("results", task_result)}
        filepath = xlsx_generator.generate_core_web_vitals_report(task_id, payload)
        if not filepath or not os.path.exists(filepath):
            return _error_response("Failed to generate report", status_code=500)
        append_task_artifact(task_id, filepath, kind="export")
        return _file_response(filepath, _XLSX, _export_filename("core_web_vitals_report", url, "xlsx"))
    except Exception as e:
        return _error_response(str(e), status_code=500)


@router.get("/tasks/core-web-vitals/{task_id}/export/docx")
async def get_core_web_vitals_docx(task_id: str):
    """GET export: Core Web Vitals → DOCX."""
    from app.reports.docx_generator import docx_generator

    try:
        task, task_result, url, err = _get_task_or_error(task_id, "core_web_vitals")
        if err:
            return err
        payload = {"url": url, "results": task_result.get("results", task_result)}
        filepath = docx_generator.generate_core_web_vitals_report(task_id, payload)
        if not filepath or not os.path.exists(filepath):
            return _error_response("Failed to generate report", status_code=500)
        append_task_artifact(task_id, filepath, kind="export")
        return _file_response(filepath, _DOCX, _export_filename("core_web_vitals_report", url, "docx"))
    except Exception as e:
        return _error_response(str(e), status_code=500)


# ═══════════════════════════════════════════════════════════════════════════
# Unified Full SEO Audit — inline report builders
# ═══════════════════════════════════════════════════════════════════════════

_SCORE_LABELS = {
    "onpage": "OnPage Audit",
    "render": "Render Audit",
    "mobile_friendly": "Mobile Friendly",
    "bot_accessibility": "Bot Accessibility",
    "redirect": "Redirect Checker",
    "cwv_mobile": "CWV Mobile",
    "cwv_desktop": "CWV Desktop",
    "cwv_avg": "CWV Average",
    "robots_ok": "Robots.txt",
}


def _score_status(value) -> str:
    """Map a numeric score to a human-readable status label."""
    try:
        v = float(value)
    except (TypeError, ValueError):
        return "N/A"
    if v >= 90:
        return "Excellent"
    if v >= 70:
        return "Good"
    if v >= 50:
        return "Needs Work"
    return "Critical"


def _build_unified_xlsx(task_id: str, url: str, results: dict) -> str:
    """Build an XLSX workbook for the Unified Full SEO Audit and return the filepath."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    from app.config import settings

    reports_dir = settings.REPORTS_DIR
    os.makedirs(reports_dir, exist_ok=True)

    wb = Workbook()

    # -- shared styles --
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="0F4C81", end_color="0F4C81", fill_type="solid")
    header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin_border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )
    zebra_fill = PatternFill(start_color="F8FBFF", end_color="F8FBFF", fill_type="solid")

    def _style_header(ws, row_idx, col_count):
        for ci in range(1, col_count + 1):
            c = ws.cell(row=row_idx, column=ci)
            c.font = header_font
            c.fill = header_fill
            c.alignment = header_align
            c.border = thin_border

    def _style_data_rows(ws, start_row, end_row, col_count):
        for ri in range(start_row, end_row + 1):
            for ci in range(1, col_count + 1):
                c = ws.cell(row=ri, column=ci)
                c.border = thin_border
                c.alignment = Alignment(vertical="center", wrap_text=True)
                if (ri - start_row) % 2 == 0:
                    c.fill = zebra_fill

    def _auto_width(ws, col_count):
        for ci in range(1, col_count + 1):
            max_len = 0
            for row in ws.iter_rows(min_col=ci, max_col=ci, values_only=False):
                for cell in row:
                    val = str(cell.value or "")
                    max_len = max(max_len, len(val))
            ws.column_dimensions[get_column_letter(ci)].width = min(max_len + 4, 60)

    generated_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")

    # ── Sheet 1: Overview ──────────────────────────────────────────────
    ws_overview = wb.active
    ws_overview.title = "Overview"
    ws_overview.sheet_properties.tabColor = "0F4C81"
    overview_headers = ["Metric", "Value"]
    for ci, h in enumerate(overview_headers, 1):
        ws_overview.cell(row=1, column=ci, value=h)
    _style_header(ws_overview, 1, 2)

    overview_rows = [
        ("URL", url),
        ("Overall Score", results.get("overall_score", "N/A")),
        ("Grade", results.get("overall_grade", "N/A")),
        ("Tools Run", results.get("tools_run", "N/A")),
        ("Tools Failed", results.get("tools_failed", "N/A")),
        ("Duration", f"{results.get('duration_ms', 'N/A')}ms"),
        ("Date", generated_at),
    ]
    for ri, (metric, value) in enumerate(overview_rows, 2):
        ws_overview.cell(row=ri, column=1, value=metric)
        ws_overview.cell(row=ri, column=2, value=str(value))
    _style_data_rows(ws_overview, 2, len(overview_rows) + 1, 2)
    _auto_width(ws_overview, 2)

    # ── Sheet 2: Scores ────────────────────────────────────────────────
    ws_scores = wb.create_sheet("Scores")
    ws_scores.sheet_properties.tabColor = "0E7490"
    score_headers = ["Tool", "Score", "Status"]
    for ci, h in enumerate(score_headers, 1):
        ws_scores.cell(row=1, column=ci, value=h)
    _style_header(ws_scores, 1, 3)

    scores = results.get("scores", {})
    row_idx = 2
    for key, label in _SCORE_LABELS.items():
        val = scores.get(key)
        ws_scores.cell(row=row_idx, column=1, value=label)
        ws_scores.cell(row=row_idx, column=2, value=val if val is not None else "N/A")
        ws_scores.cell(row=row_idx, column=3, value=_score_status(val))
        row_idx += 1
    if row_idx > 2:
        _style_data_rows(ws_scores, 2, row_idx - 1, 3)
    _auto_width(ws_scores, 3)

    # ── Sheet 3: ТЗ — Полное техническое задание ─────────────────────
    ws_tasks = wb.create_sheet("ТЗ — Техническое задание")
    ws_tasks.sheet_properties.tabColor = "C2410C"

    priority_fills = {
        "P0": PatternFill(start_color="FEE2E2", end_color="FEE2E2", fill_type="solid"),
        "P1": PatternFill(start_color="FFEDD5", end_color="FFEDD5", fill_type="solid"),
        "P2": PatternFill(start_color="FEF9C3", end_color="FEF9C3", fill_type="solid"),
        "P3": PatternFill(start_color="DCFCE7", end_color="DCFCE7", fill_type="solid"),
    }
    priority_fonts = {
        "P0": Font(bold=True, color="B91C1C"),
        "P1": Font(bold=True, color="C2410C"),
        "P2": Font(bold=False, color="92400E"),
        "P3": Font(bold=False, color="166534"),
    }

    # Title row
    ws_tasks.merge_cells("A1:J1")
    title_cell = ws_tasks.cell(row=1, column=1, value=f"Техническое задание по SEO-аудиту — {url}")
    title_cell.font = Font(bold=True, size=14, color="0F4C81")
    title_cell.alignment = Alignment(horizontal="left", vertical="center")
    ws_tasks.row_dimensions[1].height = 30

    # Subtitle
    ws_tasks.merge_cells("A2:J2")
    sub_cell = ws_tasks.cell(row=2, column=1, value=f"Оценка: {results.get('overall_score', 'N/A')}/100 ({results.get('overall_grade', '?')}) | Дата: {generated_at}")
    sub_cell.font = Font(size=10, color="666666")
    sub_cell.alignment = Alignment(horizontal="left")

    # Summary row
    dev_tasks = results.get("dev_tasks", []) or []
    p0 = sum(1 for t in dev_tasks if t.get("priority") == "P0")
    p1 = sum(1 for t in dev_tasks if t.get("priority") == "P1")
    p2 = sum(1 for t in dev_tasks if t.get("priority") == "P2")
    p3 = sum(1 for t in dev_tasks if t.get("priority") == "P3")
    ws_tasks.merge_cells("A3:J3")
    ws_tasks.cell(row=3, column=1, value=f"Всего задач: {len(dev_tasks)} | P0 (блокер): {p0} | P1 (критично): {p1} | P2 (важно): {p2} | P3 (рекомендация): {p3}")
    ws_tasks.cell(row=3, column=1).font = Font(size=10, bold=True)
    ws_tasks.row_dimensions[3].height = 22

    # Headers
    task_headers = [
        "#", "Приоритет", "Статус", "Категория", "Инструмент",
        "Что исправить", "Подробное описание", "Как исправить",
        "Ответственный", "Срок",
    ]
    for ci, h in enumerate(task_headers, 1):
        ws_tasks.cell(row=4, column=ci, value=h)
    _style_header(ws_tasks, 4, len(task_headers))

    # Priority to deadline mapping
    deadline_map = {"P0": "Немедленно", "P1": "1-3 дня", "P2": "1-2 недели", "P3": "При возможности"}

    # How-to-fix recommendations
    fix_map = {
        "Title length out of range": "Оптимизируйте <title> тег: 30-60 символов, включите основной ключ в начало.",
        "Description length out of range": "Перепишите meta description: 120-160 символов, с CTA и ключевым словом.",
        "Structured data is missing": "Добавьте JSON-LD разметку (Organization, WebPage, BreadcrumbList). Используйте Google Structured Data Testing Tool для проверки.",
        "Twitter Card type not specified": "Добавьте <meta name=\"twitter:card\" content=\"summary_large_image\"> в <head>.",
        "Twitter Card title missing": "Добавьте <meta name=\"twitter:title\" content=\"...\"> — можно дублировать OG title.",
        "Twitter Card image not set": "Добавьте <meta name=\"twitter:image\" content=\"URL_изображения\"> — мин. 800×418px.",
        "Low stopword ratio": "Увеличьте связность текста: добавьте предлоги, союзы, местоимения для естественности.",
        "Bots cannot reach URL": "Проверьте firewall/WAF, убедитесь что боты не блокируются по IP/UA. Проверьте robots.txt.",
        "Множественные слеши": "Настройте серверное правило: /page//path → 301 → /page/path.",
        "Регистр URL": "Все URL должны быть в нижнем регистре. Добавьте 301 redirect /Page → /page.",
        "Index-файлы": "Добавьте в nginx/Apache: rewrite ^/index\\.(html|php)$ / permanent;",
        "Trailing slash": "Выберите единую политику (со слешем или без) и настройте 301 для несоответствий.",
        "Старые расширения": "Настройте 301 redirect для .html/.php URL на чистые URL без расширений.",
        "Цепочки редиректов": "Замените цепочку A→B→C на прямой A→C. Обновите внутренние ссылки.",
        "Query params canonicalization": "Настройте canonical URL без UTM/tracking параметров. Используйте rel=canonical.",
        "JavaScript / Meta Refresh Redirects": "Замените JS redirect (location.href) на серверный 301 redirect.",
        "Pagination ?page=1 redirect": "Добавьте 301: ?page=1 → каноническая версия без page=1.",
        "Legacy index.html/index.php redirect": "Добавьте 301: /index.html → / (корень сайта).",
        "Слишком маленькие интерактивные элементы": "Увеличьте кнопки/ссылки до мин. 44×44px (CSS min-height/min-width/padding).",
        "Слишком мелкий текст": "Установите мин. font-size: 16px для body текста на мобильных устройствах.",
        "Low contrast text found": "Увеличьте контраст текста: мин. 4.5:1 (обычный текст) или 3:1 (крупный). Используйте WebAIM Contrast Checker.",
        "Контент появляется только после JavaScript": "Реализуйте SSR (Server-Side Rendering) или SSG для критичного контента. Боты не всегда выполняют JS.",
        "Ссылки появляются только после JavaScript": "Отрендерите навигационные ссылки на сервере (SSR). Критичные ссылки должны быть в HTML.",
        "Content hidden via CSS": "Уберите display:none / visibility:hidden с важного контента. Используйте альтернативные методы для скрытия.",
        "Links change after hydration": "Проверьте SSR гидрацию — количество ссылок не должно меняться после загрузки JS.",
        "Низкий балл рендеринга": "Основной контент должен быть доступен без JavaScript. Используйте SSR/SSG фреймворк.",
    }

    row_idx = 5
    for idx, dt in enumerate(dev_tasks, 1):
        priority = str(dt.get("priority", "P3"))
        title = str(dt.get("title", ""))
        description = str(dt.get("description", ""))
        fix = str(dt.get("fix", ""))

        # Find matching how-to-fix
        how_to_fix = fix
        if not how_to_fix:
            for pattern, recommendation in fix_map.items():
                if pattern.lower() in title.lower():
                    how_to_fix = recommendation
                    break
        if not how_to_fix:
            how_to_fix = description

        ws_tasks.cell(row=row_idx, column=1, value=idx)
        ws_tasks.cell(row=row_idx, column=2, value=priority)
        ws_tasks.cell(row=row_idx, column=3, value="Открыта")
        ws_tasks.cell(row=row_idx, column=4, value=str(dt.get("category", "")))
        ws_tasks.cell(row=row_idx, column=5, value=str(dt.get("source_tool", "")))
        ws_tasks.cell(row=row_idx, column=6, value=title)
        ws_tasks.cell(row=row_idx, column=7, value=description)
        ws_tasks.cell(row=row_idx, column=8, value=how_to_fix)
        ws_tasks.cell(row=row_idx, column=9, value=str(dt.get("owner", "")))
        ws_tasks.cell(row=row_idx, column=10, value=deadline_map.get(priority, ""))

        # Styling
        p_fill = priority_fills.get(priority)
        p_font = priority_fonts.get(priority)
        for ci in range(1, len(task_headers) + 1):
            c = ws_tasks.cell(row=row_idx, column=ci)
            c.border = thin_border
            c.alignment = Alignment(vertical="top", wrap_text=True)
            if p_fill:
                c.fill = p_fill
        if p_font:
            ws_tasks.cell(row=row_idx, column=2).font = p_font
        row_idx += 1

    # Column widths for readability
    col_widths = {1: 5, 2: 12, 3: 10, 4: 22, 5: 18, 6: 40, 7: 50, 8: 60, 9: 15, 10: 18}
    for ci, w in col_widths.items():
        ws_tasks.column_dimensions[get_column_letter(ci)].width = w

    ws_tasks.freeze_panes = "A5"

    # ── Sheet 3b: ТЗ по категориям ────────────────────────────────────
    categories = {}
    for dt in dev_tasks:
        cat = dt.get("category", "Другое")
        categories.setdefault(cat, []).append(dt)

    for cat_name, cat_tasks in categories.items():
        safe_name = re.sub(r"[^\w\s-]", "", cat_name)[:28]
        ws_cat = wb.create_sheet(safe_name)
        ws_cat.sheet_properties.tabColor = "0E7490"

        ws_cat.merge_cells("A1:H1")
        ws_cat.cell(row=1, column=1, value=f"ТЗ: {cat_name} ({len(cat_tasks)} задач)")
        ws_cat.cell(row=1, column=1).font = Font(bold=True, size=12, color="0F4C81")

        cat_headers = ["#", "Приоритет", "Что исправить", "Описание", "Как исправить", "Ответственный", "Срок", "Статус"]
        for ci, h in enumerate(cat_headers, 1):
            ws_cat.cell(row=2, column=ci, value=h)
        _style_header(ws_cat, 2, len(cat_headers))

        ri = 3
        for i, dt in enumerate(cat_tasks, 1):
            priority = str(dt.get("priority", "P3"))
            title = str(dt.get("title", ""))
            how_to_fix = str(dt.get("fix", ""))
            if not how_to_fix:
                for pattern, rec in fix_map.items():
                    if pattern.lower() in title.lower():
                        how_to_fix = rec
                        break
            ws_cat.cell(row=ri, column=1, value=i)
            ws_cat.cell(row=ri, column=2, value=priority)
            ws_cat.cell(row=ri, column=3, value=title)
            ws_cat.cell(row=ri, column=4, value=str(dt.get("description", "")))
            ws_cat.cell(row=ri, column=5, value=how_to_fix)
            ws_cat.cell(row=ri, column=6, value=str(dt.get("owner", "")))
            ws_cat.cell(row=ri, column=7, value=deadline_map.get(priority, ""))
            ws_cat.cell(row=ri, column=8, value="Открыта")
            p_fill = priority_fills.get(priority)
            for ci in range(1, len(cat_headers) + 1):
                c = ws_cat.cell(row=ri, column=ci)
                c.border = thin_border
                c.alignment = Alignment(vertical="top", wrap_text=True)
                if p_fill:
                    c.fill = p_fill
            ri += 1
        cat_widths = {1: 5, 2: 12, 3: 40, 4: 45, 5: 60, 6: 15, 7: 18, 8: 10}
        for ci, w in cat_widths.items():
            ws_cat.column_dimensions[get_column_letter(ci)].width = w
        ws_cat.freeze_panes = "A3"

    # ── Sheet 4: Errors (only if present) ──────────────────────────────
    errors = results.get("errors", {}) or {}
    if errors:
        ws_errors = wb.create_sheet("Errors")
        ws_errors.sheet_properties.tabColor = "B91C1C"
        err_headers = ["Tool", "Error"]
        for ci, h in enumerate(err_headers, 1):
            ws_errors.cell(row=1, column=ci, value=h)
        _style_header(ws_errors, 1, 2)
        row_idx = 2
        for tool_key, err_info in errors.items():
            err_msg = err_info.get("error", str(err_info)) if isinstance(err_info, dict) else str(err_info)
            ws_errors.cell(row=row_idx, column=1, value=str(tool_key))
            ws_errors.cell(row=row_idx, column=2, value=err_msg)
            row_idx += 1
        if row_idx > 2:
            _style_data_rows(ws_errors, 2, row_idx - 1, 2)
        _auto_width(ws_errors, 2)

    filepath = os.path.join(reports_dir, f"unified_audit_{task_id}.xlsx")
    wb.save(filepath)
    return filepath


def _build_unified_docx(task_id: str, url: str, results: dict) -> str:
    """Build a DOCX report for the Unified Full SEO Audit and return the filepath."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml.ns import qn
    from app.config import settings

    reports_dir = settings.REPORTS_DIR
    os.makedirs(reports_dir, exist_ok=True)

    doc = Document()
    generated_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")

    # ── Configure document typography ──────────────────────────────────
    for section in doc.sections:
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(10.5)
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.paragraph_format.line_spacing = 1.12

    for style_name, size, color in (
        ("Title", 28, RGBColor(15, 76, 129)),
        ("Heading 1", 16, RGBColor(15, 76, 129)),
        ("Heading 2", 13, RGBColor(14, 116, 144)),
        ("Heading 3", 11, RGBColor(30, 41, 59)),
    ):
        if style_name not in doc.styles:
            continue
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10 if style_name != "Heading 3" else 8)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.keep_with_next = True

    doc.core_properties.author = "SEO Tools Platform"
    doc.core_properties.title = "Full SEO Audit Report"

    # ── Cover page ─────────────────────────────────────────────────────
    doc.add_paragraph()
    title_para = doc.add_heading("Full SEO Audit Report", 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_para.runs:
        run.font.color.rgb = RGBColor(15, 76, 129)
        run.font.size = Pt(28)

    subtitle_para = doc.add_paragraph("Unified Full SEO Audit")
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if subtitle_para.runs:
        subtitle_para.runs[0].font.size = Pt(14)
        subtitle_para.runs[0].font.color.rgb = RGBColor(107, 114, 128)

    if url:
        url_para = doc.add_paragraph(url)
        url_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if url_para.runs:
            url_para.runs[0].font.color.rgb = RGBColor(14, 116, 144)

    grade_text = f"Overall Grade: {results.get('overall_grade', 'N/A')}"
    grade_para = doc.add_paragraph(grade_text)
    grade_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if grade_para.runs:
        grade_para.runs[0].font.size = Pt(18)
        grade_para.runs[0].font.bold = True
        grade_para.runs[0].font.color.rgb = RGBColor(15, 76, 129)

    ts_para = doc.add_paragraph(generated_at)
    ts_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if ts_para.runs:
        ts_para.runs[0].font.color.rgb = RGBColor(107, 114, 128)

    logo_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "static", "icon.png")
    if os.path.exists(logo_path):
        try:
            logo_para = doc.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = logo_para.add_run()
            run.add_picture(logo_path, width=Inches(1))
        except Exception:
            pass

    doc.add_page_break()

    # ── Header / Footer ────────────────────────────────────────────────
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.text = "SEO Tools Platform | Full SEO Audit Report"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if header_para.runs:
        header_para.runs[0].font.size = Pt(8)
        header_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.text = "SEO Tools Platform"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if footer_para.runs:
        footer_para.runs[0].font.size = Pt(8)
        footer_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    # ── Helper: branded table ──────────────────────────────────────────
    def _add_table(headers, rows):
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            cell = hdr_cells[i]
            cell.text = str(h)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            shading = cell._element.get_or_add_tcPr()
            shd_el = shading.makeelement(qn("w:shd"), {
                qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): "0F4C81",
            })
            shading.append(shd_el)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(255, 255, 255)
        for row_index, row_data in enumerate(rows):
            row_cells = table.add_row().cells
            for i, value in enumerate(row_data):
                row_cells[i].text = str(value)
                row_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                if row_index % 2 == 0:
                    tc_pr = row_cells[i]._element.get_or_add_tcPr()
                    shd_el = tc_pr.makeelement(qn("w:shd"), {
                        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): "F8FBFF",
                    })
                    tc_pr.append(shd_el)
        return table

    # ── Section 1: Overview ────────────────────────────────────────────
    doc.add_heading("Overview", level=1)
    overview_rows = [
        ("URL", url),
        ("Overall Score", str(results.get("overall_score", "N/A"))),
        ("Grade", str(results.get("overall_grade", "N/A"))),
        ("Tools Run", str(results.get("tools_run", "N/A"))),
        ("Tools Failed", str(results.get("tools_failed", "N/A"))),
        ("Duration", f"{results.get('duration_ms', 'N/A')}ms"),
        ("Date", generated_at),
    ]
    _add_table(["Metric", "Value"], overview_rows)
    doc.add_paragraph()

    # ── Section 2: Scores by Tool ──────────────────────────────────────
    doc.add_heading("Scores by Tool", level=1)
    scores = results.get("scores", {})
    score_rows = []
    for key, label in _SCORE_LABELS.items():
        val = scores.get(key)
        score_rows.append((label, val if val is not None else "N/A", _score_status(val)))
    _add_table(["Tool", "Score", "Status"], score_rows)
    doc.add_paragraph()

    # ── Section 3: Developer Tasks / ТЗ ────────────────────────────────
    doc.add_heading("Developer Tasks / ТЗ", level=1)
    dev_tasks = results.get("dev_tasks", []) or []
    if dev_tasks:
        priority_colors = {
            "P0": RGBColor(185, 28, 28),
            "P1": RGBColor(194, 65, 12),
            "P2": RGBColor(161, 98, 7),
            "P3": RGBColor(21, 128, 61),
        }
        task_table = doc.add_table(rows=1, cols=8)
        task_table.style = "Table Grid"
        task_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        task_table.autofit = True
        t_headers = ["#", "Priority", "Category", "Source Tool", "Title", "Description", "Owner", "Fix"]
        hdr_cells = task_table.rows[0].cells
        for i, h in enumerate(t_headers):
            cell = hdr_cells[i]
            cell.text = h
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            shading = cell._element.get_or_add_tcPr()
            shd_el = shading.makeelement(qn("w:shd"), {
                qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): "0F4C81",
            })
            shading.append(shd_el)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(255, 255, 255)
        for idx, dt in enumerate(dev_tasks, 1):
            priority = str(dt.get("priority", ""))
            row_cells = task_table.add_row().cells
            values = [
                str(idx), priority, str(dt.get("category", "")),
                str(dt.get("source_tool", "")), str(dt.get("title", "")),
                str(dt.get("description", "")), str(dt.get("owner", "")),
                str(dt.get("fix", "")),
            ]
            for i, val in enumerate(values):
                row_cells[i].text = val
                row_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            # Color-code priority cell text
            p_color = priority_colors.get(priority)
            if p_color:
                for p in row_cells[1].paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = p_color
                        r.font.bold = True
    else:
        doc.add_paragraph("No developer tasks generated.")
    doc.add_paragraph()

    # ── Section 4: Errors ──────────────────────────────────────────────
    errors = results.get("errors", {}) or {}
    if errors:
        doc.add_heading("Errors", level=1)
        error_rows = []
        for tool_key, err_info in errors.items():
            err_msg = err_info.get("error", str(err_info)) if isinstance(err_info, dict) else str(err_info)
            error_rows.append((str(tool_key), err_msg))
        _add_table(["Tool", "Error"], error_rows)

    filepath = os.path.join(reports_dir, f"unified_audit_{task_id}.docx")
    doc.save(filepath)
    return filepath


# ─── unified audit GET ─────────────────────────────────────────────────


@router.get("/tasks/unified-audit/{task_id}/export/xlsx")
async def get_unified_audit_xlsx(task_id: str):
    """GET export: Unified Full SEO Audit → XLSX."""
    try:
        task, task_result, url, err = _get_task_or_error(task_id, "unified_audit")
        if err:
            return err
        # task_result is the full run_unified_audit() output; use it directly
        results = task_result if task_result.get("overall_score") is not None else task_result.get("results", task_result) or {}
        filepath = _build_unified_xlsx(task_id, url, results)
        if not filepath or not os.path.exists(filepath):
            return _error_response("Failed to generate report", status_code=500)
        append_task_artifact(task_id, filepath, kind="export")
        return _file_response(filepath, _XLSX, _export_filename("unified_audit", url, "xlsx"))
    except Exception as e:
        return _error_response(str(e), status_code=500)


@router.get("/tasks/unified-audit/{task_id}/export/docx")
async def get_unified_audit_docx(task_id: str):
    """GET export: Unified Full SEO Audit → DOCX."""
    try:
        task, task_result, url, err = _get_task_or_error(task_id, "unified_audit")
        if err:
            return err
        results = task_result if task_result.get("overall_score") is not None else task_result.get("results", task_result) or {}
        filepath = _build_unified_docx(task_id, url, results)
        if not filepath or not os.path.exists(filepath):
            return _error_response("Failed to generate report", status_code=500)
        append_task_artifact(task_id, filepath, kind="export")
        return _file_response(filepath, _DOCX, _export_filename("unified_audit", url, "docx"))
    except Exception as e:
        return _error_response(str(e), status_code=500)


