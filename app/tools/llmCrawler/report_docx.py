from __future__ import annotations

import datetime as dt
from io import BytesIO
from typing import Any, Dict


def _safe(result: Dict[str, Any], path: list[str], default=""):
    cur = result
    for p in path:
        if isinstance(cur, dict):
            cur = cur.get(p)
        else:
            return default
    return cur if cur is not None else default


def _num(value: Any, default: float = 0.0) -> float:
    try:
        return float(value)
    except Exception:
        return float(default)


def _pct(value: Any, default: str = "-") -> str:
    try:
        return f"{round(float(value), 2)}%"
    except Exception:
        return default


def _yesno(value: Any) -> str:
    return "✅ Yes" if bool(value) else "❌ No"


def _words(text: Any, limit: int = 25) -> str:
    raw = str(text or "").strip()
    if not raw:
        return ""
    parts = raw.split()
    if len(parts) <= limit:
        return raw
    return " ".join(parts[:limit]) + "..."


def _progress_bar(value: Any, width: int = 18) -> str:
    try:
        v = max(0.0, min(100.0, float(value)))
    except Exception:
        return "—"
    filled = int(round((v / 100.0) * width))
    return f"{'█' * filled}{'░' * (width - filled)} {round(v, 1)}%"


def _score_badge(value: Any) -> str:
    try:
        v = float(value)
    except Exception:
        return "⚪ Not evaluated"
    if v >= 80:
        return "✅ Excellent"
    if v >= 50:
        return "⚠ Needs work"
    return "❌ Critical"


def build_docx_v2(job: Dict[str, Any], job_id: str, wow_enabled: bool = True) -> BytesIO:
    try:
        from docx import Document  # type: ignore
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # type: ignore
        from docx.oxml import parse_xml  # type: ignore
        from docx.oxml.ns import nsdecls  # type: ignore
        from docx.shared import Inches, Pt, RGBColor  # type: ignore
    except Exception as exc:
        raise RuntimeError(f"python-docx unavailable: {exc}")

    result = job.get("result") or {}
    doc = Document()
    normal = doc.styles["Normal"].font
    normal.name = "Calibri"
    normal.size = Pt(11)

    for section in doc.sections:
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)

    def add_heading(title: str, level: int = 1, color: str = "0F4C81") -> None:
        p = doc.add_heading("", level=level)
        run = p.add_run(title)
        run.font.color.rgb = RGBColor.from_string(color)

    def add_subtitle(text: str) -> None:
        p = doc.add_paragraph(text)
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.color.rgb = RGBColor.from_string("475569")

    def add_kv(label: str, value: Any) -> None:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{label}: ")
        r1.bold = True
        r1.font.color.rgb = RGBColor.from_string("0F172A")
        r2 = p.add_run(str(value))
        r2.font.color.rgb = RGBColor.from_string("111827")

    def shade_cell(cell: Any, fill_hex: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_pr.append(parse_xml(rf'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'))

    def add_callout(title: str, lines: list[str], fill_hex: str = "EEF6FF", accent_hex: str = "0F4C81") -> None:
        table = doc.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        cell = table.cell(0, 0)
        shade_cell(cell, fill_hex)
        p = cell.paragraphs[0]
        h = p.add_run(title)
        h.bold = True
        h.font.color.rgb = RGBColor.from_string(accent_hex)
        for line in lines:
            lp = cell.add_paragraph(str(line))
            if lp.runs:
                lp.runs[0].font.size = Pt(10)

    def add_kpi_cards(items: list[dict[str, Any]]) -> None:
        if not items:
            return
        cols = 2
        rows = (len(items) + cols - 1) // cols
        table = doc.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"
        idx = 0
        for r in range(rows):
            for c in range(cols):
                cell = table.cell(r, c)
                shade_cell(cell, "F8FBFF")
                if idx >= len(items):
                    cell.text = ""
                    continue
                item = items[idx]
                idx += 1
                p = cell.paragraphs[0]
                t = p.add_run(str(item.get("title", "")))
                t.bold = True
                t.font.color.rgb = RGBColor.from_string("475569")
                v = cell.add_paragraph(str(item.get("value", "-")))
                if v.runs:
                    v.runs[0].bold = True
                    v.runs[0].font.size = Pt(15)
                    v.runs[0].font.color.rgb = RGBColor.from_string("0F172A")
                if item.get("note"):
                    n = cell.add_paragraph(str(item.get("note")))
                    if n.runs:
                        n.runs[0].font.size = Pt(9)
                        n.runs[0].font.color.rgb = RGBColor.from_string("64748B")

    def add_table(headers, rows):
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = str(h)
            if hdr_cells[i].paragraphs and hdr_cells[i].paragraphs[0].runs:
                hdr_cells[i].paragraphs[0].runs[0].bold = True
        for row in rows:
            cells = table.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)
        return table

    final_url = _safe(result, ["final_url"], _safe(result, ["requested_url"], "-"))
    score_total = _num(_safe(result, ["score", "total"], 0))
    projected = _num(_safe(result, ["projected_score_after_fixes"], score_total), score_total)
    citation = _num(_safe(result, ["citation_probability"], 0))
    eeat = _num(_safe(result, ["eeat_score", "score"], 0))
    trust = _num(_safe(result, ["trust_signal_score"], 0))
    ai_understanding = result.get("ai_understanding") or {}
    citation_breakdown = result.get("citation_breakdown") or {}
    discoverability = result.get("discoverability") or {}
    preview = result.get("ai_answer_preview") or {}
    llm = result.get("llm") or {}
    diff = result.get("diff") or {}
    nojs = result.get("nojs") or {}
    nojs_content = nojs.get("content") or {}
    rendered = result.get("rendered") or {}
    render_status = result.get("render_status") or {}
    rendered_content = (rendered.get("content") or {}) if isinstance(rendered, dict) else {}
    segmentation = (result.get("segmentation") or nojs.get("segmentation") or {})
    noise_breakdown = (result.get("noise_breakdown") or segmentation.get("noise_breakdown") or nojs_content.get("noise_breakdown") or {})
    main_confidence = (result.get("main_content_confidence") or segmentation.get("main_content_confidence") or nojs_content.get("main_content_confidence") or {})
    chunk_dedupe = (result.get("chunk_dedupe") or nojs_content.get("chunk_dedupe") or {})
    page_type = str(result.get("page_type") or nojs.get("page_type") or "-")
    schema = nojs.get("schema") or {}
    resources = nojs.get("resources") or {}
    recommendations = result.get("recommendations") or []
    top_issues = (result.get("score") or {}).get("top_issues") or []
    bot_matrix = result.get("bot_matrix") or []
    diagnostics = (rendered.get("render_debug") or {}) if isinstance(rendered, dict) else {}
    js_dependency = result.get("js_dependency") or {}
    entity_graph = result.get("entity_graph") or {}
    metrics_bytes = result.get("metrics_bytes") or {}
    snippet_library = result.get("snippet_library") or {}
    projected_wf = result.get("projected_score_waterfall") or {}
    cloaking = result.get("cloaking") or {}
    ai_blocks = result.get("ai_blocks") or {}
    critical_blocks = result.get("critical_blocks") or []
    ai_directives = result.get("ai_directives") or {}
    improvement_library = result.get("improvement_library") or {}
    detection_issues = result.get("detection_issues") or []
    content_loss = _num(result.get("content_loss_percent"), 0)
    main_text_len = int(_num(nojs_content.get("main_text_length"), 0))
    rendered_text_len = int(_num(rendered_content.get("main_text_length"), main_text_len))

    # PAGE 1: COVER
    add_heading("AI Crawlability & LLM Visibility Report", 0, color="0B3B63")
    add_subtitle("Enterprise LLM Crawler Simulation Report")
    doc.add_paragraph("")
    add_kv("URL analyzed", final_url)
    add_kv("Generated", f"{dt.datetime.now(dt.timezone.utc).isoformat()}")
    add_kv("Job ID", job_id)
    doc.add_paragraph("")
    score_p = doc.add_paragraph()
    score_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    score_run = score_p.add_run(f"{round(score_total, 2)} / 100")
    score_run.bold = True
    score_run.font.size = Pt(34)
    score_run.font.color.rgb = RGBColor.from_string("0F4C81")
    lbl = doc.add_paragraph("AI Visibility Score")
    lbl.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if lbl.runs:
        lbl.runs[0].font.color.rgb = RGBColor.from_string("475569")
    kpi_items = [
        {"title": "Current score", "value": f"{round(score_total, 2)} / 100", "note": _score_badge(score_total)},
        {"title": "Projected score", "value": f"{round(projected, 2)} / 100", "note": _score_badge(projected)},
        {"title": "Citation likelihood", "value": _pct(citation), "note": _progress_bar(citation)},
        {"title": "Trust completeness", "value": _pct(trust), "note": _progress_bar(trust)},
    ]
    add_kpi_cards(kpi_items)
    doc.add_page_break()

    # PAGE 2: EXECUTIVE SUMMARY
    add_heading("⚡ Executive Summary", 1, color="0B3B63")
    add_kv("Overall score", f"{round(score_total, 2)} / 100")
    add_kv("Projected score after fixes", f"{round(projected, 2)} / 100")
    add_kv("Citation likelihood", _pct(citation))
    if (result.get("eeat_score") or {}).get("status") == "not_evaluated":
        add_kv("EEAT score", f"— Not evaluated ({(result.get('eeat_score') or {}).get('reason', 'unknown')})")
    else:
        add_kv("EEAT score", _pct(eeat))
    add_kv("Trust completeness", _pct(trust))
    add_callout(
        "Snapshot",
        [
            f"Quality status: {_score_badge(score_total)}",
            f"Projected lift: +{round(max(0.0, projected - score_total), 2)} points",
            f"Content loss: {_pct(content_loss)}",
        ],
        fill_hex="EEF8FF",
        accent_hex="0E7490",
    )
    if projected_wf and wow_enabled:
        add_heading("Projected Impact Waterfall", 2, color="0E7490")
        rows = [["Baseline", projected_wf.get("baseline", score_total)]]
        for step in projected_wf.get("steps") or []:
            rows.append([f"{step.get('label', 'Step')} (+{step.get('delta', 0)})", step.get("value", "-")])
        rows.append(["Projected", projected_wf.get("target", projected)])
        add_table(["Step", "Score"], rows)
    doc.add_paragraph("")
    add_heading("⚠ Top Findings", 2, color="9A3412")
    if top_issues:
        for issue in top_issues[:6]:
            doc.add_paragraph(str(issue), style="List Bullet")
    else:
        doc.add_paragraph("No critical findings detected.")
    doc.add_page_break()

    # PAGE 3: WHAT AI UNDERSTANDS
    add_heading("🧠 AI Understanding Summary", 1, color="0B3B63")
    add_kv("Topic detected", _safe(ai_understanding, ["topic"], llm.get("summary", "Not detected")))
    add_kv("Page type", page_type)
    add_kv("AI understanding score", _pct(_safe(ai_understanding, ["score"], 0)))
    add_kv("Topic confidence", _pct(_safe(ai_understanding, ["topic_confidence"], 0)))
    add_kv("Topic fallback used", _yesno(_safe(ai_understanding, ["topic_fallback_used"], False)))
    add_kv("Primary intent", _safe(ai_understanding, ["intent"], "informational"))
    if _safe(ai_understanding, ["content_clarity_status"], "") == "not_evaluated":
        add_kv("Content clarity", f"— Not evaluated ({_safe(ai_understanding, ['content_clarity_reason'], 'unknown')})")
    else:
        add_kv("Content clarity", _pct(_safe(ai_understanding, ["content_clarity"], 0)))
    add_kv("Discoverability score", _pct(discoverability.get("discoverability_score", 0)))
    add_kv("Click depth estimate", discoverability.get("click_depth_estimate", "-"))
    add_callout(
        "AI Readiness Indicators",
        [
            f"Understanding: {_progress_bar(_safe(ai_understanding, ['score'], 0))}",
            f"Content clarity: {_progress_bar(_safe(ai_understanding, ['content_clarity'], 0)) if _safe(ai_understanding, ['content_clarity'], None) is not None else '—'}",
            f"Discoverability: {_progress_bar(discoverability.get('discoverability_score', 0))}",
        ],
        fill_hex="F0F9FF",
        accent_hex="0369A1",
    )
    doc.add_paragraph("")
    add_heading("Detected Entities", 2, color="0E7490")
    entities = _safe(ai_understanding, ["entities"], [])
    if entities:
        for entity in entities[:20]:
            doc.add_paragraph(str(entity), style="List Bullet")
    else:
        doc.add_paragraph("No stable entities detected.")
    doc.add_page_break()

    # PAGE 4: CONTENT LOSS ANALYSIS
    add_heading("📉 Content Loss Analysis", 1, color="0B3B63")
    add_kv("Total HTML text length", rendered_text_len)
    add_kv("Extracted main text length", main_text_len)
    add_kv("Lost content percent", _pct(content_loss))
    add_kv("HTML bytes", metrics_bytes.get("html_bytes", rendered_text_len))
    add_kv("Text bytes", metrics_bytes.get("text_bytes", main_text_len))
    add_kv("Text/HTML ratio", _pct(_num(metrics_bytes.get("text_html_ratio"), 0) * 100))
    add_kv("Main content ratio", _pct(_num(nojs_content.get("main_content_ratio"), 0) * 100))
    add_kv("Boilerplate ratio", _pct(_num(nojs_content.get("boilerplate_ratio"), 0) * 100))
    add_kv("Chunks count", len(nojs_content.get("chunks") or []))
    add_kv("Unique chunks", int(_num(chunk_dedupe.get("chunks_unique"), len(nojs_content.get("chunks") or []))))
    add_kv("Removed duplicate chunks", int(_num(chunk_dedupe.get("removed_duplicates"), 0)))
    add_kv("Dedupe ratio", _pct(chunk_dedupe.get("dedupe_ratio"), "0%"))
    add_callout(
        "Content Ratios",
        [
            f"Text/HTML: {_progress_bar(_num(metrics_bytes.get('text_html_ratio'), 0) * 100)}",
            f"Main content: {_progress_bar(_num(nojs_content.get('main_content_ratio'), 0) * 100)}",
            f"Boilerplate: {_progress_bar(_num(nojs_content.get('boilerplate_ratio'), 0) * 100)}",
        ],
        fill_hex="F8FAFC",
        accent_hex="475569",
    )
    if wow_enabled and (nojs_content.get("chunks") or []):
        add_heading("Top Chunk Excerpts", 2, color="0E7490")
        for ch in (nojs_content.get("chunks") or [])[:3]:
            doc.add_paragraph(f"Chunk {ch.get('idx', '-')}: {_words(ch.get('text', ''), 35)}", style="List Bullet")
    add_heading("Extracted Text Preview", 2, color="0E7490")
    doc.add_paragraph(str(nojs_content.get("main_text_preview", ""))[:1200] or "No preview.")
    doc.add_page_break()

    # PAGE 5: NOISE & SEGMENTATION
    add_heading("🧹 Noise & Segmentation", 1, color="0B3B63")
    add_kv("Main content percent", _pct(noise_breakdown.get("main_pct"), "-"))
    add_kv("Ads percent", _pct(noise_breakdown.get("ads_pct"), "-"))
    add_kv("Live scores percent", _pct(noise_breakdown.get("live_pct"), "-"))
    add_kv("Navigation/footer percent", _pct(noise_breakdown.get("nav_pct"), "-"))
    add_kv("Utility percent", _pct(noise_breakdown.get("utility_pct"), "-"))
    add_kv("Main-content confidence", str(main_confidence.get("level") or "unknown"))
    reasons = main_confidence.get("reasons") or []
    if reasons:
        add_heading("Confidence reasons", 2, color="0E7490")
        for reason in reasons[:6]:
            doc.add_paragraph(str(reason), style="List Bullet")
    doc.add_page_break()

    # PAGE 6: CITATION READINESS BREAKDOWN
    add_heading("📚 Citation Readiness Breakdown", 1, color="0B3B63")
    add_table(
        ["Factor", "Score"],
        [
            ["Schema", _pct(citation_breakdown.get("schema", 0))],
            ["Author", _pct(citation_breakdown.get("author", 0))],
            ["Content clarity", _pct(citation_breakdown.get("content_clarity", 0))],
            ["Bot accessibility", _pct(citation_breakdown.get("bot_accessibility", 0))],
            ["Structure", _pct(citation_breakdown.get("structure", 0))],
        ],
    )
    add_heading("Projected Improvement", 2, color="0E7490")
    add_kv("Current score", f"{round(score_total, 2)} / 100")
    add_kv("Projected score", f"{round(projected, 2)} / 100")
    doc.add_page_break()

    # PAGE 7: BOT ACCESS
    add_heading("🤖 Bot Access Matrix", 1, color="0B3B63")
    matrix_rows = []
    for m in bot_matrix:
        matrix_rows.append(
            [
                m.get("profile", "-"),
                "✅ Allowed" if m.get("allowed") else "❌ Blocked",
                m.get("reason", "-"),
            ]
        )
    add_table(["Bot", "Access", "Reason"], matrix_rows or [["-", "-", "-"]])
    add_heading("Directives", 2, color="0E7490")
    add_kv("meta robots", _safe(result, ["policies", "meta", "meta_robots"], "-"))
    add_kv("x-robots-tag", _safe(result, ["policies", "meta", "x_robots_tag"], "-"))
    doc.add_page_break()

    # PAGE 8: LLM SIMULATION
    add_heading("💬 LLM Simulation", 1, color="0B3B63")
    add_kv("Summary", llm.get("summary", "Not available"))
    key_facts = llm.get("key_facts") or []
    if key_facts:
        add_heading("Key Facts", 2, color="0E7490")
        for fact in key_facts[:10]:
            doc.add_paragraph(str(fact), style="List Bullet")
    add_heading("Citation Spans", 2, color="0E7490")
    spans = llm.get("citation_spans") or []
    if spans:
        for span in spans[:12]:
            quoted = _words((span or {}).get("text", ""), 25)
            if quoted:
                doc.add_paragraph(f"\"{quoted}\"", style="List Bullet")
    else:
        doc.add_paragraph("No citation spans.")
    add_heading("AI Answer Preview", 2, color="0E7490")
    add_kv("Question", preview.get("question", "What is this page about?"))
    add_kv("Answer", preview.get("answer", "Not enough content"))
    add_kv("Confidence", _pct(preview.get("confidence"), "-"))
    if preview.get("warning"):
        add_kv("Warning", preview.get("warning"))
    bullets = preview.get("bullets") or []
    fix_steps = preview.get("fix_steps") or []
    if bullets:
        add_heading("Key bullets", 2, color="0E7490")
        for b in bullets[:5]:
            doc.add_paragraph(str(b), style="List Bullet")
    if fix_steps:
        add_heading("How to fix preview reliability", 2, color="9A3412")
        for step in fix_steps[:3]:
            doc.add_paragraph(str(step), style="List Bullet")
    doc.add_page_break()

    # PAGE 9: STRUCTURED DATA
    add_heading("🧩 Structured Data & Site Structure", 1, color="0B3B63")
    add_kv("Schema coverage", _pct(schema.get("coverage_score"), "-"))
    add_kv("Organization schema", _yesno("Organization" in (schema.get("jsonld_types") or [])))
    add_kv("Article schema", _yesno("Article" in (schema.get("jsonld_types") or [])))
    add_kv("Product schema", _yesno("Product" in (schema.get("jsonld_types") or [])))
    add_kv("Breadcrumb schema", _yesno("BreadcrumbList" in (schema.get("jsonld_types") or [])))
    add_kv("Microdata types", ", ".join((schema.get("microdata_types") or [])[:8]) or "-")
    add_kv("RDFa types", ", ".join((schema.get("rdfa_types") or [])[:8]) or "-")
    doc.add_page_break()

    # PAGE 10: TECHNICAL DIAGNOSTICS
    add_heading("🛠 Technical Diagnostics", 1, color="0B3B63")
    add_kv("Render snapshot status", render_status.get("status", "not_executed"))
    add_kv("Render snapshot reason", render_status.get("reason", "render_not_executed"))
    if str(js_dependency.get("status") or "") == "executed":
        add_kv("JS dependency score", _pct(js_dependency.get("score"), "-"))
        add_kv("Failed resources", js_dependency.get("failures", 0))
        add_kv("Blocked scripts/styles", js_dependency.get("blocked", 0))
    else:
        add_kv("JS dependency score", "Not executed")
        add_kv("JS dependency reason", js_dependency.get("reason", "render_not_executed"))
    add_kv("Cloaking status", cloaking.get("status", "not_executed"))
    h1_cons = diff.get("h1Consistency") or {}
    if bool(h1_cons.get("h1_appears_only_after_js")):
        add_kv("H1 consistency", "H1 appears only after JS")
    if cloaking.get("status") == "executed":
        add_kv("Cloaking risk", cloaking.get("risk", "unknown"))
        add_kv("Similarity browser vs gptbot", _pct(_safe(cloaking, ["similarity_scores", "browser_vs_gptbot"], "-")))
        add_kv("Similarity browser vs googlebot", _pct(_safe(cloaking, ["similarity_scores", "browser_vs_googlebot"], "-")))
    else:
        add_kv("Cloaking reason", cloaking.get("reason", "not_run"))
    if cloaking.get("status") == "executed":
        add_callout(
            "Cloaking quick view",
            [
                f"Risk: {cloaking.get('risk', 'unknown')}",
                f"Browser vs GPTBot: {_pct(_safe(cloaking, ['similarity_scores', 'browser_vs_gptbot'], '-'))}",
                f"Browser vs Googlebot: {_pct(_safe(cloaking, ['similarity_scores', 'browser_vs_googlebot'], '-'))}",
            ],
            fill_hex="F0FDF4",
            accent_hex="166534",
        )
    add_kv("Console errors", len(diagnostics.get("console_errors") or []))
    add_kv("Failed requests", len(diagnostics.get("failed_requests") or []))
    if diagnostics.get("console_errors"):
        add_heading("Top Console Errors", 2, color="9A3412")
        for row in (diagnostics.get("console_errors") or [])[:8]:
            doc.add_paragraph(str(row)[:220], style="List Bullet")
    doc.add_page_break()

    # PAGE 11: ACCESS BARRIERS
    add_heading("🔒 Access Barriers", 1, color="0B3B63")
    add_kv("Cookie wall", _yesno(resources.get("cookie_wall")))
    add_kv("Paywall", _yesno(resources.get("paywall")))
    add_kv("Login wall", _yesno(resources.get("login_wall")))
    add_kv("Strict CSP", _yesno(resources.get("csp_strict")))
    add_kv("Mixed content count", resources.get("mixed_content_count", 0))
    add_page = doc.add_paragraph("")
    if add_page.runs:
        add_page.runs[0].font.size = Pt(1)
    doc.add_page_break()

    # PAGE 12: ENTITY + RECOMMENDATIONS
    add_heading("🧭 Entity Graph & Recommendations", 1, color="0B3B63")
    add_kv("Organizations", ", ".join((entity_graph.get("organizations") or [])[:12]) or "-")
    add_kv("Persons", ", ".join((entity_graph.get("persons") or [])[:12]) or "-")
    add_kv("Products", ", ".join((entity_graph.get("products") or [])[:12]) or "-")
    add_kv("Locations", ", ".join((entity_graph.get("locations") or [])[:12]) or "-")
    doc.add_paragraph("")
    add_heading("🔧 Prioritized Fix Backlog", 2, color="9A3412")
    if recommendations:
        for rec in recommendations[:20]:
            pri = str(rec.get("priority", "P2")).upper()
            area = rec.get("area", "-")
            title = rec.get("title", "-")
            expected = rec.get("expected_lift") or ("+12 score" if pri == "P0" else "+7 score" if pri == "P1" else "+3 score")
            citation_effect = rec.get("expected_citation_effect") or "-"
            doc.add_paragraph(
                f"{pri} | {area} | {title} | expected impact {expected} | citation effect {citation_effect}",
                style="List Bullet",
            )
            evidence = rec.get("evidence") or []
            if evidence:
                for ev in evidence[:3]:
                    doc.add_paragraph(f"Evidence: {ev}", style="List Bullet")
    else:
        doc.add_paragraph("No recommendations.")

    if recommendations:
        p0 = len([r for r in recommendations if str(r.get("priority", "")).upper() == "P0"])
        p1 = len([r for r in recommendations if str(r.get("priority", "")).upper() == "P1"])
        p2 = len([r for r in recommendations if str(r.get("priority", "")).upper() == "P2"])
        add_callout(
            "Recommendation mix",
            [
                f"P0 critical: {p0}",
                f"P1 recommended: {p1}",
                f"P2 optional: {p2}",
            ],
            fill_hex="FFF7ED",
            accent_hex="9A3412",
        )

    add_heading("AI Detection Coverage", 2, color="0E7490")
    add_kv("Coverage percent", _pct(ai_blocks.get("coverage_percent"), "-"))
    detected_blocks = ai_blocks.get("detected") or []
    if detected_blocks:
        for b in detected_blocks[:10]:
            doc.add_paragraph(f"{b.get('label', b.get('id', '-'))}: {b.get('confidence', '-')}", style="List Bullet")
    missing_critical = ai_blocks.get("missing_critical") or []
    if missing_critical:
        add_heading("Missing critical blocks", 3, color="9A3412")
        for m in missing_critical[:8]:
            doc.add_paragraph(str(m), style="List Bullet")
    if critical_blocks:
        add_heading("Critical block checklist", 2, color="0E7490")
        for item in critical_blocks[:12]:
            status = "✅" if str(item.get("status")) == "present" else "❌"
            line = f"{status} {item.get('label', item.get('id', '-'))}: {item.get('evidence', '-')}"
            doc.add_paragraph(line, style="List Bullet")

    directives_profiles = (ai_directives.get("profiles") or {})
    if directives_profiles:
        add_heading("AI directives audit", 2, color="0E7490")
        rows = [[k, v.get("status", "-"), v.get("reason", "-")] for k, v in directives_profiles.items()]
        add_table(["Bot", "Status", "Reason"], rows)
    if detection_issues:
        add_heading("Detection issues", 2, color="9A3412")
        for issue in detection_issues[:10]:
            doc.add_paragraph(str(issue), style="List Bullet")

    missing_improvements = improvement_library.get("missing") or []
    if missing_improvements:
        add_heading("Improvement Library", 2, color="0E7490")
        for item in missing_improvements[:10]:
            doc.add_paragraph(f"{item.get('title', item.get('id', '-'))}: {item.get('reason', '-')}", style="List Bullet")

    if wow_enabled and snippet_library:
        doc.add_page_break()
        add_heading("Appendix: Copy-Paste Snippets", 1, color="0B3B63")
        for key, snippet in snippet_library.items():
            add_heading(str(key), 2, color="0E7490")
            doc.add_paragraph(str(snippet)[:1800])

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio
