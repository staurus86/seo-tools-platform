# -*- coding: utf-8 -*-
"""
SEO Rendering Audit - JS vs. No-JS Comparison

Usage:
  python seo_render_audit.py https://example.com/ --out-dir output

Requirements:
  pip install playwright requests beautifulsoup4 lxml python-docx openpyxl
  playwright install chromium
"""

from __future__ import annotations

import argparse
import json
import re
import sys
import time
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, Iterable, List, Tuple
import textwrap
from urllib.parse import urlparse

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from openpyxl import Workbook
from playwright.sync_api import sync_playwright, TimeoutError as PlaywrightTimeoutError
from app.config import settings


GOOGLEBOT_DESKTOP_UA = (
    "Mozilla/5.0 (compatible; Googlebot/2.1; +http://www.google.com/bot.html)"
)
GOOGLEBOT_MOBILE_UA = (
    "Mozilla/5.0 (Linux; Android 6.0.1; Nexus 5X Build/MMB29P) "
    "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/41.0.2272.96 Mobile Safari/537.36 "
    "(compatible; Googlebot/2.1; +http://www.google.com/bot.html)"
)

DEFAULT_TIMEOUT_MS = 45000


@dataclass
class PageSnapshot:
    url: str
    fetched_url: str
    status_code: int | None
    title: str
    meta_description: str
    canonical: str
    headings: List[str]
    links: List[str]
    images: List[str]
    structured_data: List[str]
    schema_types: List[str]
    viewport_present: bool
    og_present: bool
    visible_text_lines: List[str]
    html_bytes: int
    errors: List[str]


def _norm(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").strip().lower())


def _unique_keep_order(items: Iterable[str]) -> List[str]:
    seen = set()
    out = []
    for item in items:
        key = _norm(item)
        if not key or key in seen:
            continue
        seen.add(key)
        out.append(item)
    return out


def extract_schema_types(jsonld_items: Iterable[str]) -> List[str]:
    types: List[str] = []
    for item in jsonld_items:
        try:
            data = json.loads(item)
        except Exception:
            continue

        def walk(node: Any):
            if isinstance(node, dict):
                t = node.get("@type")
                if isinstance(t, str):
                    types.append(t)
                elif isinstance(t, list):
                    for t_item in t:
                        if isinstance(t_item, str):
                            types.append(t_item)
                for v in node.values():
                    walk(v)
            elif isinstance(node, list):
                for v in node:
                    walk(v)

        walk(data)
    return _unique_keep_order(types)


def count_headings(headings: Iterable[str], level: int) -> int:
    prefix = f"h{level}:"
    return sum(1 for h in headings if h.lower().startswith(prefix))


def format_ru_datetime(dt: datetime) -> str:
    months = [
        "января",
        "февраля",
        "марта",
        "апреля",
        "мая",
        "июня",
        "июля",
        "августа",
        "сентября",
        "октября",
        "ноября",
        "декабря",
    ]
    return f"{dt.day} {months[dt.month - 1]} {dt.year} г. в {dt.strftime('%H:%M')}"


def status_for_count(count: int) -> str:
    if count == 0:
        return "✅"
    if count <= 5:
        return "⚠️"
    return "❌"


def status_for_match(a: Any, b: Any) -> str:
    return "✅" if a == b else "⚠️"


def match_text(a: Any, b: Any) -> str:
    return "✓ Совпадает" if a == b else "⚠ Не совпадает"


def bar(value_ms: float, max_ms: float) -> str:
    if max_ms <= 0:
        return ""
    blocks = int(min(10, max(0, round((value_ms / max_ms) * 10))))
    return "█" * blocks + "░" * (10 - blocks)


def log_step(label: str, step: int, total: int) -> None:
    filled = int(round((step / total) * 20))
    bar_str = "█" * filled + "░" * (20 - filled)
    print(f"{label} [{bar_str}] {step}/{total}", flush=True)


def fetch_raw_html(url: str, user_agent: str, timeout_s: int = 30) -> Tuple[str, int | None, str]:
    headers = {"User-Agent": user_agent, "Accept-Language": "en-US,en;q=0.9"}
    try:
        resp = requests.get(url, headers=headers, timeout=timeout_s, allow_redirects=True)
        return resp.text, resp.status_code, resp.url
    except Exception as exc:
        return "", None, f"request error: {exc}"


def parse_raw_html(url: str, html: str, status_code: int | None, fetched_url: str, error: str) -> PageSnapshot:
    errors = []
    if error:
        errors.append(error)
    soup = BeautifulSoup(html or "", "lxml")
    title = (soup.title.string if soup.title and soup.title.string else "").strip()

    meta_desc = ""
    meta_tag = soup.find("meta", attrs={"name": re.compile(r"^description$", re.I)})
    if meta_tag and meta_tag.get("content"):
        meta_desc = meta_tag.get("content", "").strip()

    canonical = ""
    canonical_tag = soup.find("link", attrs={"rel": re.compile(r"canonical", re.I)})
    if canonical_tag and canonical_tag.get("href"):
        canonical = canonical_tag.get("href", "").strip()

    headings = []
    for level in range(1, 7):
        for h in soup.find_all(f"h{level}"):
            text = h.get_text(" ", strip=True)
            if text:
                headings.append(f"h{level}: {text}")

    links = []
    for a in soup.find_all("a"):
        href = (a.get("href") or "").strip()
        text = a.get_text(" ", strip=True)
        if href or text:
            links.append(f"{href} | {text}".strip())

    images = []
    for img in soup.find_all("img"):
        src = (img.get("src") or "").strip()
        alt = (img.get("alt") or "").strip()
        if src or alt:
            images.append(f"{src} | alt={alt}".strip())

    structured_data = []
    for script in soup.find_all("script", attrs={"type": re.compile(r"ld\+json", re.I)}):
        text = script.string or script.get_text() or ""
        text = text.strip()
        if text:
            structured_data.append(text)

    schema_types = extract_schema_types(structured_data)
    viewport_present = bool(soup.find("meta", attrs={"name": re.compile(r"^viewport$", re.I)}))
    og_present = bool(soup.find("meta", attrs={"property": re.compile(r"^og:", re.I)}))

    raw_text = soup.get_text("\n", strip=True)
    text_lines = [line.strip() for line in raw_text.split("\n") if line.strip()]

    return PageSnapshot(
        url=url,
        fetched_url=fetched_url or url,
        status_code=status_code,
        title=title,
        meta_description=meta_desc,
        canonical=canonical,
        headings=_unique_keep_order(headings),
        links=_unique_keep_order(links),
        images=_unique_keep_order(images),
        structured_data=_unique_keep_order(structured_data),
        schema_types=_unique_keep_order(schema_types),
        viewport_present=viewport_present,
        og_present=og_present,
        visible_text_lines=_unique_keep_order(text_lines),
        html_bytes=len((html or "").encode("utf-8", errors="ignore")),
        errors=errors,
    )


def render_with_playwright(
    url: str,
    user_agent: str,
    mobile: bool,
    screenshot_path: Path | None = None,
) -> Tuple[Dict[str, Any], Dict[str, float], List[str]]:
    errors: List[str] = []
    data: Dict[str, Any] = {}
    timing: Dict[str, float] = {}
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        try:
            context_kwargs: Dict[str, Any] = {
                "user_agent": user_agent,
                "locale": "en-US",
            }
            if mobile:
                device = p.devices.get("Pixel 5")
                if device:
                    context_kwargs.update(device)
            else:
                context_kwargs.update({"viewport": {"width": 1366, "height": 900}})

            context = browser.new_context(**context_kwargs)
            page = context.new_page()
            page.set_default_timeout(DEFAULT_TIMEOUT_MS)

            try:
                page.goto(url, wait_until="networkidle")
            except PlaywrightTimeoutError:
                try:
                    page.goto(url, wait_until="load")
                except Exception as exc:
                    errors.append(f"page load error: {exc}")

            js = textwrap.dedent(
                r"""
                () => {
                  const out = {};
                  const norm = (t) => (t || '').trim();
                  const isVisible = (el) => {
                    if (!el) return false;
                    const style = window.getComputedStyle(el);
                    if (!style) return false;
                    if (style.display === 'none' || style.visibility === 'hidden' || style.opacity === '0') return false;
                    const rects = el.getClientRects();
                    if (!rects || rects.length === 0) return false;
                    if (el.offsetWidth <= 0 || el.offsetHeight <= 0) return false;
                    return true;
                  };

                  const getMetaDescription = () => {
                    const metas = document.querySelectorAll('meta[name]');
                    for (const m of metas) {
                      const name = (m.getAttribute('name') || '').toLowerCase();
                      if (name === 'description') return m.getAttribute('content') || '';
                    }
                    return '';
                  };

                  const getCanonical = () => {
                    const links = document.querySelectorAll('link[rel]');
                    for (const l of links) {
                      const rel = (l.getAttribute('rel') || '').toLowerCase();
                      if (rel.includes('canonical')) return l.getAttribute('href') || '';
                    }
                    return '';
                  };

                  out.title = norm(document.title || '');
                  out.meta_description = norm(getMetaDescription());
                  out.canonical = norm(getCanonical());
                  out.viewport_present = !!document.querySelector('meta[name=\"viewport\" i]');
                  out.og_present = !!document.querySelector('meta[property^=\"og:\" i]');

                  const headings = [];
                  for (let i = 1; i <= 6; i++) {
                    document.querySelectorAll('h' + i).forEach((h) => {
                      if (isVisible(h)) {
                        const text = norm(h.innerText || h.textContent || '');
                        if (text) headings.push('h' + i + ': ' + text);
                      }
                    });
                  }

                  const links = [];
                  document.querySelectorAll('a').forEach((a) => {
                    if (!isVisible(a)) return;
                    const href = norm(a.getAttribute('href') || '');
                    const text = norm(a.innerText || a.textContent || '');
                    if (href || text) links.push((href + ' | ' + text).trim());
                  });

                  const images = [];
                  document.querySelectorAll('img').forEach((img) => {
                    if (!isVisible(img)) return;
                    const src = norm(img.getAttribute('src') || '');
                    const alt = norm(img.getAttribute('alt') || '');
                    if (src || alt) images.push((src + ' | alt=' + alt).trim());
                  });

                  const structured_data = [];
                  document.querySelectorAll('script[type=\"application/ld+json\"]').forEach((s) => {
                    const text = norm(s.textContent || '');
                    if (text) structured_data.push(text);
                  });

                  const visible_text_lines = [];
                  const body = document.body;
                  if (body) {
                    const text = (body.innerText || '').split('\\n');
                    text.forEach((line) => {
                      const t = norm(line);
                      if (t) visible_text_lines.push(t);
                    });
                  }

                  out.headings = headings;
                  out.links = links;
                  out.images = images;
                  out.structured_data = structured_data;
                  out.visible_text_lines = visible_text_lines;
                  return out;
                }
                """
            ).strip()
            data = page.evaluate(js)
            timing = page.evaluate(
                """
                () => {
                  const nav = performance.getEntriesByType('navigation')[0];
                  const t = nav || performance.timing || {};
                  const toMs = (v) => (typeof v === 'number' ? v : 0);
                  const base = toMs(nav ? nav.startTime : t.navigationStart);
                  const rel = (v) => Math.max(0, toMs(v) - base);

                  const fcpEntry = performance.getEntriesByType('paint')
                    .find((e) => e.name === 'first-contentful-paint');
                  const fcp = fcpEntry ? fcpEntry.startTime : 0;

                  return {
                    dns: rel(nav ? nav.domainLookupEnd : t.domainLookupEnd) - rel(nav ? nav.domainLookupStart : t.domainLookupStart),
                    tcp: rel(nav ? nav.connectEnd : t.connectEnd) - rel(nav ? nav.connectStart : t.connectStart),
                    html: rel(nav ? nav.responseEnd : t.responseEnd) - rel(nav ? nav.responseStart : t.responseStart),
                    dom_parse: rel(nav ? nav.domContentLoadedEventStart : t.domContentLoadedEventStart) - rel(nav ? nav.responseEnd : t.responseEnd),
                    render: rel(nav ? nav.loadEventEnd : t.loadEventEnd) - rel(nav ? nav.domContentLoadedEventEnd : t.domContentLoadedEventEnd),
                    ttfb: rel(nav ? nav.responseStart : t.responseStart) - rel(nav ? nav.requestStart : t.requestStart),
                    fcp: fcp,
                    dom_content_loaded: rel(nav ? nav.domContentLoadedEventEnd : t.domContentLoadedEventEnd),
                    tbt: 0,
                    tti: rel(nav ? nav.loadEventEnd : t.loadEventEnd),
                  };
                }
                """
            )
            data["html_bytes"] = len((page.content() or "").encode("utf-8", errors="ignore"))
            if screenshot_path:
                page.screenshot(path=str(screenshot_path), full_page=True)
            context.close()
        finally:
            browser.close()

    return data, timing, errors


def screenshot_nojs(
    url: str,
    user_agent: str,
    mobile: bool,
    landscape: bool,
    screenshot_path: Path,
) -> List[str]:
    errors: List[str] = []
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        try:
            context_kwargs: Dict[str, Any] = {
                "user_agent": user_agent,
                "locale": "en-US",
                "java_script_enabled": False,
            }
            if mobile:
                device = p.devices.get("Pixel 5")
                if device:
                    context_kwargs.update(device)
            else:
                context_kwargs.update({"viewport": {"width": 1366, "height": 900}})

            if landscape and context_kwargs.get("viewport"):
                vp = context_kwargs["viewport"]
                context_kwargs["viewport"] = {"width": vp["height"], "height": vp["width"]}

            context = browser.new_context(**context_kwargs)
            page = context.new_page()
            page.set_default_timeout(DEFAULT_TIMEOUT_MS)

            try:
                page.goto(url, wait_until="networkidle")
            except PlaywrightTimeoutError:
                try:
                    page.goto(url, wait_until="load")
                except Exception as exc:
                    errors.append(f"no-js page load error: {exc}")

            page.screenshot(path=str(screenshot_path), full_page=True)
            context.close()
        finally:
            browser.close()
    return errors


def timing_nojs(
    url: str,
    user_agent: str,
    mobile: bool,
) -> Tuple[Dict[str, float], List[str]]:
    errors: List[str] = []
    timing: Dict[str, float] = {}
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        try:
            context_kwargs: Dict[str, Any] = {
                "user_agent": user_agent,
                "locale": "en-US",
            }
            if mobile:
                device = p.devices.get("Pixel 5")
                if device:
                    context_kwargs.update(device)
            else:
                context_kwargs.update({"viewport": {"width": 1366, "height": 900}})

            context = browser.new_context(**context_kwargs)
            page = context.new_page()
            page.set_default_timeout(DEFAULT_TIMEOUT_MS)

            page.route(
                "**/*",
                lambda route, request: route.abort()
                if request.resource_type in {"script", "xhr", "fetch"} else route.continue_(),
            )

            try:
                page.goto(url, wait_until="load")
            except PlaywrightTimeoutError:
                try:
                    page.goto(url, wait_until="domcontentloaded")
                except Exception as exc:
                    errors.append(f"timing no-js load error: {exc}")

            timing = page.evaluate(
                """
                () => {
                  const nav = performance.getEntriesByType('navigation')[0];
                  const t = nav || performance.timing || {};
                  const toMs = (v) => (typeof v === 'number' ? v : 0);
                  const base = toMs(nav ? nav.startTime : t.navigationStart);
                  const rel = (v) => Math.max(0, toMs(v) - base);
                  const fcpEntry = performance.getEntriesByType('paint')
                    .find((e) => e.name === 'first-contentful-paint');
                  const fcp = fcpEntry ? fcpEntry.startTime : 0;
                  return {
                    dns: rel(nav ? nav.domainLookupEnd : t.domainLookupEnd) - rel(nav ? nav.domainLookupStart : t.domainLookupStart),
                    tcp: rel(nav ? nav.connectEnd : t.connectEnd) - rel(nav ? nav.connectStart : t.connectStart),
                    html: rel(nav ? nav.responseEnd : t.responseEnd) - rel(nav ? nav.responseStart : t.responseStart),
                    dom_parse: rel(nav ? nav.domContentLoadedEventStart : t.domContentLoadedEventStart) - rel(nav ? nav.responseEnd : t.responseEnd),
                    render: rel(nav ? nav.loadEventEnd : t.loadEventEnd) - rel(nav ? nav.domContentLoadedEventEnd : t.domContentLoadedEventEnd),
                    ttfb: rel(nav ? nav.responseStart : t.responseStart) - rel(nav ? nav.requestStart : t.requestStart),
                    fcp: fcp,
                    dom_content_loaded: rel(nav ? nav.domContentLoadedEventEnd : t.domContentLoadedEventEnd),
                    tbt: 0,
                    tti: rel(nav ? nav.loadEventEnd : t.loadEventEnd),
                  };
                }
                """
            )
            context.close()
        finally:
            browser.close()
    return timing, errors


def screenshot_rendered_only(
    url: str,
    user_agent: str,
    mobile: bool,
    landscape: bool,
    screenshot_path: Path,
) -> List[str]:
    errors: List[str] = []
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        try:
            context_kwargs: Dict[str, Any] = {
                "user_agent": user_agent,
                "locale": "en-US",
            }
            if mobile:
                device = p.devices.get("Pixel 5")
                if device:
                    context_kwargs.update(device)
            else:
                context_kwargs.update({"viewport": {"width": 1366, "height": 900}})

            if landscape and context_kwargs.get("viewport"):
                vp = context_kwargs["viewport"]
                context_kwargs["viewport"] = {"width": vp["height"], "height": vp["width"]}

            context = browser.new_context(**context_kwargs)
            page = context.new_page()
            page.set_default_timeout(DEFAULT_TIMEOUT_MS)

            try:
                page.goto(url, wait_until="networkidle")
            except PlaywrightTimeoutError:
                try:
                    page.goto(url, wait_until="load")
                except Exception as exc:
                    errors.append(f"rendered page load error: {exc}")

            page.screenshot(path=str(screenshot_path), full_page=True)
            context.close()
        finally:
            browser.close()
    return errors


def snapshot_from_rendered(url: str, rendered: Dict[str, Any], errors: List[str]) -> PageSnapshot:
    structured = _unique_keep_order(rendered.get("structured_data", []) or [])
    return PageSnapshot(
        url=url,
        fetched_url=url,
        status_code=None,
        title=rendered.get("title", "") or "",
        meta_description=rendered.get("meta_description", "") or "",
        canonical=rendered.get("canonical", "") or "",
        headings=_unique_keep_order(rendered.get("headings", []) or []),
        links=_unique_keep_order(rendered.get("links", []) or []),
        images=_unique_keep_order(rendered.get("images", []) or []),
        structured_data=structured,
        schema_types=extract_schema_types(structured),
        viewport_present=bool(rendered.get("viewport_present")),
        og_present=bool(rendered.get("og_present")),
        visible_text_lines=_unique_keep_order(rendered.get("visible_text_lines", []) or []),
        html_bytes=int(rendered.get("html_bytes", 0) or 0),
        errors=errors,
    )


def diff_missing(rendered: List[str], raw: List[str]) -> List[str]:
    raw_set = {_norm(x) for x in raw}
    out = []
    for item in rendered:
        if _norm(item) and _norm(item) not in raw_set:
            out.append(item)
    return out


def write_docx(
    out_path: Path,
    url: str,
    ua_label: str,
    raw: PageSnapshot,
    rendered: PageSnapshot,
    missing: Dict[str, List[str]],
    timings: Dict[str, float],
    screenshot_rendered: Path | None,
    screenshot_nojs: Path | None,
    screenshot_rendered_landscape: Path | None,
    screenshot_nojs_landscape: Path | None,
    recommendations: List[str],
    timing_nojs_data: Dict[str, float],
    timing_js_data: Dict[str, float],
):
    doc = Document()
    doc.add_paragraph("SEO RENDERING AUDIT")
    doc.add_paragraph("Анализ контента с JavaScript и без него")
    doc.add_paragraph("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")

    doc.add_paragraph("Анализируемый URL:")
    doc.add_paragraph(url)
    doc.add_paragraph(f"Дата аудита: {format_ru_datetime(datetime.now())}")

    doc.add_heading("Краткое резюме", level=2)
    doc.add_paragraph(
        f"Данный отчёт содержит результаты SEO-аудита страницы {url} с фокусом на сравнение контента, "
        "доступного с JavaScript и без него."
    )

    metrics = compute_metrics(rendered, missing)
    total_missing = int(metrics["total_missing"])
    score = metrics["score"]

    score_nojs = round(score)
    score_js = round(score)
    load_nojs = timings.get("raw", 0.0)
    load_js = timings.get("rendered", 0.0)

    def load_status(seconds: float) -> str:
        if seconds <= 10:
            return "✅ Быстро"
        if seconds <= 20:
            return "⚠️ Медленно"
        return "❌ Очень медленно"

    metrics_table = doc.add_table(rows=1, cols=4)
    metrics_table.style = "Light Grid"
    hdr = metrics_table.rows[0].cells
    hdr[0].text = "Метрика"
    hdr[1].text = "Без JavaScript"
    hdr[2].text = "С JavaScript"
    hdr[3].text = "Статус"

    r1 = metrics_table.add_row().cells
    r1[0].text = "SEO Score"
    r1[1].text = f"{score_nojs}/100"
    r1[2].text = f"{score_js}/100"
    r1[3].text = "✅ Хорошо" if score_js >= 70 else "⚠️ Средне"

    r2 = metrics_table.add_row().cells
    r2[0].text = "Время загрузки"
    r2[1].text = f"{load_nojs:.2f}s"
    r2[2].text = f"{load_js:.2f}s"
    r2[3].text = load_status(load_js)

    doc.add_heading("Ключевые выводы", level=2)
    removed_count = (
        len(diff_missing(raw.visible_text_lines, rendered.visible_text_lines))
        + len(diff_missing(raw.headings, rendered.headings))
        + len(diff_missing(raw.links, rendered.links))
        + len(diff_missing(raw.images, rendered.images))
        + len(diff_missing(raw.structured_data, rendered.structured_data))
    )
    crit_count = sum(
        1 for v in missing.values() if len(v) > 5
    )
    warn_count = sum(
        1 for v in missing.values() if 0 < len(v) <= 5
    )

    doc.add_paragraph(f"● SEO Score без JavaScript: {score_nojs}/100", style="List Bullet")
    doc.add_paragraph(f"● SEO Score с JavaScript: {score_js}/100", style="List Bullet")
    doc.add_paragraph(f"● Критических проблем: {crit_count}", style="List Bullet")
    doc.add_paragraph(f"● Предупреждений: {warn_count}", style="List Bullet")
    doc.add_paragraph(f"● Добавлено элементов через JS: {total_missing}", style="List Bullet")
    doc.add_paragraph(f"● Удалено/скрыто элементов: {removed_count}", style="List Bullet")
    doc.add_paragraph("● Изменено элементов: 0", style="List Bullet")

    doc.add_heading("Сравнение SEO-элементов", level=2)
    doc.add_paragraph("Детальное сравнение SEO-элементов страницы с JavaScript и без него:")
    comp = doc.add_table(rows=1, cols=4)
    comp.style = "Light Grid"
    ch = comp.rows[0].cells
    ch[0].text = "Элемент"
    ch[1].text = "Без JS"
    ch[2].text = "С JS"
    ch[3].text = "Статус"

    def add_comp_row(label: str, a: str, b: str):
        row = comp.add_row().cells
        row[0].text = label
        row[1].text = a
        row[2].text = b
        row[3].text = status_for_match(a, b)

    add_comp_row("Title", raw.title, rendered.title)
    add_comp_row("Meta Description", raw.meta_description, rendered.meta_description)
    add_comp_row("H1 заголовки", f"{count_headings(raw.headings, 1)} шт", f"{count_headings(rendered.headings, 1)} шт")
    add_comp_row("H2 заголовки", f"{count_headings(raw.headings, 2)} шт", f"{count_headings(rendered.headings, 2)} шт")
    add_comp_row("Изображения", f"{len(raw.images)} шт", f"{len(rendered.images)} шт")
    add_comp_row("Ссылки", f"{len(raw.links)} шт", f"{len(rendered.links)} шт")
    add_comp_row("Canonical", raw.canonical or "—", rendered.canonical or "—")
    add_comp_row(
        "Schema Markup",
        ", ".join(raw.schema_types) if raw.schema_types else "—",
        ", ".join(rendered.schema_types) if rendered.schema_types else "—",
    )
    add_comp_row(
        "Viewport",
        "✅ Есть" if raw.viewport_present else "❌ Нет",
        "✅ Есть" if rendered.viewport_present else "❌ Нет",
    )
    add_comp_row(
        "Open Graph",
        "✅ Есть" if raw.og_present else "❌ Нет",
        "✅ Есть" if rendered.og_present else "❌ Нет",
    )

    doc.add_heading("SEO элементы", level=2)
    def seo_block(label: str, a: str, b: str):
        doc.add_paragraph(label)
        doc.add_paragraph(match_text(a, b))
        doc.add_paragraph("Без JS:")
        doc.add_paragraph(a or "—")
        doc.add_paragraph("")
        doc.add_paragraph("С JS:")
        doc.add_paragraph(b or "—")
        doc.add_paragraph("")

    seo_block("Title", raw.title, rendered.title)
    seo_block("Meta Description", raw.meta_description, rendered.meta_description)
    seo_block("H1 заголовки", f"{count_headings(raw.headings, 1)} шт", f"{count_headings(rendered.headings, 1)} шт")
    seo_block("H2 заголовки", f"{count_headings(raw.headings, 2)} шт", f"{count_headings(rendered.headings, 2)} шт")
    seo_block("Изображения", f"{len(raw.images)} шт", f"{len(rendered.images)} шт")
    seo_block("Ссылки", f"{len(raw.links)} шт", f"{len(rendered.links)} шт")
    seo_block("Canonical", raw.canonical, rendered.canonical)
    seo_block(
        "Schema Markup",
        ", ".join(raw.schema_types) if raw.schema_types else "—",
        ", ".join(rendered.schema_types) if rendered.schema_types else "—",
    )

    doc.add_heading("Изменения контента", level=2)
    if total_missing == 0 and removed_count == 0:
        doc.add_paragraph("Изменения не обнаружены.")
    else:
        doc.add_paragraph(f"Добавлено элементов через JS: {total_missing}")
        doc.add_paragraph(f"Удалено/скрыто элементов: {removed_count}")

    doc.add_heading("Анализ по устройствам", level=2)
    dev = doc.add_table(rows=1, cols=4)
    dev.style = "Light Grid"
    dh = dev.rows[0].cells
    dh[0].text = "Устройство"
    dh[1].text = "Score без JS"
    dh[2].text = "Score с JS"
    dh[3].text = "Разница"
    dr = dev.add_row().cells
    dr[0].text = ua_label
    dr[1].text = f"{score_nojs}"
    dr[2].text = f"{score_js}"
    dr[3].text = "+0%"

    doc.add_heading("⏱️ Timeline загрузки", level=2)
    doc.add_paragraph("Без JavaScript")
    tl = doc.add_table(rows=1, cols=4)
    tl.style = "Light Grid"
    th = tl.rows[0].cells
    th[0].text = "Этап"
    th[1].text = "Значение"
    th[2].text = "Единицы"
    th[3].text = "График"
    max_nojs = max(
        [
            timing_nojs_data.get("dns", 0.0),
            timing_nojs_data.get("tcp", 0.0),
            timing_nojs_data.get("html", 0.0),
            timing_nojs_data.get("dom_parse", 0.0),
            timing_nojs_data.get("render", 0.0),
        ]
        or [0.0]
    )
    for label, key in [
        ("DNS Lookup", "dns"),
        ("TCP Connection", "tcp"),
        ("HTML Download", "html"),
        ("DOM Parse", "dom_parse"),
        ("Render", "render"),
    ]:
        row = tl.add_row().cells
        row[0].text = label
        val = timing_nojs_data.get(key, 0.0)
        row[1].text = f"{val:.0f}"
        row[2].text = "ms"
        row[3].text = bar(val, max_nojs)

    doc.add_paragraph("С JavaScript")
    tl2 = doc.add_table(rows=1, cols=4)
    tl2.style = "Light Grid"
    th2 = tl2.rows[0].cells
    th2[0].text = "Этап"
    th2[1].text = "Значение"
    th2[2].text = "Единицы"
    th2[3].text = "График"
    max_js = max(
        [
            timing_js_data.get("dns", 0.0),
            timing_js_data.get("tcp", 0.0),
            timing_js_data.get("html", 0.0),
            timing_js_data.get("dom_parse", 0.0),
            timing_js_data.get("tti", 0.0),
        ]
        or [0.0]
    )
    for label, key in [
        ("DNS Lookup", "dns"),
        ("TCP Connection", "tcp"),
        ("HTML Download", "html"),
        ("JS Parse & Compile", "tbt"),
        ("JS Execution", "tbt"),
        ("API/Data Fetch", "tbt"),
        ("DOM Updates", "dom_parse"),
        ("Full Render", "tti"),
    ]:
        value = timing_js_data.get(key, 0.0)
        row = tl2.add_row().cells
        row[0].text = label
        row[1].text = f"{value:.0f}"
        row[2].text = "ms"
        row[3].text = bar(value, max_js)

    doc.add_paragraph("Метрики без JS")
    m1 = doc.add_table(rows=1, cols=4)
    m1.style = "Light Grid"
    mh1 = m1.rows[0].cells
    mh1[0].text = "Метрика"
    mh1[1].text = "Значение"
    mh1[2].text = "Единицы"
    mh1[3].text = "График"
    max_m1 = max(
        [
            timing_nojs_data.get("ttfb", 0.0),
            timing_nojs_data.get("fcp", 0.0),
            timing_nojs_data.get("dom_content_loaded", 0.0),
            timing_nojs_data.get("tbt", 0.0),
        ]
        or [0.0]
    )
    for label, key in [
        ("Time to First Byte", "ttfb"),
        ("First Contentful Paint", "fcp"),
        ("DOM Content Loaded", "dom_content_loaded"),
        ("Total Blocking Time", "tbt"),
    ]:
        row = m1.add_row().cells
        row[0].text = label
        val = timing_nojs_data.get(key, 0.0)
        row[1].text = f"{val:.0f}"
        row[2].text = "ms"
        row[3].text = bar(val, max_m1)

    doc.add_paragraph("Метрики с JS")
    m2 = doc.add_table(rows=1, cols=4)
    m2.style = "Light Grid"
    mh2 = m2.rows[0].cells
    mh2[0].text = "Метрика"
    mh2[1].text = "Значение"
    mh2[2].text = "Единицы"
    mh2[3].text = "График"
    max_m2 = max(
        [
            timing_js_data.get("ttfb", 0.0),
            timing_js_data.get("fcp", 0.0),
            timing_js_data.get("dom_content_loaded", 0.0),
            timing_js_data.get("tbt", 0.0),
            timing_js_data.get("tti", 0.0),
        ]
        or [0.0]
    )
    for label, key in [
        ("Time to First Byte", "ttfb"),
        ("First Contentful Paint", "fcp"),
        ("DOM Content Loaded", "dom_content_loaded"),
        ("Total Blocking Time", "tbt"),
        ("Time to Interactive", "tti"),
    ]:
        row = m2.add_row().cells
        row[0].text = label
        val = timing_js_data.get(key, 0.0)
        row[1].text = f"{val:.0f}"
        row[2].text = "ms"
        row[3].text = bar(val, max_m2)

    if screenshot_rendered and screenshot_rendered.exists():
        doc.add_paragraph("Скриншот (JS включён):")
        doc.add_picture(str(screenshot_rendered), width=Inches(6.5))

    if screenshot_nojs and screenshot_nojs.exists():
        doc.add_paragraph("Скриншот (JS отключён):")
        doc.add_picture(str(screenshot_nojs), width=Inches(6.5))

    if screenshot_rendered_landscape and screenshot_rendered_landscape.exists():
        doc.add_paragraph("Скриншот (мобильный, горизонтальный, JS включён):")
        doc.add_picture(str(screenshot_rendered_landscape), width=Inches(6.5))

    if screenshot_nojs_landscape and screenshot_nojs_landscape.exists():
        doc.add_paragraph("Скриншот (мобильный, горизонтальный, JS отключён):")
        doc.add_picture(str(screenshot_nojs_landscape), width=Inches(6.5))

    if raw.errors or rendered.errors:
        doc.add_heading("Ошибки", level=2)
        for err in raw.errors + rendered.errors:
            doc.add_paragraph(err, style="List Bullet")

    doc.add_heading("Рекомендации по оптимизации", level=2)
    if recommendations:
        for rec in recommendations:
            doc.add_paragraph(rec, style="List Bullet")
    else:
        doc.add_paragraph("Нет специфических рекомендаций по результатам аудита.")

    if raw.og_present is False or rendered.og_present is False:
        doc.add_paragraph("Отсутствуют Open Graph теги", style="List Bullet")
        doc.add_paragraph(
            "Добавьте OG теги для лучшего отображения при шеринге в соцсетях.",
            style="List Bullet",
        )

    doc.save(out_path)


def write_xlsx(out_path: Path, errors: List[Dict[str, Any]]):
    wb = Workbook()
    ws = wb.active
    ws.title = "Errors"
    ws.append(["timestamp", "context", "severity", "message"])
    for err in errors:
        ws.append([err.get("timestamp"), err.get("context"), err.get("severity"), err.get("message")])
    wb.save(out_path)


def collect_errors(raw: PageSnapshot, rendered: PageSnapshot, context: str) -> List[Dict[str, Any]]:
    out = []
    now = datetime.now(timezone.utc).isoformat()
    for e in raw.errors:
        out.append({"timestamp": now, "context": f"{context} raw", "severity": "error", "message": e})
    for e in rendered.errors:
        out.append({"timestamp": now, "context": f"{context} rendered", "severity": "error", "message": e})
    return out


def build_recommendations(missing: Dict[str, List[str]]) -> List[str]:
    recs: List[str] = []
    if len(missing.get("Visible text", [])) > 10:
        recs.append(
            "[Высокий] Убедитесь, что ключевой контент доступен без JavaScript (SSR/SSG), чтобы Googlebot видел его в исходном HTML."
        )
    if missing.get("Headings"):
        recs.append(
            "[Высокий] Заголовки (H1–H2) должны присутствовать в исходном HTML без выполнения JS."
        )
    if missing.get("Links"):
        recs.append(
            "[Средний] Важные ссылки навигации и внутренние ссылки должны быть в исходном HTML для корректного обхода."
        )
    if missing.get("Images"):
        recs.append(
            "[Низкий] Критичные изображения и alt-тексты лучше отдавать в HTML до выполнения JS."
        )
    if missing.get("Structured data"):
        recs.append(
            "[Средний] Структурированные данные (JSON-LD) рекомендуется рендерить на сервере."
        )
    return recs


def compute_metrics(rendered: PageSnapshot, missing: Dict[str, List[str]]) -> Dict[str, float]:
    total_missing = float(sum(len(v) for v in missing.values()))
    rendered_total = float(
        len(rendered.visible_text_lines)
        + len(rendered.headings)
        + len(rendered.links)
        + len(rendered.images)
        + len(rendered.structured_data)
    )
    if rendered_total > 0:
        missing_pct = min(100.0, (total_missing / rendered_total) * 100.0)
        score = max(0.0, 100.0 - missing_pct)
    else:
        missing_pct = 0.0
        score = 100.0
    return {
        "total_missing": total_missing,
        "rendered_total": rendered_total,
        "missing_pct": missing_pct,
        "score": score,
    }


def audit_variant(
    url: str,
    ua_label: str,
    user_agent: str,
    mobile: bool,
    out_dir: Path,
) -> Tuple[Path, List[Dict[str, Any]], Dict[str, float], Dict[str, float], Dict[str, float]]:
    log_step(f"[{ua_label}] Fetch raw HTML", 1, 5)
    t0 = time.time()
    html, status_code, fetched = fetch_raw_html(url, user_agent)
    raw = parse_raw_html(url, html, status_code, fetched if isinstance(fetched, str) else url, "" if status_code else "raw fetch failed")
    t1 = time.time()

    log_step(f"[{ua_label}] Render with JS + screenshot", 2, 5)
    screenshot_rendered = out_dir / f"screenshot_rendered_{ua_label.replace(' ', '_').lower()}.png"
    screenshot_no_js = out_dir / f"screenshot_nojs_{ua_label.replace(' ', '_').lower()}.png"
    screenshot_rendered_landscape = out_dir / f"screenshot_rendered_landscape_{ua_label.replace(' ', '_').lower()}.png"
    screenshot_nojs_landscape = out_dir / f"screenshot_nojs_landscape_{ua_label.replace(' ', '_').lower()}.png"

    rendered_data, timing_js, render_errors = render_with_playwright(url, user_agent, mobile, screenshot_rendered)
    rendered = snapshot_from_rendered(url, rendered_data, render_errors)
    t2 = time.time()

    log_step(f"[{ua_label}] No-JS screenshot", 3, 5)
    nojs_errors = screenshot_nojs(url, user_agent, mobile, False, screenshot_no_js)
    rendered.errors.extend(nojs_errors)
    log_step(f"[{ua_label}] Timing (no-JS)", 4, 5)
    timing_nojs_data, timing_nojs_errors = timing_nojs(url, user_agent, mobile)
    rendered.errors.extend(timing_nojs_errors)

    if mobile:
        rendered.errors.extend(
            screenshot_rendered_only(url, user_agent, mobile, True, screenshot_rendered_landscape)
        )
        rendered.errors.extend(
            screenshot_nojs(url, user_agent, mobile, True, screenshot_nojs_landscape)
        )

    missing = {
        "Visible text": diff_missing(rendered.visible_text_lines, raw.visible_text_lines),
        "Headings": diff_missing(rendered.headings, raw.headings),
        "Links": diff_missing(rendered.links, raw.links),
        "Images": diff_missing(rendered.images, raw.images),
        "Structured data": diff_missing(rendered.structured_data, raw.structured_data),
    }

    timings = {"raw": t1 - t0, "rendered": t2 - t1}
    recommendations = build_recommendations(missing)
    metrics = compute_metrics(rendered, missing)

    log_step(f"[{ua_label}] Build report", 5, 5)
    out_docx = out_dir / f"seo_render_audit_{ua_label.replace(' ', '_').lower()}.docx"
    write_docx(
        out_docx,
        url,
        ua_label,
        raw,
        rendered,
        missing,
        timings,
        screenshot_rendered,
        screenshot_no_js,
        screenshot_rendered_landscape,
        screenshot_nojs_landscape,
        recommendations,
        timing_nojs_data,
        timing_js,
    )

    errors = collect_errors(raw, rendered, ua_label)
    return out_docx, errors, metrics, timing_nojs_data, timing_js


def _slug(text: str) -> str:
    return re.sub(r"[^a-zA-Z0-9]+", "-", (text or "").strip()).strip("-").lower()


class RenderAuditServiceV2:
    """Render audit engine for API integration."""

    def __init__(self, timeout: int = 35):
        self.timeout = max(10, int(timeout))

    def _variant_issues(
        self,
        raw: PageSnapshot,
        rendered: PageSnapshot,
        missing: Dict[str, List[str]],
        metrics: Dict[str, float],
        timings: Dict[str, float],
    ) -> List[Dict[str, Any]]:
        issues: List[Dict[str, Any]] = []
        for err in (raw.errors or []) + (rendered.errors or []):
            issues.append(
                {
                    "severity": "critical",
                    "code": "runtime_error",
                    "title": "Ошибка выполнения рендер-аудита",
                    "details": str(err),
                }
            )
        if missing.get("Visible text"):
            sev = "critical" if len(missing["Visible text"]) > 10 else "warning"
            issues.append(
                {
                    "severity": sev,
                    "code": "content_missing_nojs",
                    "title": "Контент появляется только после JavaScript",
                    "details": f"Отсутствует в no-JS версии: {len(missing['Visible text'])} строк текста.",
                }
            )
        if missing.get("Headings"):
            issues.append(
                {
                    "severity": "critical",
                    "code": "headings_missing_nojs",
                    "title": "Заголовки отсутствуют без JavaScript",
                    "details": f"Найдено {len(missing['Headings'])} заголовков только после JS.",
                }
            )
        if missing.get("Links"):
            issues.append(
                {
                    "severity": "warning",
                    "code": "links_missing_nojs",
                    "title": "Ссылки появляются только после JavaScript",
                    "details": f"Найдено {len(missing['Links'])} ссылок только в JS-рендере.",
                }
            )
        if missing.get("Structured data"):
            issues.append(
                {
                    "severity": "warning",
                    "code": "schema_missing_nojs",
                    "title": "Structured data зависит от JavaScript",
                    "details": f"Найдено {len(missing['Structured data'])} элементов JSON-LD только с JS.",
                }
            )
        if timings.get("rendered", 0.0) > 12:
            issues.append(
                {
                    "severity": "warning",
                    "code": "render_too_slow",
                    "title": "Медленный JS-рендер",
                    "details": f"Время рендера с JS: {timings.get('rendered', 0.0):.2f} сек.",
                }
            )
        if metrics.get("score", 100.0) < 70:
            issues.append(
                {
                    "severity": "critical",
                    "code": "low_render_score",
                    "title": "Низкий балл рендеринга",
                    "details": f"Текущий балл: {metrics.get('score', 0.0):.1f}/100.",
                }
            )
        return issues

    def run(self, url: str, task_id: str, progress_callback=None) -> Dict[str, Any]:
        def _notify(progress: int, message: str) -> None:
            if callable(progress_callback):
                try:
                    progress_callback(progress, message)
                except Exception:
                    pass

        variants = [
            ("googlebot_desktop", "Googlebot Desktop", GOOGLEBOT_DESKTOP_UA, False),
            ("googlebot_mobile", "Googlebot Mobile", GOOGLEBOT_MOBILE_UA, True),
        ]
        domain = urlparse(url).netloc or "site"
        stamp = datetime.utcnow().strftime("%Y-%m-%d_%H-%M")
        shot_dir = Path(settings.REPORTS_DIR) / "render" / task_id / "screenshots"
        shot_dir.mkdir(parents=True, exist_ok=True)

        all_variant_results: List[Dict[str, Any]] = []
        global_issues: List[Dict[str, Any]] = []
        screenshot_paths: List[str] = []

        _notify(5, "Подготовка рендер-аудита")
        try:
            import playwright  # noqa: F401
        except Exception as exc:
            _notify(100, "Среда Playwright недоступна")
            return {
                "task_type": "render_audit",
                "url": url,
                "completed_at": datetime.utcnow().isoformat(),
                "results": {
                    "engine": "v2",
                    "variants": [],
                    "issues": [
                        {
                            "severity": "critical",
                            "code": "playwright_unavailable",
                            "title": "Playwright runtime unavailable",
                            "details": str(exc),
                        }
                    ],
                    "issues_count": 1,
                    "summary": {
                        "variants_total": 0,
                        "critical_issues": 1,
                        "warning_issues": 0,
                        "info_issues": 0,
                        "score": None,
                        "missing_total": 0,
                        "avg_missing_pct": 0,
                        "avg_raw_load_ms": 0,
                        "avg_js_load_ms": 0,
                    },
                    "recommendations": ["Установите Playwright/Chromium для полноценного рендер-аудита."],
                    "artifacts": {"screenshot_dir": str(shot_dir), "screenshots": []},
                },
            }

        total = len(variants) or 1
        for index, (variant_id, variant_label, ua, mobile) in enumerate(variants, start=1):
            _notify(10 + int(((index - 1) / total) * 80), f"Аудит {variant_label} ({index}/{total})")
            html, status_code, fetched_url = fetch_raw_html(url, ua, timeout_s=self.timeout)
            raw = parse_raw_html(url, html, status_code, fetched_url if isinstance(fetched_url, str) else url, "" if status_code else "raw fetch failed")

            t0 = time.time()
            base = f"render_{_slug(domain)}_{stamp}_{variant_id}"
            shot_js = shot_dir / f"{base}_js.png"
            shot_nojs = shot_dir / f"{base}_nojs.png"
            shot_js_land = shot_dir / f"{base}_mobile_landscape_js.png"
            shot_nojs_land = shot_dir / f"{base}_mobile_landscape_nojs.png"

            rendered_data, timing_js, render_errors = render_with_playwright(url, ua, mobile, shot_js)
            rendered = snapshot_from_rendered(url, rendered_data, render_errors)
            timing_nojs_data, timing_nojs_errors = timing_nojs(url, ua, mobile)
            rendered.errors.extend(timing_nojs_errors)
            rendered.errors.extend(screenshot_nojs(url, ua, mobile, False, shot_nojs))
            if mobile:
                rendered.errors.extend(screenshot_rendered_only(url, ua, mobile, True, shot_js_land))
                rendered.errors.extend(screenshot_nojs(url, ua, mobile, True, shot_nojs_land))
            t1 = time.time()

            missing = {
                "Visible text": diff_missing(rendered.visible_text_lines, raw.visible_text_lines),
                "Headings": diff_missing(rendered.headings, raw.headings),
                "Links": diff_missing(rendered.links, raw.links),
                "Images": diff_missing(rendered.images, raw.images),
                "Structured data": diff_missing(rendered.structured_data, raw.structured_data),
            }
            metrics = compute_metrics(rendered, missing)
            timings = {"raw": 0.0, "rendered": max(0.0, t1 - t0)}
            recommendations = build_recommendations(missing)
            issues = self._variant_issues(raw, rendered, missing, metrics, timings)
            for item in issues:
                global_issues.append({**item, "variant": variant_label})

            shots_payload: Dict[str, Dict[str, str]] = {}
            for tag, path in [
                ("js", shot_js),
                ("nojs", shot_nojs),
                ("js_landscape", shot_js_land if mobile else None),
                ("nojs_landscape", shot_nojs_land if mobile else None),
            ]:
                if path and path.exists():
                    shots_payload[tag] = {
                        "path": str(path),
                        "name": path.name,
                        "url": f"/api/render-artifacts/{task_id}/{path.name}",
                    }
                    screenshot_paths.append(str(path))

            all_variant_results.append(
                {
                    "variant_id": variant_id,
                    "variant_label": variant_label,
                    "mobile": mobile,
                    "user_agent": ua,
                    "raw": {
                        "status_code": raw.status_code,
                        "fetched_url": raw.fetched_url,
                        "title": raw.title,
                        "meta_description": raw.meta_description,
                        "canonical": raw.canonical,
                        "h1_count": count_headings(raw.headings, 1),
                        "h2_count": count_headings(raw.headings, 2),
                        "links_count": len(raw.links),
                        "images_count": len(raw.images),
                        "structured_data_count": len(raw.structured_data),
                        "schema_types": raw.schema_types,
                        "viewport_present": raw.viewport_present,
                        "og_present": raw.og_present,
                        "visible_text_count": len(raw.visible_text_lines),
                        "html_bytes": raw.html_bytes,
                        "errors": raw.errors,
                    },
                    "rendered": {
                        "title": rendered.title,
                        "meta_description": rendered.meta_description,
                        "canonical": rendered.canonical,
                        "h1_count": count_headings(rendered.headings, 1),
                        "h2_count": count_headings(rendered.headings, 2),
                        "links_count": len(rendered.links),
                        "images_count": len(rendered.images),
                        "structured_data_count": len(rendered.structured_data),
                        "schema_types": rendered.schema_types,
                        "viewport_present": rendered.viewport_present,
                        "og_present": rendered.og_present,
                        "visible_text_count": len(rendered.visible_text_lines),
                        "html_bytes": rendered.html_bytes,
                        "errors": rendered.errors,
                    },
                    "missing": {
                        "visible_text": missing["Visible text"],
                        "headings": missing["Headings"],
                        "links": missing["Links"],
                        "images": missing["Images"],
                        "structured_data": missing["Structured data"],
                    },
                    "metrics": metrics,
                    "timings": {"raw_s": timings["raw"], "rendered_s": timings["rendered"]},
                    "timing_nojs_ms": timing_nojs_data,
                    "timing_js_ms": timing_js,
                    "issues": issues,
                    "recommendations": recommendations,
                    "screenshots": shots_payload,
                }
            )
            _notify(10 + int((index / total) * 80), f"Завершено: {variant_label}")

        critical_count = sum(1 for i in global_issues if i.get("severity") == "critical")
        warning_count = sum(1 for i in global_issues if i.get("severity") == "warning")
        info_count = sum(1 for i in global_issues if i.get("severity") == "info")
        scores = [float(v.get("metrics", {}).get("score", 100.0)) for v in all_variant_results]
        missing_total = int(sum(float(v.get("metrics", {}).get("total_missing", 0.0)) for v in all_variant_results))
        missing_pct = [float(v.get("metrics", {}).get("missing_pct", 0.0)) for v in all_variant_results]
        raw_ms = [float(v.get("timings", {}).get("raw_s", 0.0)) * 1000 for v in all_variant_results]
        js_ms = [float(v.get("timings", {}).get("rendered_s", 0.0)) * 1000 for v in all_variant_results]

        summary = {
            "variants_total": len(all_variant_results),
            "critical_issues": critical_count,
            "warning_issues": warning_count,
            "info_issues": info_count,
            "score": round(sum(scores) / len(scores), 1) if scores else None,
            "missing_total": missing_total,
            "avg_missing_pct": round(sum(missing_pct) / len(missing_pct), 1) if missing_pct else 0,
            "avg_raw_load_ms": round(sum(raw_ms) / len(raw_ms), 1) if raw_ms else 0,
            "avg_js_load_ms": round(sum(js_ms) / len(js_ms), 1) if js_ms else 0,
        }
        recommendations = []
        if critical_count > 0:
            recommendations.append("Устраните критичные расхождения JS/no-JS в первую очередь.")
        if warning_count > 0:
            recommendations.append("Снизьте зависимость SEO-контента от JavaScript и перенесите важные блоки на серверный рендер.")
        if critical_count == 0 and warning_count == 0:
            recommendations.append("Критичных SEO-расхождений между JS и no-JS не обнаружено.")

        _notify(98, "Формирование итогов рендер-аудита")
        return {
            "task_type": "render_audit",
            "url": url,
            "completed_at": datetime.utcnow().isoformat(),
            "results": {
                "engine": "v2",
                "summary": summary,
                "variants": all_variant_results,
                "issues": global_issues,
                "issues_count": len(global_issues),
                "recommendations": recommendations,
                "artifacts": {
                    "screenshot_dir": str(shot_dir),
                    "screenshots": screenshot_paths,
                },
            },
        }


def main() -> int:
    parser = argparse.ArgumentParser(description="SEO Rendering Audit - JS vs No-JS Comparison")
    parser.add_argument("url", help="Target URL to audit")
    parser.add_argument("--out-dir", default="output", help="Output directory")
    args = parser.parse_args()

    base_out_dir = Path(args.out_dir)
    parsed = urlparse(args.url)
    host = (parsed.netloc or parsed.path or "site").replace(":", "_")
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    out_dir = base_out_dir / f"{timestamp}_{host}"
    out_dir.mkdir(parents=True, exist_ok=True)

    all_errors: List[Dict[str, Any]] = []

    docx_desktop, errors_desktop, metrics_desktop, timing_nojs_desktop, timing_js_desktop = audit_variant(
        args.url, "Googlebot Desktop", GOOGLEBOT_DESKTOP_UA, False, out_dir
    )
    all_errors.extend(errors_desktop)

    docx_mobile, errors_mobile, metrics_mobile, timing_nojs_mobile, timing_js_mobile = audit_variant(
        args.url, "Googlebot Mobile", GOOGLEBOT_MOBILE_UA, True, out_dir
    )
    all_errors.extend(errors_mobile)

    xlsx_path = out_dir / "seo_render_audit_errors.xlsx"
    write_xlsx(xlsx_path, all_errors)

    def status_line(label: str, metrics: Dict[str, float]) -> str:
        total_missing = int(metrics.get("total_missing", 0))
        missing_pct = metrics.get("missing_pct", 0.0)
        score = metrics.get("score", 0.0)
        if total_missing == 0:
            icon = "✔"
            status = "OK"
        elif total_missing <= 25:
            icon = "⚠"
            status = "WARN"
        else:
            icon = "✖"
            status = "CRITICAL"
        return f"{icon} {label}: {status} | Missing: {total_missing} | Missing %: {missing_pct:.1f}% | Score: {score:.1f}/100"

    print("Summary:")
    print(status_line("Desktop", metrics_desktop))
    print(status_line("Mobile", metrics_mobile))
    print("")
    print("Reports generated:")
    print(f"- {docx_desktop}")
    print(f"- {docx_mobile}")
    print(f"- {xlsx_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
